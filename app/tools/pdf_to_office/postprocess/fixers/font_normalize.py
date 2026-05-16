"""字型正規化 fixer。

修正 pdf2docx 帶過來的 PDF 字型名稱（含 subset prefix）→ 系統可用名稱，
且 CJK 字型套到 w:eastAsia，避免「中文段落套到 Times New Roman」變方框字。
有 PDFTruth alignment 時用真值校正。
"""
from __future__ import annotations

import logging
from copy import deepcopy

from docx.oxml.ns import qn

from ...pdf_truth.aligner import DocxToPdfAlignment
from ..config import (
    CJK_FALLBACK_BY_LANG, FALLBACK_ASCII_FONT, FALLBACK_CJK_FONT, FONT_MAPPING,
    MONOSPACE_FALLBACK, MONOSPACE_HINTS,
)

log = logging.getLogger(__name__)


def _strip_subset_prefix(name: str) -> str:
    """`BAAAAA+PingFangTC-Regular` → `PingFangTC-Regular`。"""
    if not name:
        return ""
    if "+" in name and name.split("+", 1)[0].isupper() and len(name.split("+", 1)[0]) == 6:
        return name.split("+", 1)[1]
    return name


def _is_monospace(name: str) -> bool:
    """偵測字型名稱看起來是不是等寬。"""
    base = _strip_subset_prefix(name) or ""
    bl = base.lower()
    return any(h.lower() in bl for h in MONOSPACE_HINTS)


def _resolve_font(pdf_font_name: str) -> tuple[str, str]:
    """PDF 字型名 → (eastAsia 字型, ASCII 字型)。

    優先順序：
    1. 等寬字型（Courier / Mono / Consolas / Menlo 等）→ Courier New，避免 code 變
       proportional 對不齊
    2. FONT_MAPPING 完整匹配
    3. prefix 匹配（PingFangTC-Regular / PingFangTC-Semibold 都對 PingFangTC）
    4. fallback 新細明體 + Times New Roman
    """
    base = _strip_subset_prefix(pdf_font_name)
    if _is_monospace(base):
        return MONOSPACE_FALLBACK
    if base in FONT_MAPPING:
        return FONT_MAPPING[base]
    for key, value in FONT_MAPPING.items():
        if base.startswith(key + "-") or base.startswith(key):
            return value
    return (FALLBACK_CJK_FONT, FALLBACK_ASCII_FONT)


def _set_run_fonts(run, eastasia: str, ascii_font: str) -> None:
    """在 docx run 上同時設 ascii / hAnsi / eastAsia 字型。
    python-docx 高層 API 只能設一個 .name；要同時設多個 region 得直接動 XML。"""
    rPr = run._element.get_or_add_rPr()
    # 移除舊 rFonts，重建
    for old in rPr.findall(qn("w:rFonts")):
        rPr.remove(old)
    rFonts = deepcopy(rPr.makeelement(qn("w:rFonts"), {}))
    rFonts.set(qn("w:ascii"), ascii_font)
    rFonts.set(qn("w:hAnsi"), ascii_font)
    rFonts.set(qn("w:eastAsia"), eastasia)
    rFonts.set(qn("w:cs"), ascii_font)
    rPr.insert(0, rFonts)


def _walk_paragraphs(doc):
    for p in doc.paragraphs:
        yield p
    for tbl in doc.tables:
        for row in tbl.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    yield p


def fix_font_normalize(docx_doc, pdf_truth, alignment: DocxToPdfAlignment) -> dict:
    """套用字型正規化。回 changelog dict。"""
    al_by_di = {a.docx_para_index: a for a in alignment.alignments}
    changes = 0
    pdf_truth_used = 0
    # 依 PDF 文種覆寫 CJK fallback（v1.8.57+）— 簡中 / 日文 / 韓文 PDF 別用繁中字型
    lang_specific_cjk = None
    if pdf_truth and getattr(pdf_truth, "language_guess", None):
        lang_specific_cjk = CJK_FALLBACK_BY_LANG.get(pdf_truth.language_guess)

    for di, p in enumerate(_walk_paragraphs(docx_doc)):
        a = al_by_di.get(di)
        if a and a.pdf_dominant_font:
            pdf_font = a.pdf_dominant_font
            pdf_truth_used += 1
        else:
            # 沒對應到 PDF block — 用 docx 自己的 run 字型決定
            run_fonts = [r.font.name for r in p.runs if r.font.name]
            pdf_font = run_fonts[0] if run_fonts else ""

        if not pdf_font:
            continue
        eastasia, ascii_font = _resolve_font(pdf_font)
        # 若 PDF 文種跟 fallback 對不上（例如 SC PDF 卻 fallback 給繁中新細明體），
        # 用 lang-specific CJK 蓋掉
        if lang_specific_cjk and eastasia == FALLBACK_CJK_FONT:
            eastasia = lang_specific_cjk
        for run in p.runs:
            _set_run_fonts(run, eastasia, ascii_font)
            changes += 1

    return {
        "fixer": "font_normalize",
        "changes": changes,
        "pdf_truth_used": pdf_truth_used,
        "fallback_used": changes - pdf_truth_used,
        "lang_specific_cjk": lang_specific_cjk or "",
    }
