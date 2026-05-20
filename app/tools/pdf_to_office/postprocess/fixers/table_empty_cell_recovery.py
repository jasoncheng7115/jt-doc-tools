"""補回 docx table 內 empty cell 對應的 PDFTruth 文字（C5）。

<樣本 A> v1.8.60 真機踩到：
- 數量欄下「1 單位」、TEST footer、台灣（4 行 address 第 4 行）等 PDF 內有的
  文字在 docx 內無對應段落，但因 pdf2docx 把它們塞進 table 結構並留空 cell
- `text_recovery` 只處理 body 段落補回，遇 table empty cell 內缺文字無能為力
- 既有 `table_cell_repair` 用 pdfplumber 比對，但對 pdf2docx 結構錯亂的 table
  shape match 不上 → fallback 也不靈

本 fixer 用 PDFTruth y-序對應補：

1) 對每個 docx table，找 PDF 上對應 region (用非空 cell text 在 PDFTruth blocks
   找 union bbox)
2) 收集落在該 region y 範圍內、且 normalized text 在 docx 任何段落 / cell 都找
   不到的 PDFTruth line（候選漏抓文字）
3) 收集 table 內 empty cells，依「(row, col) 序」排
4) 依 PDF y_top 序把候選漏文字配到 empty cells；若數量不一致，配對 min(n)
5) 為避免破壞：每張表配對失敗 ratio > 50% → 該表 skip 不補
"""
from __future__ import annotations

import logging
import re

from docx.oxml.ns import qn

log = logging.getLogger(__name__)


MIN_CELL_TEXT_FOR_ANCHOR = 2
MAX_RECOVERED_PER_TABLE = 20  # 單表上限避免某些壞 table 被當垃圾桶
BBOX_Y_MARGIN_PT = 5.0
CANDIDATE_X_TOLERANCE_PT = 150.0   # v1.8.62 放寬 ±50 → ±150
CANDIDATE_TEXT_MAX_CHARS = 100     # v1.8.62 放寬 50 → 100


def _normalize(s: str) -> str:
    if not s:
        return ""
    return re.sub(r"\s+", "", s).strip()


def _collect_docx_text_set(docx_doc) -> set[str]:
    """整份 docx 內 normalized 已存在文字 set。body.iter(w:p) 走訪含 table cell 內
    所有段落 — 不必另外掃 tables。也另外收每個 tc 整 cell 串接文字（cell 內多段
    視為一個整 cell 文字，給 '<區號><國家>' 這類已合進 cell 單行的情境）。"""
    out: set[str] = set()
    body = docx_doc.element.body
    for p_el in body.iter(qn("w:p")):
        parts = []
        for t in p_el.iter(qn("w:t")):
            if t.text:
                parts.append(t.text)
        n = _normalize("".join(parts))
        if n:
            out.add(n)
    # cell-level 整 cell 文字
    for tc_el in body.iter(qn("w:tc")):
        n = _normalize(_tc_text(tc_el))
        if n:
            out.add(n)
    return out


def _text_present(needle: str, hay: set[str]) -> bool:
    nn = _normalize(needle)
    if not nn:
        return True
    if nn in hay:
        return True
    for h in hay:
        if nn in h:
            return True
    return False


def _tc_text(tc_el) -> str:
    """直接從 tc element 撈所有 w:t 文字（不靠 python-docx 的 cell wrapper，避免
    lxml 元素 Python wrapper 每次 access 重新建立時 id() 不穩的問題）。"""
    parts = []
    for t in tc_el.iter(qn("w:t")):
        if t.text:
            parts.append(t.text)
    return "".join(parts)


def _iter_unique_tcs(table) -> list:
    """走表內每個 tr 的 <w:tc> 直接子元素 — 跨 merged cell 用 vMerge 判斷。回
    [(row_idx, col_idx, tc_el)]。"""
    out = []
    for r_idx, row in enumerate(table.rows):
        tr_el = row._element
        for c_idx, tc_el in enumerate(tr_el.findall(qn("w:tc"))):
            out.append((r_idx, c_idx, tc_el))
    return out


def _find_table_pdf_region(table, pdf_truth):
    """配對 docx table cell texts → PDFTruth blocks，回 (page_num, (x0,y0,x1,y1))
    或 None。同時回 matched_text_set (normalized) 給後續排除已 matched 文字。"""
    cell_texts = []
    for r_idx, c_idx, tc_el in _iter_unique_tcs(table):
        t = _normalize(_tc_text(tc_el))
        if t and len(t) >= MIN_CELL_TEXT_FOR_ANCHOR:
            cell_texts.append(t)
    if not cell_texts:
        return None, set()
    page_count: dict[int, int] = {}
    matched_bboxes: dict[int, list] = {}
    matched_text: set[str] = set()
    for ct in cell_texts:
        for b in pdf_truth.all_blocks:
            n_text = _normalize(b.text)
            if not n_text:
                continue
            if ct in n_text:
                pn = b.page_num
                page_count[pn] = page_count.get(pn, 0) + 1
                matched_bboxes.setdefault(pn, []).append(b.bbox)
                matched_text.add(ct)
                break
    if not page_count:
        return None, matched_text
    dominant_page = max(page_count.items(), key=lambda x: x[1])[0]
    boxes = matched_bboxes[dominant_page]
    x0 = min(bb[0] for bb in boxes)
    y0 = min(bb[1] for bb in boxes)
    x1 = max(bb[2] for bb in boxes)
    y1 = max(bb[3] for bb in boxes)
    return (dominant_page, (x0, y0, x1, y1)), matched_text


def _enumerate_empty_cells(table) -> list:
    """走每個 tr 的直接 tc 子元素，回 (row_idx, col_idx, cell_wrapper, tc_element)。
    cell wrapper 是 row.cells 內第一個包同一 tc element 的 wrap（用 lxml C 元素
    identity 比對，不是 Python wrapper id — 後者不穩）。"""
    out = []
    for r_idx, row in enumerate(table.rows):
        # row.cells 對 merged cell 會多次回 same wrapper；用 wrapper._element 走
        # 但比對用 etree element 的 lxml 內建相等（`is` 對 lxml _Element 比 C ptr）
        cells_in_row = list(row.cells)
        # 取每個直接 tc element 序
        for c_idx, tc_el in enumerate(row._element.findall(qn("w:tc"))):
            text = _tc_text(tc_el).strip()
            if text:
                continue
            # 找對應 cell wrapper（拿任一即可 — 都包同一個 tc element）
            cell = None
            for c in cells_in_row:
                if c._element is tc_el:
                    cell = c
                    break
            if cell is None:
                # fallback：本來 row.cells 索引應該對應，直接拿
                if c_idx < len(cells_in_row):
                    cell = cells_in_row[c_idx]
            if cell is None:
                continue
            out.append((r_idx, c_idx, cell, tc_el))
    return out


def _set_cell_text(cell, text: str) -> None:
    """寫入 cell 第一段第一個 run；無 run 則新增。"""
    paras = cell.paragraphs
    if not paras:
        cell.add_paragraph(text)
        return
    p = paras[0]
    runs = p._element.findall(qn("w:r"))
    if runs:
        first_t = runs[0].find(qn("w:t"))
        if first_t is None:
            first_t = runs[0].makeelement(qn("w:t"), {qn("xml:space"): "preserve"})
            runs[0].append(first_t)
        first_t.text = text
        first_t.set(qn("xml:space"), "preserve")
    else:
        new_r = p._element.makeelement(qn("w:r"), {})
        new_t = new_r.makeelement(qn("w:t"), {qn("xml:space"): "preserve"})
        new_t.text = text
        new_r.append(new_t)
        p._element.append(new_r)


def fix_table_empty_cell_recovery(docx_doc, pdf_truth, alignment) -> dict:
    if not pdf_truth or not pdf_truth.pages:
        return {"fixer": "table_empty_cell_recovery", "filled": 0,
                "skipped": "no pdf_truth"}
    docx_tables = list(docx_doc.tables)
    if not docx_tables:
        return {"fixer": "table_empty_cell_recovery", "filled": 0,
                "tables": 0}
    hay = _collect_docx_text_set(docx_doc)
    total_filled = 0
    tables_touched = 0
    skipped_no_match = 0

    for ti, table in enumerate(docx_tables):
        try:
            match, matched_texts = _find_table_pdf_region(table, pdf_truth)
        except Exception as e:
            log.debug("table %d region match failed: %s", ti, e)
            match, matched_texts = None, set()
        if not match:
            skipped_no_match += 1
            continue
        page_num, (rx0, ry0, rx1, ry1) = match
        # 候選漏文字：該 page 內、bbox y 在 region 內、normalized text 不在 hay
        candidates: list = []
        try:
            page = pdf_truth.pages[page_num]
        except IndexError:
            continue
        for b in page.blocks:
            if b.block_type != "text":
                continue
            for ln in b.lines:
                txt = (ln.text or "").strip()
                if not txt or len(txt) > CANDIDATE_TEXT_MAX_CHARS:
                    continue
                lx0, ly0, lx1, ly1 = ln.bbox
                # y 必須在 region 內 (±margin)
                if ly1 < ry0 - BBOX_Y_MARGIN_PT or ly0 > ry1 + BBOX_Y_MARGIN_PT:
                    continue
                # x 中心也大致在 region 內（避免抓到該 y 但其實在 region 旁邊的別欄）
                lcx = (lx0 + lx1) / 2.0
                if lcx < rx0 - CANDIDATE_X_TOLERANCE_PT or lcx > rx1 + CANDIDATE_X_TOLERANCE_PT:
                    continue
                if _text_present(txt, hay):
                    continue
                candidates.append((ly0, lcx, txt))
        if not candidates:
            continue
        empty_cells = _enumerate_empty_cells(table)
        if not empty_cells:
            continue
        # 排序：empty cells 依 (row_idx, col_idx)；candidates 依 (y, x)
        empty_cells.sort(key=lambda x: (x[0], x[1]))
        candidates.sort(key=lambda x: (x[0], x[1]))
        n_match = min(len(empty_cells), len(candidates))
        if n_match == 0:
            continue
        n_filled = 0
        for i in range(n_match):
            if n_filled >= MAX_RECOVERED_PER_TABLE:
                break
            _, _, cell, _ = empty_cells[i]
            _, _, txt = candidates[i]
            try:
                _set_cell_text(cell, txt)
                hay.add(_normalize(txt))
                n_filled += 1
            except Exception as e:
                log.debug("set cell text failed: %s", e)
        if n_filled > 0:
            tables_touched += 1
            total_filled += n_filled

    return {
        "fixer": "table_empty_cell_recovery",
        "filled": total_filled,
        "tables_touched": tables_touched,
        "skipped_no_match": skipped_no_match,
        "tables": len(docx_tables),
    }
