"""雜訊清理 fixer。

Sprint 1 範圍：
- 連續空段落壓成 1 個
- 移除全空白字元的段落（若超過 max_consecutive_empty）
- 移除孤立單字元段落（除非有意義標點）
- 移除過小圖片 (< tiny_image_threshold_pt)
"""
from __future__ import annotations

import logging

from docx.oxml.ns import qn

from ..config import CLEANUP

log = logging.getLogger(__name__)

_MEANINGFUL_SINGLE = set("。．，,．、？！?!:：；;()（）「」『』『』「」")


def _is_blank(p) -> bool:
    return not (p.text or "").strip()


def fix_cleanup(docx_doc, pdf_truth, alignment) -> dict:
    removed_empty = 0
    removed_tiny_image = 0

    # ----- 連續空段落壓縮 -----
    if CLEANUP.get("compress_empty_paragraphs", True):
        max_consec = CLEANUP.get("max_consecutive_empty", 1)
        consecutive = 0
        # python-docx 的 doc.paragraphs 不含表格內 — Sprint 1 只處理 body 段落
        for p in list(docx_doc.paragraphs):
            if _is_blank(p):
                consecutive += 1
                if consecutive > max_consec:
                    p._element.getparent().remove(p._element)
                    removed_empty += 1
            else:
                consecutive = 0

    # ----- 移除過小圖片 -----
    if CLEANUP.get("remove_tiny_images", True):
        threshold = CLEANUP.get("tiny_image_threshold_pt", 10.0)
        # docx 內嵌圖片在 inline drawings 裡，size 是 EMU (1 pt = 12700 EMU)
        threshold_emu = threshold * 12700
        for p in list(docx_doc.paragraphs):
            for r in list(p.runs):
                drawings = r._element.findall(qn("w:drawing"))
                for drw in drawings:
                    extents = drw.findall(".//" + qn("a:ext"))
                    for ext in extents:
                        try:
                            cx = int(ext.get("cx") or 0)
                            cy = int(ext.get("cy") or 0)
                            if 0 < cx < threshold_emu and 0 < cy < threshold_emu:
                                r._element.remove(drw)
                                removed_tiny_image += 1
                                break
                        except (TypeError, ValueError):
                            pass

    return {
        "fixer": "cleanup",
        "removed_empty_paragraphs": removed_empty,
        "removed_tiny_images": removed_tiny_image,
    }
