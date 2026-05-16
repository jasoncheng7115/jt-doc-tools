"""頁首頁尾識別 fixer。

把每頁重複出現的頁首頁尾從內文移到 Word 的 header/footer。

策略：
- 多頁 PDF：跨頁聚合 — ≥ 50% 頁面同 y 範圍出現相似文字 → 視為頁首頁尾
- 單頁 PDF（v1.8.42+ 加）：用 footer pattern 啟發式 — 最下方 (y > 0.85*page_h) 的
  block 若含 contact info pattern (電話 / Email / 統編 / Page X/Y / @ / .com / 頁次)
  → 視為頁尾移到 docx footer

保守策略：
- ≤ 1 個 candidate 不處理
- 文字過長 (> 200 字) 不視為頁首頁尾（誤判風險高）
"""
from __future__ import annotations

import logging
import re

from docx.shared import Pt

log = logging.getLogger(__name__)

HEADER_TOP_RATIO = 0.1
FOOTER_BOTTOM_RATIO = 0.9
MIN_PAGE_RATIO_FOR_DETECTION = 0.5

# 單頁 PDF footer 啟發式 — 文字 contains any of these → likely footer
_FOOTER_HINTS = (
    "電話", "Tel", "TEL", "Phone",
    "郵件", "Email", "@",
    "統編", "統一編號", "VAT",
    "頁:", "頁：", "Page", "Page ",
    "/1", "/2", "/3", "頁次",
    "傳真", "Fax", "FAX",
)
MAX_FOOTER_CHARS = 200


def _normalize(s: str) -> str:
    return " ".join((s or "").split())


def _normalize_loose(s: str) -> str:
    """寬鬆 normalize — 移除 PUA 字符 (FontAwesome icon glyph) + 全部空白合併。"""
    if not s:
        return ""
    # PUA private use area (E000-F8FF) 是常見 icon font 範圍 — strip
    out = []
    for ch in s:
        if "" <= ch <= "":
            continue
        out.append(ch)
    return "".join(out).split() and " ".join("".join(out).split()) or ""


def _candidates(pages, top: bool) -> list[str]:
    """收集所有頁的頁首 / 頁尾候選文字。"""
    out: list[str] = []
    for p in pages:
        for b in p.blocks:
            if b.block_type != "text":
                continue
            x0, y0, x1, y1 = b.bbox
            if top:
                if y1 <= p.height * HEADER_TOP_RATIO:
                    out.append(_normalize(b.text))
            else:
                if y0 >= p.height * FOOTER_BOTTOM_RATIO:
                    out.append(_normalize(b.text))
    return [t for t in out if t]


def _common_texts(candidates: list[str], page_count: int) -> list[str]:
    """找出在 ≥ 50% 頁面出現的文字 — 相同字串視為「同一個頁首」。"""
    if not candidates or page_count < 2:
        return []
    from collections import Counter
    c = Counter(candidates)
    threshold = max(2, int(page_count * MIN_PAGE_RATIO_FOR_DETECTION))
    return [text for text, n in c.items() if n >= threshold]


def _remove_paragraphs_with_text(docx_doc, texts: set[str]) -> int:
    if not texts:
        return 0
    removed = 0
    for p in list(docx_doc.paragraphs):
        if _normalize(p.text) in texts:
            try:
                p._element.getparent().remove(p._element)
                removed += 1
            except Exception:
                pass
    return removed


def _set_section_header(section, text: str) -> None:
    try:
        hdr = section.header
        # 清空現有段落內容
        for p in hdr.paragraphs:
            p.text = ""
        if hdr.paragraphs:
            hdr.paragraphs[0].text = text
        else:
            hdr.add_paragraph(text)
    except Exception:
        pass


def _set_section_footer(section, text: str) -> None:
    try:
        ftr = section.footer
        for p in ftr.paragraphs:
            p.text = ""
        if ftr.paragraphs:
            ftr.paragraphs[0].text = text
        else:
            ftr.add_paragraph(text)
    except Exception:
        pass


def _looks_like_footer(text: str) -> bool:
    """單頁 PDF 用 — 文字像 contact info / 頁碼 / 法律宣告 → 視為 footer。"""
    if not text or len(text) > MAX_FOOTER_CHARS:
        return False
    return any(h in text for h in _FOOTER_HINTS)


# Footer 自動偵測 regex — 連續含 contact info pattern 的 substring 視為 footer
# (含可選的 PUA 圖示字符 - — 例如 fontawesome 字型)
_FOOTER_INLINE_RE = re.compile(
    r"[-\s]*"                          # 開頭可能的 PUA icon
    r"(?:電話|Tel|TEL|Phone|傳真|Fax|FAX)\s*[:：]?\s*[\d\-\+\(\) ]{6,}"  # 電話
    r".*?"
    r"(?:郵件|Email|email|E-?mail)?\s*[:：]?\s*[\w\.\-+]+@[\w\.\-]+\.[A-Za-z]{2,}"  # email
    r".*?"
    r"(?:統編|統一編號|VAT)\s*[:：]?\s*\d{6,10}"   # 統編
    r"(?:.*?(?:頁\s*[:：]?\s*\d+\s*/\s*\d+|Page\s*\d+))?",  # 可選 Page 1/N
    re.UNICODE | re.DOTALL,
)


_PAGE_NUM_RE = re.compile(
    r"(?:頁|Page|p\.?|Pg)\s*[:：]?\s*\d+\s*[/／]?\s*\d*",
    re.IGNORECASE,
)
# 「strong footer evidence」— 單段內含 ≥ 2 個獨立 contact pattern（電話 / email
# / 統編 / 頁碼），認定為純粹頁尾段落
_PHONE_RE = re.compile(r"\b\d{2,4}[-\s]?\d{3,4}[-\s]?\d{4}\b|\b886[-\s]?\d")
_EMAIL_RE = re.compile(r"[\w\.\-+]+@[\w\.\-]+\.[A-Za-z]{2,}")
_VAT_RE = re.compile(r"(?:統編|統一編號|VAT|統一編號)\s*[:：]?\s*\d{6,10}")


def _looks_like_pure_footer(txt: str) -> bool:
    """段落整段就是 footer — ≥ 2 個 contact 證據 + 全段不長 (純 footer，無正文)。

    重要：不能只看 contact pattern 數量。pdf2docx 會把 footer 黏到收件方資訊
    或正文段內，整段同時含「公司名 + 地址 + 電話 + 信箱」這類聯絡欄位但
    **前面有真實業務內容**。這種 case 必須剝離 footer 不是整段刪。

    判定為「純 footer 段」(可整段刪)：
    - 長度 ≤ 80 字 (短，沒有實質正文夾雜)
    - score ≥ 2 contact patterns
    """
    if not txt:
        return False
    n = _normalize(txt)
    if len(n) > 80:
        return False  # 太長 — 一定有正文，不可整段刪
    score = 0
    if _PHONE_RE.search(n): score += 1
    if _EMAIL_RE.search(n): score += 1
    if _VAT_RE.search(n): score += 1
    if _PAGE_NUM_RE.search(n): score += 1
    return score >= 2


def _strip_footer_substring(docx_doc, footer_text: str) -> int:
    """在 docx 內找含 footer pattern 的段落，把 footer substring 剝離。

    支援多種 fallback：
    1. 直接 footer_text 子字串 in para text → replace
    2. loose normalize 比對（去 PUA 字符 + 空白）→ 找出 footer 起點剝離
    3. _looks_like_pure_footer — 整段都是 footer evidence → 整段刪
    4. heuristic: para 含 footer 半段 + 結尾頁碼 → 從 contact hint 切走
    """
    if not footer_text:
        return 0
    target_loose = _normalize_loose(footer_text)
    for p in list(docx_doc.paragraphs):
        txt = (p.text or "")
        if not txt.strip():
            continue
        # 1. exact substring
        if footer_text in txt:
            new_text = txt.replace(footer_text, "").strip()
            _replace_paragraph_text(p, new_text)
            return 1
        # 2. loose normalize 比對（PUA / 空白差異）
        if target_loose and target_loose in _normalize_loose(txt):
            # 找 docx para 內 footer 起點 — 用第一個 contact pattern 位置切
            for matcher in (_PHONE_RE, _EMAIL_RE, _VAT_RE):
                m = matcher.search(txt)
                if m and m.start() > 0:
                    # 往前找最近的 PUA 字符 / 標點當切點
                    cut = m.start()
                    for i in range(m.start() - 1, max(0, m.start() - 5), -1):
                        if "" <= txt[i] <= "" or txt[i] in " \t":
                            cut = i
                            break
                    new_text = txt[:cut].rstrip()
                    _replace_paragraph_text(p, new_text)
                    return 1
            # 找不到切點 → 整段刪
            p._element.getparent().remove(p._element)
            return 1
        # 3. 整段是 footer evidence
        if _looks_like_pure_footer(txt):
            p._element.getparent().remove(p._element)
            return 1
        # 4. heuristic: para 結尾段含 footer + 頁碼
        if _PAGE_NUM_RE.search(txt) and (_EMAIL_RE.search(txt) or _VAT_RE.search(txt)):
            for kw in ("電話", "Tel", "TEL", "Phone", "傳真", "Fax", "FAX"):
                idx = txt.find(kw)
                if idx > 0:
                    new_text = txt[:idx].rstrip()
                    _replace_paragraph_text(p, new_text)
                    return 1
            # 沒中文 keyword — 用 PUA / 第一 contact 切
            for matcher in (_PHONE_RE, _EMAIL_RE):
                m = matcher.search(txt)
                if m and m.start() > 0:
                    new_text = txt[:m.start()].rstrip()
                    # 移掉 trailing PUA chars
                    while new_text and "" <= new_text[-1] <= "":
                        new_text = new_text[:-1]
                    _replace_paragraph_text(p, new_text.rstrip())
                    return 1
    return 0


def _replace_paragraph_text(p, new_text: str) -> None:
    """把段落內容換成 new_text，保留第一個 run 屬性；空字串則刪除整段。"""
    if not new_text.strip():
        try:
            p._element.getparent().remove(p._element)
        except Exception:
            pass
        return
    if p.runs:
        p.runs[0].text = new_text
        for r in p.runs[1:]:
            r.text = ""
    else:
        p.add_run(new_text)


def fix_header_footer(docx_doc, pdf_truth, alignment) -> dict:
    if pdf_truth is None or pdf_truth.total_pages < 1:
        return {"fixer": "header_footer", "moved_to_header": 0,
                "moved_to_footer": 0, "skipped": "no pdf"}

    # 單頁 PDF — 用啟發式找底部 footer-like block
    if pdf_truth.total_pages == 1:
        page = pdf_truth.pages[0]
        candidates = []
        for b in page.blocks:
            if b.block_type != "text":
                continue
            x0, y0, x1, y1 = b.bbox
            if y0 < page.height * FOOTER_BOTTOM_RATIO:
                continue
            n = _normalize(b.text)
            if _looks_like_footer(n):
                candidates.append(n)
        moved_f = 0
        footer_text_used = ""
        if candidates:
            footer_text = candidates[0]
            footer_text_used = footer_text
            moved_f = _remove_paragraphs_with_text(docx_doc, {footer_text})
            if moved_f == 0:
                # 整段精確比對失敗 — 用 substring + fuzzy 剝離
                # （pdf2docx 把 footer 黏到別段內、或 unicode 私有區字符差異）
                moved_f = _strip_footer_substring(docx_doc, footer_text)
            if moved_f:
                _set_section_footer(docx_doc.sections[0], footer_text)
        return {
            "fixer": "header_footer",
            "moved_to_header": 0,
            "moved_to_footer": moved_f,
            "single_page_heuristic": True,
            "footer_candidates": len(candidates),
            "footer_text_preview": (footer_text_used[:50] + "...") if footer_text_used else "",
        }

    # 多頁 PDF — 跨頁聚合
    headers = _common_texts(_candidates(pdf_truth.pages, top=True), pdf_truth.total_pages)
    footers = _common_texts(_candidates(pdf_truth.pages, top=False), pdf_truth.total_pages)

    moved_h = 0
    moved_f = 0

    if headers:
        # 取最常見的一個當主要頁首；其餘保留在內文
        header_text = headers[0]
        moved_h = _remove_paragraphs_with_text(docx_doc, {header_text})
        if moved_h:
            _set_section_header(docx_doc.sections[0], header_text)
    if footers:
        footer_text = footers[0]
        moved_f = _remove_paragraphs_with_text(docx_doc, {footer_text})
        if moved_f:
            _set_section_footer(docx_doc.sections[0], footer_text)

    return {
        "fixer": "header_footer",
        "moved_to_header": moved_h,
        "moved_to_footer": moved_f,
        "header_candidates": len(headers),
        "footer_candidates": len(footers),
    }
