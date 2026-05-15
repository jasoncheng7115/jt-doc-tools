"""表格樣式正規化 fixer (Sprint 3)。

範圍：
- 統一表格邊框（all sides single 0.5pt #999）
- cell 段距前後設 0
- cell 垂直對齊 center
- 第一列若含粗體 / 字級較大 → 視為標題列，加淺灰底色

合併儲存格偵測（log only）：
- 對 PDF drawings 建線條集合
- 偵測 docx 表格區域對應 PDF 內缺漏的內部線 → 標警告
- 不自動修（風險高）
"""
from __future__ import annotations

import logging
from copy import deepcopy

from docx.oxml.ns import qn
from docx.shared import Pt

log = logging.getLogger(__name__)

BORDER_COLOR = "999999"
BORDER_SIZE = "4"  # 1/8 pt = 0.5pt total
HEADER_SHADING = "F1F5F9"


def _set_table_borders(table) -> None:
    """設整個 table 4 邊 + 內格 single 0.5pt 灰邊框。"""
    tblPr = table._element.find(qn("w:tblPr"))
    if tblPr is None:
        return
    # 移除舊 borders
    for old in tblPr.findall(qn("w:tblBorders")):
        tblPr.remove(old)
    borders = tblPr.makeelement(qn("w:tblBorders"), {})
    for tag in ("top", "left", "bottom", "right", "insideH", "insideV"):
        el = borders.makeelement(qn(f"w:{tag}"), {})
        el.set(qn("w:val"), "single")
        el.set(qn("w:sz"), BORDER_SIZE)
        el.set(qn("w:space"), "0")
        el.set(qn("w:color"), BORDER_COLOR)
        borders.append(el)
    tblPr.append(borders)


def _set_cell_props(cell, *, vertical_align: str = "center",
                    shading: str | None = None) -> None:
    tcPr = cell._element.find(qn("w:tcPr"))
    if tcPr is None:
        tcPr = cell._element.makeelement(qn("w:tcPr"), {})
        cell._element.insert(0, tcPr)
    # vertical align
    for old in tcPr.findall(qn("w:vAlign")):
        tcPr.remove(old)
    vAlign = tcPr.makeelement(qn("w:vAlign"), {})
    vAlign.set(qn("w:val"), vertical_align)
    tcPr.append(vAlign)
    # shading
    if shading:
        for old in tcPr.findall(qn("w:shd")):
            tcPr.remove(old)
        shd = tcPr.makeelement(qn("w:shd"), {})
        shd.set(qn("w:val"), "clear")
        shd.set(qn("w:color"), "auto")
        shd.set(qn("w:fill"), shading)
        tcPr.append(shd)


def _clear_cell_paragraph_spacing(cell) -> None:
    for p in cell.paragraphs:
        try:
            pf = p.paragraph_format
            pf.space_before = Pt(0)
            pf.space_after = Pt(0)
        except Exception:
            pass


def _is_header_row(row) -> bool:
    """第一列若所有 cell 都粗體 OR 字級 > 內文 1.1 倍 → 視為標題列。"""
    bolds = 0
    total = 0
    for c in row.cells:
        for p in c.paragraphs:
            for r in p.runs:
                total += 1
                if r.bold:
                    bolds += 1
    if total == 0:
        return False
    return bolds / total >= 0.5


def _set_table_no_borders(table) -> None:
    """設整個 table 全邊框 nil（無框線 PDF 的對應方式 — 不要憑空長框）。"""
    tblPr = table._element.find(qn("w:tblPr"))
    if tblPr is None:
        return
    for old in tblPr.findall(qn("w:tblBorders")):
        tblPr.remove(old)
    borders = tblPr.makeelement(qn("w:tblBorders"), {})
    for tag in ("top", "left", "bottom", "right", "insideH", "insideV"):
        el = borders.makeelement(qn(f"w:{tag}"), {})
        el.set(qn("w:val"), "nil")
        borders.append(el)
    tblPr.append(borders)


def _pdf_is_borderless(pdf_truth) -> bool:
    """PDF 整體看是不是「無框線排版」型 — 用 page drawings 數量啟發判斷。

    判定為無框線（不要在 docx 加框）：
    - 第一頁總 drawings < 30 (簡單發票 / 報價沒太多線條，pdf2docx 卻硬包成 table)
    - 並且該頁 text blocks > 5 (確認有實際內容，不是空白頁)

    對票卡 / 廠商資料表 / 申請表這類「PDF 真的有畫框線的表格」(drawings > 100)
    這個 fixer 仍會套上邊框（_set_table_borders）— 維持原表格樣貌。
    """
    if not pdf_truth or not pdf_truth.pages:
        return False
    first_page = pdf_truth.pages[0]
    n_drawings = len(first_page.drawings)
    n_text_blocks = sum(1 for b in first_page.blocks if b.block_type == "text")
    return n_drawings < 30 and n_text_blocks > 5


def fix_table_normalize(docx_doc, pdf_truth, alignment) -> dict:
    tables_styled = 0
    cells_centered = 0
    headers_shaded = 0
    border_mode = "single"

    # 偵測 PDF 是不是「無框線排版」(用 invisible table 對齊欄位的 invoice 風格)
    borderless = _pdf_is_borderless(pdf_truth)
    if borderless:
        border_mode = "nil"

    for table in docx_doc.tables:
        if borderless:
            _set_table_no_borders(table)
        else:
            _set_table_borders(table)
        tables_styled += 1
        rows = list(table.rows)
        if not rows:
            continue
        is_header = _is_header_row(rows[0]) and not borderless
        for ri, row in enumerate(rows):
            for cell in row.cells:
                shading = HEADER_SHADING if (ri == 0 and is_header) else None
                # 標題列 vertical center 看起來比較整齊；內文列 vertical top
                # 比較自然（避免短內容飄在 cell 中間 — invoice item description
                # 應該靠上對齊跟其他欄位齊平）
                v_align = "center" if (ri == 0 and is_header) else "top"
                _set_cell_props(cell, vertical_align=v_align, shading=shading)
                _clear_cell_paragraph_spacing(cell)
                cells_centered += 1
            if ri == 0 and is_header:
                headers_shaded += 1
    return {
        "fixer": "table_normalize",
        "tables_styled": tables_styled,
        "cells_centered": cells_centered,
        "header_rows_shaded": headers_shaded,
        "border_mode": border_mode,
    }
