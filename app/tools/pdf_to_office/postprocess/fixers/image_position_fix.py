"""圖片位置 / 大小校正 fixer（Sprint B #9 強化）。

策略分三層配對 docx 內嵌圖 ↔ PDFTruth.images：

1) **content-hash 配對**（既有）：SHA1 byte hash 完全一致 → 100% 同一張
2) **bbox-size 配對**（Sprint B 新加 fallback）：hash 失敗時，按「PDF 像素長寬
   比 + bbox pt 大小」近似配對。pdf2docx 抽出圖片時可能重壓 / 改 encoding，
   SHA1 會錯，但像素比例 + bbox 大小通常守得住
3) **段落水平對齊**（Sprint B 新加）：依 PDF bbox 在頁面 X 軸位置推斷 image 應
   align center / right / left，套到 image 所在段落的 alignment

並非試圖把 docx inline image 變 absolute-positioned（那會破壞文字流），而是把
「大小 + 水平對齊」校到接近 PDF 真值。
"""
from __future__ import annotations

import hashlib
import logging

from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml.ns import qn

log = logging.getLogger(__name__)

EMU_PER_PT = 12700
SIZE_DIFF_TOLERANCE = 0.10        # cx / cy 差異 > 10% 才動
BBOX_SIZE_MATCH_TOL = 0.15        # 寬高比 ± 15% 視為同圖（fallback 配對用）
PARA_ALIGN_LEFT_RIGHT_GAP = 0.20  # bbox 中心離頁中軸 > 頁寬 20% 視為 left/right


def _emu_to_pt(emu: int) -> float:
    return float(emu) / EMU_PER_PT


def _safe_int(x, default=0) -> int:
    try:
        return int(x)
    except (TypeError, ValueError):
        return default


def _aspect(w: float, h: float) -> float:
    if h <= 0:
        return 0.0
    return w / h


def _collect_docx_inline_images(docx_doc) -> list[dict]:
    """走訪 docx 內所有 inline drawings — 同時把所在 paragraph 也記下來（給
    水平對齊用）。"""
    out: list[dict] = []
    image_part_map: dict[str, str] = {}
    try:
        for rel in docx_doc.part.rels.values():
            if "image" in (rel.reltype or "").lower():
                target = rel.target_part
                try:
                    blob = target.blob
                    h = hashlib.sha1(blob, usedforsecurity=False).hexdigest()[:16]
                    image_part_map[rel.rId] = h
                except Exception:
                    pass
    except Exception as e:
        log.debug("collect image parts failed: %s", e)

    def _walk_paras(paras):
        for p in paras:
            for r in p.runs:
                for drw in r._element.findall(qn("w:drawing")):
                    blip = drw.find(".//" + qn("a:blip"))
                    rid = blip.get(qn("r:embed")) if blip is not None else None
                    h = image_part_map.get(rid, "")
                    ext = drw.find(".//" + qn("a:ext"))
                    cx = _safe_int(ext.get("cx")) if ext is not None else 0
                    cy = _safe_int(ext.get("cy")) if ext is not None else 0
                    out.append({
                        "paragraph": p, "run": r, "drawing": drw, "ext": ext,
                        "rid": rid, "hash": h, "cx": cx, "cy": cy,
                    })

    _walk_paras(docx_doc.paragraphs)
    for tbl in docx_doc.tables:
        for row in tbl.rows:
            for cell in row.cells:
                _walk_paras(cell.paragraphs)
    return out


def _match_by_bbox_size(docx_img: dict, pdf_images_unmatched: list,
                        used: set) -> object | None:
    """fallback：用 PDF 像素比 + bbox 大小匹配。回 PDFImage 或 None。"""
    dx_cx_pt = _emu_to_pt(docx_img["cx"]) if docx_img["cx"] > 0 else 0
    dx_cy_pt = _emu_to_pt(docx_img["cy"]) if docx_img["cy"] > 0 else 0
    if dx_cx_pt <= 0 or dx_cy_pt <= 0:
        return None
    dx_aspect = _aspect(dx_cx_pt, dx_cy_pt)
    best = None
    best_score = float("inf")
    for pdf_im in pdf_images_unmatched:
        if id(pdf_im) in used:
            continue
        x0, y0, x1, y1 = pdf_im.bbox
        pdf_w = max(0.0, x1 - x0)
        pdf_h = max(0.0, y1 - y0)
        if pdf_w <= 0 or pdf_h <= 0:
            continue
        pdf_aspect = _aspect(pdf_w, pdf_h)
        if pdf_aspect <= 0 or dx_aspect <= 0:
            continue
        # 比例差
        ar_diff = abs(pdf_aspect - dx_aspect) / max(pdf_aspect, dx_aspect)
        if ar_diff > BBOX_SIZE_MATCH_TOL:
            continue
        # 寬度差作為次要 score
        w_diff = abs(pdf_w - dx_cx_pt) / max(pdf_w, dx_cx_pt)
        score = ar_diff + w_diff
        if score < best_score:
            best_score = score
            best = pdf_im
    return best


def _infer_alignment(pdf_im, page_width: float) -> str | None:
    """從 PDF bbox X 中心推 horizontal alignment。回 'center' / 'right' / 'left' / None。"""
    if not pdf_im or page_width <= 0:
        return None
    x0, _, x1, _ = pdf_im.bbox
    cx = (x0 + x1) / 2.0
    page_cx = page_width / 2.0
    offset = (cx - page_cx) / page_width  # -0.5 .. 0.5
    if abs(offset) < 0.05:
        return "center"
    if offset > PARA_ALIGN_LEFT_RIGHT_GAP:
        return "right"
    if offset < -PARA_ALIGN_LEFT_RIGHT_GAP:
        return "left"
    return None  # 接近 center 但沒對齊到，不動


def _apply_alignment(paragraph, align: str) -> bool:
    """套段落水平對齊。回 True 表示有改。"""
    if not align:
        return False
    mapping = {
        "center": WD_PARAGRAPH_ALIGNMENT.CENTER,
        "right": WD_PARAGRAPH_ALIGNMENT.RIGHT,
        "left": WD_PARAGRAPH_ALIGNMENT.LEFT,
    }
    target = mapping.get(align)
    if target is None:
        return False
    if paragraph.alignment == target:
        return False
    try:
        paragraph.alignment = target
        return True
    except Exception as e:
        log.debug("set paragraph alignment failed: %s", e)
        return False


def _build_page_width_index(pdf_truth) -> dict[int, float]:
    return {p.page_num: float(p.width) for p in pdf_truth.pages}


def fix_image_position_fix(docx_doc, pdf_truth, alignment) -> dict:
    if pdf_truth is None:
        return {"fixer": "image_position_fix", "adjusted": 0,
                "skipped_no_pdftruth": True}

    # 整理 PDF images：按 hash 索引 + 扁平 list 給 fallback
    pdf_imgs_by_hash: dict[str, list] = {}
    all_pdf_imgs: list = []
    for p in pdf_truth.pages:
        for im in p.images:
            if im.image_hash:
                pdf_imgs_by_hash.setdefault(im.image_hash, []).append(im)
            all_pdf_imgs.append(im)
    page_widths = _build_page_width_index(pdf_truth)

    docx_imgs = _collect_docx_inline_images(docx_doc)
    if not docx_imgs or not all_pdf_imgs:
        return {"fixer": "image_position_fix", "adjusted": 0,
                "docx_images": len(docx_imgs),
                "pdf_images": len(all_pdf_imgs)}

    size_adjusted = 0
    aligned = 0
    matched_hash = 0
    matched_bbox = 0
    used: set = set()

    for di in docx_imgs:
        pdf_im = None
        # Layer 1: SHA1 hash
        if di["hash"]:
            candidates = pdf_imgs_by_hash.get(di["hash"]) or []
            for c in candidates:
                if id(c) not in used:
                    pdf_im = c
                    break
            if pdf_im is not None:
                matched_hash += 1
        # Layer 2: bbox-size fallback
        if pdf_im is None:
            pdf_im = _match_by_bbox_size(di, all_pdf_imgs, used)
            if pdf_im is not None:
                matched_bbox += 1
        if pdf_im is None:
            continue
        used.add(id(pdf_im))

        # ---- 大小調整 ----
        x0, y0, x1, y1 = pdf_im.bbox
        pdf_w_pt = max(0.0, x1 - x0)
        pdf_h_pt = max(0.0, y1 - y0)
        if pdf_w_pt > 0 and pdf_h_pt > 0:
            target_cx = int(pdf_w_pt * EMU_PER_PT)
            target_cy = int(pdf_h_pt * EMU_PER_PT)
            ext = di["ext"]
            if di["cx"] > 0 and abs(di["cx"] - target_cx) / di["cx"] > SIZE_DIFF_TOLERANCE:
                if ext is not None:
                    ext.set("cx", str(target_cx))
                    size_adjusted += 1
            if di["cy"] > 0 and abs(di["cy"] - target_cy) / di["cy"] > SIZE_DIFF_TOLERANCE:
                if ext is not None:
                    ext.set("cy", str(target_cy))

        # ---- 段落水平對齊 ----
        page_w = page_widths.get(pdf_im.page_num, 0.0)
        align = _infer_alignment(pdf_im, page_w)
        if align and _apply_alignment(di["paragraph"], align):
            aligned += 1

    # ---- 補插：同 hash PDF 圖出現在 N 個 bbox 但 docx 只放了 < N 張 ----
    # 例：兩個相同 starburst shape 在 PDF 各自位置；pdf2docx 只插了一張 → 補一張。
    # 策略：複製既有 inline drawing element 整個（保留同 rId）插在「位置最接近
    # 缺漏 PDF bbox 的 docx anchor」之後。位置不準，但內容回來。
    inserted = 0
    if pdf_imgs_by_hash:
        # 建 docx 內每張圖出現的「對應 PDF image」累計（hash → list of pdf_im）
        # used set 內已是「被任一 docx_imgs match 到的 pdf_im id」；剩下沒 match
        # 到的 pdf_im 若 hash 跟某 docx_img 一致 → 補插
        leftover_pdf: dict[str, list] = {}
        for h, ims in pdf_imgs_by_hash.items():
            for im in ims:
                if id(im) not in used:
                    leftover_pdf.setdefault(h, []).append(im)
        if leftover_pdf:
            # docx 內各 hash 的「典型 drawing element」拿來 clone
            from copy import deepcopy
            tmpl_by_hash: dict[str, dict] = {}
            for di in docx_imgs:
                if di["hash"] and di["hash"] not in tmpl_by_hash:
                    tmpl_by_hash[di["hash"]] = di
            for h, ims in leftover_pdf.items():
                tmpl = tmpl_by_hash.get(h)
                if not tmpl:
                    continue
                for im in ims:
                    # 防爆量：單一 hash 最多補 5 張
                    if inserted >= 10:
                        break
                    try:
                        run_el = tmpl["run"]._element
                        # 找 run_el 內 w:drawing 整個 clone
                        old_drw = tmpl["drawing"]
                        new_drw = deepcopy(old_drw)
                        # 套用 PDF 真值大小
                        x0, y0, x1, y1 = im.bbox
                        w_pt = max(0.0, x1 - x0)
                        h_pt = max(0.0, y1 - y0)
                        if w_pt > 0 and h_pt > 0:
                            new_ext = new_drw.find(".//" + qn("a:ext"))
                            if new_ext is not None:
                                new_ext.set("cx", str(int(w_pt * EMU_PER_PT)))
                                new_ext.set("cy", str(int(h_pt * EMU_PER_PT)))
                        # 在 template paragraph 後新增一段，把新圖放進去
                        anchor_p = tmpl["paragraph"]._element
                        new_p_el = anchor_p.makeelement(qn("w:p"), {})
                        new_r_el = new_p_el.makeelement(qn("w:r"), {})
                        new_r_el.append(new_drw)
                        new_p_el.append(new_r_el)
                        parent = anchor_p.getparent()
                        if parent is None:
                            continue
                        idx = list(parent).index(anchor_p)
                        parent.insert(idx + 1, new_p_el)
                        inserted += 1
                    except Exception as e:
                        log.debug("insert missing duplicate image failed: %s", e)

    return {
        "fixer": "image_position_fix",
        "adjusted": size_adjusted,
        "aligned_paragraphs": aligned,
        "matched_by_hash": matched_hash,
        "matched_by_bbox_size": matched_bbox,
        "inserted_duplicates": inserted,
        "docx_images": len(docx_imgs),
        "pdf_images": len(all_pdf_imgs),
    }
