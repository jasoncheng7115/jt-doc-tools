"""中文排版修正 fixer。

Sprint 2 範圍（保守）：
- CJK 字之間的單一半形空白（PDF 字距渲染殘留）→ 移除
- 中英文之間的單一空白保留
- 段落 autoSpaceDE / autoSpaceDN 開啟（Word 中英自動字距）

不做（避免修改原文）：
- 標點全形化（會改原文，user 可能不希望）
- 段首縮排調整（依文件慣例不一致時容易過度修）
"""
from __future__ import annotations

import logging
import re

from docx.oxml.ns import qn

log = logging.getLogger(__name__)

# CJK char range — Hiragana / Katakana / CJK Ext-A + main / Hangul syllables。
# 用 \u 顯式定義避免重複範圍（CodeQL js/regex/overly-large-range）。
# - U+3040–U+30FF: 平假名 + 片假名
# - U+3400–U+9FFF: CJK Extension A + 主 CJK Unified Ideographs
# - U+AC00–U+D7AF: 韓文 Hangul Syllables
_CJK_CHARS = "\u3040-\u30ff\u3400-\u9fff\uac00-\ud7af"
_CJK_RE = re.compile(f"[{_CJK_CHARS}]")
# 連續 CJK + 單空白 + CJK → 移除空白
_INTER_CJK_SPACE = re.compile(f"([{_CJK_CHARS}])[\xa0 \t]+([{_CJK_CHARS}])")
# Private Use Area glyphs (U+E000–U+F8FF, BMP PUA) — 多為 PDF 嵌字 icon font
# (FontAwesome / Material / 自訂 dingbat) 沒映 ToUnicode 留下的殘留亂碼，對純文字
# 讀者沒意義；轉成 Word 後直接移除（保留前後空白避免黏字）。
# CodeQL「overly large range」是 by design — PUA 整個區段都要清。
# lgtm[py/regex/overly-large-range]
_PUA_RE = re.compile("[\ue000-\uf8ff]+")


def _is_listy(p) -> bool:
    style_name = (p.style.name if p.style else "") or ""
    # 表格 cell 內、List/Heading 不動
    return any(k in style_name for k in ("List", "Heading", "Title"))


def _is_codey(p) -> bool:
    """code 段落不動 — 用第一個 run 字型判斷。"""
    if not p.runs:
        return False
    name = (p.runs[0].font.name or "").lower()
    return any(h in name for h in ("courier", "mono", "consolas", "menlo"))


def _clean_inter_cjk_spaces(text: str) -> tuple[str, int]:
    """回 (清理後 text, 移除空白數)。

    保留：日期、地址、電話等常用「字 + 多空白 + 字」表單欄位骨架，
    避免「年  月  日」「年  月  日 自年月日起」被吞成「年月日」失去填寫間距。
    """
    # 表單欄位佔位 — 保留多 (≥2) 半形空白
    PRESERVE = re.compile(
        r"(年[  \t]{1,}月[  \t]{1,}日"
        r"|時[  \t]{1,}分"
        r"|公[  \t]{1,}里[  \t]{1,}公[  \t]{1,}尺"
        r")"
    )
    # 把 PRESERVE 區段抽出 → 用 placeholder 暫存 → 清理 → 還原
    placeholders: list[str] = []
    def _stash(m):
        placeholders.append(m.group(0))
        return f"\x02PRESERVE{len(placeholders)-1}\x02"
    masked = PRESERVE.sub(_stash, text)

    n = 0
    new = masked
    while True:
        replaced, count = _INTER_CJK_SPACE.subn(r"\1\2", new)
        if count == 0:
            break
        new = replaced
        n += count

    # 還原 PRESERVE 區段
    for i, ph in enumerate(placeholders):
        new = new.replace(f"\x02PRESERVE{i}\x02", ph)

    return (new if n else text, n)


def _walk_paragraphs(doc):
    for p in doc.paragraphs:
        yield p
    for tbl in doc.tables:
        for row in tbl.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    yield p


def fix_cjk_typography(docx_doc, pdf_truth, alignment) -> dict:
    spaces_removed = 0
    pua_removed = 0
    paragraphs_touched = 0

    for p in _walk_paragraphs(docx_doc):
        if _is_listy(p) or _is_codey(p):
            continue
        if not p.runs:
            continue
        # 先合併 run 文字 → 清理 → 攤回去
        full = "".join(r.text or "" for r in p.runs)
        # 1) PUA glyph 清理（跨段都做，不限 CJK 段）
        pua_n = 0
        if _PUA_RE.search(full):
            full_no_pua, pua_n = _PUA_RE.subn(" ", full)
            # 連空白壓回單空白；段首尾空白 strip
            import re as _re
            full_no_pua = _re.sub(r"[  \t]{2,}", " ", full_no_pua).strip()
            full = full_no_pua
        # 2) inter-CJK 空白清理（僅 CJK 段才跑）
        cleaned, n = (full, 0)
        if _CJK_RE.search(full):
            cleaned, n = _clean_inter_cjk_spaces(full)
        if pua_n == 0 and n == 0:
            continue
        # Naive 攤回去：把所有文字塞回第一個 run，後續 run 清空
        # （犧牲 per-run 屬性精度，換中文段落乾淨 — 多數案例 run 屬性一致沒差）
        p.runs[0].text = cleaned
        for r in p.runs[1:]:
            r.text = ""
        spaces_removed += n
        pua_removed += pua_n
        paragraphs_touched += 1

    # autoSpaceDE / autoSpaceDN — 全文 default 開啟
    try:
        styles_element = docx_doc.styles.element
        doc_defaults = styles_element.find(qn("w:docDefaults"))
        if doc_defaults is not None:
            pPrDefault = doc_defaults.find(qn("w:pPrDefault"))
            if pPrDefault is None:
                pPrDefault = doc_defaults.makeelement(qn("w:pPrDefault"), {})
                doc_defaults.append(pPrDefault)
            pPr = pPrDefault.find(qn("w:pPr"))
            if pPr is None:
                pPr = pPrDefault.makeelement(qn("w:pPr"), {})
                pPrDefault.append(pPr)
            for tag in ("w:autoSpaceDE", "w:autoSpaceDN"):
                if pPr.find(qn(tag)) is None:
                    el = pPr.makeelement(qn(tag), {})
                    el.set(qn("w:val"), "1")
                    pPr.append(el)
    except Exception as e:
        log.debug("autoSpaceDE/DN set failed: %s", e)

    return {
        "fixer": "cjk_typography",
        "inter_cjk_spaces_removed": spaces_removed,
        "pua_glyphs_removed": pua_removed,
        "paragraphs_touched": paragraphs_touched,
    }
