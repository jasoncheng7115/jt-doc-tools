"""表格 cell 內重複段落 / 重複行去重（Sprint B 2 強化）。

pdf2docx 在處理用「stroke + fill 多層渲染粗體效果」的 PDF 時，會把同一段文字
重複塞進同一 cell 多次（在使用者看來就是字疊字 / 重複文字）。例如：

  原 PDF 一個 cell 顯示「技術服務」
  pdf2docx 抽出 cell 內含 2 個段落都是「技術服務」（粗體 stroke + fill 兩次）

extractor 端的 `_dedup_consecutive_lines` 在 PDFLine 級已處理過，但 pdf2docx
寫 docx 時不一定吃這層，docx cell 內仍可能重複。本 fixer 在 docx 層收尾：

- 對每個 table cell 內的段落集合掃重複（normalized 文字一致）
- 連續重複 → 保留第一個，後面的清空文字（保留 element 與 style 不破壞 cell 結構）
- 非連續重複 → 不動（可能是有意的同名兩列）
"""
from __future__ import annotations

import logging
import re

from docx.oxml.ns import qn

log = logging.getLogger(__name__)


def _normalize(s: str) -> str:
    if not s:
        return ""
    return re.sub(r"\s+", "", s)


def _paragraph_text(p_el) -> str:
    parts = []
    for r in p_el.findall(qn("w:r")):
        for t in r.findall(qn("w:t")):
            if t.text:
                parts.append(t.text)
    return "".join(parts)


def _clear_paragraph_text(p_el) -> None:
    """把段落內所有 w:t 文字清空，但保留段落 element + run 屬性（避免破壞表格
    結構與 row height）。"""
    for r in p_el.findall(qn("w:r")):
        for t in r.findall(qn("w:t")):
            t.text = ""


def _dedup_cell_paragraphs(cell) -> int:
    """對單一 cell 內段落掃連續重複文字 → 後者清空。回清掉的段落數。"""
    paras = cell.paragraphs
    if len(paras) < 2:
        return 0
    cleared = 0
    seen_prev = ""
    for p in paras:
        norm = _normalize(_paragraph_text(p._element))
        if not norm:
            seen_prev = ""
            continue
        if norm == seen_prev:
            _clear_paragraph_text(p._element)
            cleared += 1
        else:
            seen_prev = norm
    return cleared


def fix_table_cell_dedup_text(docx_doc, pdf_truth, alignment) -> dict:
    if not docx_doc.tables:
        return {"fixer": "table_cell_dedup_text", "cleared_paragraphs": 0,
                "tables": 0}
    cleared = 0
    cells_touched = 0
    for tbl in docx_doc.tables:
        seen_cells: set = set()
        for row in tbl.rows:
            for cell in row.cells:
                # merged cell 重覆映射到同一個 element，只處理一次
                cid = id(cell._element)
                if cid in seen_cells:
                    continue
                seen_cells.add(cid)
                n = _dedup_cell_paragraphs(cell)
                if n > 0:
                    cleared += n
                    cells_touched += 1
    return {
        "fixer": "table_cell_dedup_text",
        "cleared_paragraphs": cleared,
        "cells_touched": cells_touched,
        "tables": len(docx_doc.tables),
    }
