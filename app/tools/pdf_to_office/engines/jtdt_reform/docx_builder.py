"""DocumentModel → docx 寫入（B4）。

把 paragraph_grouper 建好的 DocumentModel 寫成 docx 檔。

特性：
- 表格按 region 寬高建表 → 套對應 cell_blocks 內容
- free 段落按 Y 序加入 body
- 圖片用 InlineShapes 補到對應頁段落
- 超連結用 add_hyperlink 套 link annotation 的 uri
- 頁面尺寸 / margins 從 PDFTruth 第一頁拿（後續頁不同要 split section，這版簡化）
- CJK 字型用 body_font_name 或 fallback Noto Sans CJK TC
- 表格 cell vAlign 統一 center；無框線 (drawings 沒線) 的 cell 不畫 border

座標單位：PDF 點 → docx EMU (1 pt = 12700 EMU) 圖片用；docx 邊距用 Inches/Cm。
"""
from __future__ import annotations

import logging
from io import BytesIO
from pathlib import Path

import fitz  # PyMuPDF
from docx import Document
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Emu, Pt

from .paragraph_grouper import DocumentModel, FreeBlock, PageModel, TableModel

log = logging.getLogger(__name__)


EMU_PER_PT = 12700
HYPERLINK_COLOR = "0563C1"


def _block_aggregated_text(block) -> str:
    """合併 block 內所有 lines 文字（line 間以 \\n）"""
    if not block.lines:
        return (block.text or "").strip()
    return "\n".join((ln.text or "") for ln in block.lines if ln.text).strip()


def _block_dom_font_size(block) -> float:
    return float(block.dominant_size or 0)


def _block_dom_font_name(block) -> str:
    return (block.dominant_font or "").lstrip("+")


def _set_page_geometry(doc: Document, page_w_pt: float, page_h_pt: float,
                        margin_left_pt: float = 0, margin_top_pt: float = 0,
                        margin_right_pt: float = 0,
                        margin_bottom_pt: float = 0) -> None:
    """套頁面尺寸 + 邊距（從 PDFTruth 取），到 sectPr。"""
    sect = doc.sections[0]
    sect.page_width = Emu(int(page_w_pt * EMU_PER_PT))
    sect.page_height = Emu(int(page_h_pt * EMU_PER_PT))
    # 邊距：若 PDFTruth 沒給有效值（< 5pt）退回 1cm 預設
    _MIN_VALID_MARGIN = 5.0    # < 5pt 視為 PDFTruth 沒抓到，用預設
    _MAX_VALID_MARGIN_R = 0.4  # > 頁尺寸 40% 視為異常
    def _pick(m_pt, page_size_pt, default_cm=1.0):
        if m_pt < _MIN_VALID_MARGIN or m_pt > page_size_pt * _MAX_VALID_MARGIN_R:
            return Cm(default_cm)
        return Emu(int(m_pt * EMU_PER_PT))
    sect.left_margin = _pick(margin_left_pt, page_w_pt)
    sect.right_margin = _pick(margin_right_pt, page_w_pt)
    sect.top_margin = _pick(margin_top_pt, page_h_pt)
    sect.bottom_margin = _pick(margin_bottom_pt, page_h_pt)


def _wrap_run_as_hyperlink(paragraph, run_text: str, uri: str,
                            size_pt: float = 0, font_name: str = "") -> None:
    """在 paragraph 內加一段 w:hyperlink — 透過 doc.part 註冊 relationship。
    回傳 None，已寫入 paragraph。"""
    from docx.oxml import OxmlElement
    from docx.opc.constants import RELATIONSHIP_TYPE as RT
    part = paragraph.part
    r_id = part.relate_to(uri, RT.HYPERLINK, is_external=True)
    hyper = OxmlElement("w:hyperlink")
    hyper.set(qn("r:id"), r_id)
    new_r = OxmlElement("w:r")
    rPr = OxmlElement("w:rPr")
    color = OxmlElement("w:color")
    color.set(qn("w:val"), HYPERLINK_COLOR)
    rPr.append(color)
    u = OxmlElement("w:u")
    u.set(qn("w:val"), "single")
    rPr.append(u)
    if size_pt > 0:
        sz = OxmlElement("w:sz")
        sz.set(qn("w:val"), str(int(size_pt * 2)))
        rPr.append(sz)
    if font_name:
        # normalize 同 _set_run_font — 避免 PDF embedded font 名稱在客戶端 OS 找不到
        ascii_f, ea_f = _normalize_font_name(font_name) if "_normalize_font_name" in globals() else (font_name, font_name)
        rFonts = OxmlElement("w:rFonts")
        rFonts.set(qn("w:ascii"), ascii_f)
        rFonts.set(qn("w:hAnsi"), ascii_f)
        rFonts.set(qn("w:eastAsia"), ea_f)
        rPr.append(rFonts)
    new_r.append(rPr)
    t = OxmlElement("w:t")
    t.set(qn("xml:space"), "preserve")
    t.text = run_text
    new_r.append(t)
    hyper.append(new_r)
    paragraph._element.append(hyper)


_BOLD_NAME_HINTS = ("bold", "black", "heavy", "demibold", "semibold")
_ITALIC_NAME_HINTS = ("italic", "oblique")


def _font_name_implies_bold_italic(name: str) -> tuple[bool, bool]:
    """從 font name 推 bold / italic — 通用 PDF 字型命名規則 (Helvetica-Bold 等)。"""
    n = (name or "").lower()
    if not n:
        return False, False
    bold = any(s in n for s in _BOLD_NAME_HINTS)
    italic = any(s in n for s in _ITALIC_NAME_HINTS)
    return bold, italic


def _line_dominant_bold_italic(line) -> tuple[bool, bool]:
    """從 line.chars 推 is_bold / is_italic — 結合 font flag bit + font name hint。"""
    if not getattr(line, "chars", None):
        # fallback：用 line.dominant_font 推
        return _font_name_implies_bold_italic(line.dominant_font or "")
    n = len(line.chars)
    if n == 0:
        return _font_name_implies_bold_italic(line.dominant_font or "")
    # 投票：flag bit 或 font name 任一 → bold
    n_bold = 0
    n_italic = 0
    for c in line.chars:
        flag_b = bool(getattr(c, "is_bold", False))
        flag_i = bool(getattr(c, "is_italic", False))
        name_b, name_i = _font_name_implies_bold_italic(getattr(c, "font_name", "") or "")
        if flag_b or name_b:
            n_bold += 1
        if flag_i or name_i:
            n_italic += 1
    # 30% 閾值（多數投票 50% 對「中英混排 + 數字 + 空格」line 太嚴）
    return (n_bold >= max(1, n // 3)), (n_italic >= max(1, n // 3))


def _line_dominant_color(line) -> str:
    """從 line.chars 取多數 color hex。回 '' 表示無資料 / 黑色（不用設）。"""
    if not getattr(line, "chars", None):
        return ""
    from collections import Counter
    cnt = Counter()
    for ch in line.chars:
        c = (ch.color or "").lstrip("#")
        if c:
            cnt[c] += 1
    if not cnt:
        return ""
    most = cnt.most_common(1)[0][0]
    # 黑色 (000000) 是預設，不用設
    if most.lower() in ("000000", "000"):
        return ""
    return most


def _normalize_font_name(font_name: str) -> tuple[str, str]:
    """把 PDF embedded font (TT*) 之類非標準字型替換成系統通用 CJK 字型。

    回 (ascii_font, eastAsia_font)。PDF 內 embedded font 之名稱如 "TT6120o00"
    在 user 系統不存在，docx render 會 fallback 字型且 metric 不同，cell 寬會被
    擠縮（fixed layout 也救不了）— 直接用通用名稱避免此問題。

    通用名稱：'Times New Roman' (Western) + 'Noto Sans CJK TC' (CJK)。
    這些名稱在大部分 OS（Windows / macOS / Linux）有 fallback 機制，不會 render
    錯亂。
    """
    f = (font_name or "").lstrip("+").strip()
    DEFAULT_W = "Times New Roman"
    DEFAULT_CJK = "Noto Sans CJK TC"
    if not f:
        return DEFAULT_W, DEFAULT_CJK
    # PDF embedded font 特徵：TT 開頭 + 數字 / 純亂碼短名 / 含「o00」之類
    fl = f.lower()
    is_pdf_embedded = (
        f.startswith("TT") and any(c.isdigit() for c in f)
        or "o00" in fl
        or (len(f) <= 12 and f[:1].isupper()
            and not any(k in fl for k in ("arial", "times", "courier", "noto",
                                            "ming", "song", "kai", "hei", "yi",
                                            "yuan", "fang", "jheng", "ping",
                                            "helvet", "verd", "georg")))
    )
    if is_pdf_embedded:
        return DEFAULT_W, DEFAULT_CJK
    return f, f


def _set_run_font(run, size_pt: float, font_name: str, color_hex: str = "",
                   bold: bool = False, italic: bool = False) -> None:
    if size_pt > 0:
        run.font.size = Pt(size_pt)
    if font_name:
        ascii_font, ea_font = _normalize_font_name(font_name)
        run.font.name = ascii_font
        rPr = run._element.get_or_add_rPr()
        rFonts = rPr.find(qn("w:rFonts"))
        if rFonts is None:
            rFonts = rPr.makeelement(qn("w:rFonts"), {})
            rPr.append(rFonts)
        rFonts.set(qn("w:ascii"), ascii_font)
        rFonts.set(qn("w:hAnsi"), ascii_font)
        rFonts.set(qn("w:eastAsia"), ea_font)
    if color_hex:
        try:
            from docx.shared import RGBColor
            r = int(color_hex[0:2], 16)
            g = int(color_hex[2:4], 16)
            b = int(color_hex[4:6], 16)
            run.font.color.rgb = RGBColor(r, g, b)
        except Exception:
            pass
    if bold:
        run.bold = True
    if italic:
        run.italic = True


def _infer_line_alignment(line, page_width: float) -> str | None:
    """從 line bbox X 中心 vs 頁中軸推 alignment。回 'center'/'right'/'left'/None。"""
    if page_width <= 0 or not line.bbox:
        return None
    x0, _, x1, _ = line.bbox
    cx = (x0 + x1) / 2.0
    page_cx = page_width / 2.0
    offset_ratio = (cx - page_cx) / page_width  # -0.5 .. 0.5
    # line 寬度（X 範圍 / 頁寬）— short line + 居中可信度高
    line_width_ratio = (x1 - x0) / page_width
    if abs(offset_ratio) < 0.05 and line_width_ratio < 0.7:
        return "center"
    if offset_ratio > 0.20 and line_width_ratio < 0.5:
        return "right"
    # 預設 left（不顯式設，讓 docx 用預設）
    return None


def _apply_paragraph_alignment(p, align: str) -> None:
    from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
    mapping = {
        "center": WD_PARAGRAPH_ALIGNMENT.CENTER,
        "right": WD_PARAGRAPH_ALIGNMENT.RIGHT,
        "left": WD_PARAGRAPH_ALIGNMENT.LEFT,
    }
    target = mapping.get(align)
    if target is not None:
        try:
            p.alignment = target
        except Exception:
            pass


def _add_paragraph_for_block(doc: Document, block, is_heading: bool = False,
                              is_footer: bool = False,
                              page_width: float = 0.0) -> list:
    """為 free block 加 docx 段落，**每 PDF line 一段**。回新增的 paragraph list。"""
    out = []
    lines = list(block.lines) if block.lines else []
    if not lines:
        text = (block.text or "").strip()
        if not text:
            return out
        p = doc.add_paragraph()
        if is_heading:
            try:
                p.style = doc.styles["Heading 2"]
            except KeyError:
                pass
        run = p.add_run(text)
        _set_run_font(run, _block_dom_font_size(block), _block_dom_font_name(block))
        out.append(p)
        return out
    for ln in lines:
        txt = (ln.text or "").strip()
        if not txt:
            continue
        p = doc.add_paragraph()
        if is_heading and ln is lines[0]:
            try:
                p.style = doc.styles["Heading 2"]
            except KeyError:
                pass
        # 對齊推算（用本 line 的 bbox 中心 vs 頁寬）
        align = _infer_line_alignment(ln, page_width)
        if align:
            _apply_paragraph_alignment(p, align)
        link = _find_link_for_line(ln)
        size_pt = float(ln.dominant_size or 0)
        font_name = (ln.dominant_font or "").lstrip("+")
        color_hex = _line_dominant_color(ln)
        if link and link.get("uri"):
            try:
                _wrap_run_as_hyperlink(p, txt, link["uri"],
                                         size_pt=size_pt, font_name=font_name)
                out.append(p)
                continue
            except Exception as e:
                log.debug("wrap hyperlink (free para) failed: %s", e)
        run = p.add_run(txt)
        bold, italic = _line_dominant_bold_italic(ln)
        _set_run_font(run, size_pt, font_name, color_hex=color_hex,
                       bold=bold, italic=italic)
        out.append(p)
    return out


# OOXML CT_TcPrInner schema 順序 — vAlign / tcBorders / shd / gridSpan 等都要照此排
# 順序錯 OxOffice / LibreOffice 不認 cell width，會 auto-shrink cell 至 content 寬。
_TCPR_CHILD_ORDER = [
    "cnfStyle", "tcW", "gridSpan", "hMerge", "vMerge",
    "tcBorders", "shd", "noWrap", "tcMar",
    "textDirection", "tcFitText", "vAlign", "hideMark", "headers",
]


def _reorder_tcPr(tcPr) -> None:
    """把 tcPr 子元素 reorder 到 OOXML schema 規範順序。"""
    if tcPr is None:
        return
    existing = list(tcPr)
    for ch in existing:
        tcPr.remove(ch)
    def _idx(elem):
        tag = elem.tag.split("}")[-1]
        return _TCPR_CHILD_ORDER.index(tag) if tag in _TCPR_CHILD_ORDER else 999
    for ch in sorted(existing, key=_idx):
        tcPr.append(ch)


def _set_cell_vertical_center(cell) -> None:
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    # python-docx 把 vAlign 直接 append，可能順序錯 — reorder 確保 schema 合規
    tcPr = cell._element.find(qn("w:tcPr"))
    _reorder_tcPr(tcPr)


def _set_cell_shading(cell, fill_hex: str) -> None:
    """設 cell 背景色（從 PDF rect fill 推）。空 / 白色 / 接近白 → 不設。"""
    h = (fill_hex or "").lstrip("#").lower()
    if not h or len(h) < 6:
        return
    try:
        r = int(h[0:2], 16)
        g = int(h[2:4], 16)
        b = int(h[4:6], 16)
    except Exception:
        return
    # 接近白 (250,250,250+) 不設 — 視覺等同無底色
    if r >= 250 and g >= 250 and b >= 250:
        return
    tcPr = cell._element.find(qn("w:tcPr"))
    if tcPr is None:
        tcPr = cell._element.makeelement(qn("w:tcPr"), {})
        cell._element.insert(0, tcPr)
    for old in tcPr.findall(qn("w:shd")):
        tcPr.remove(old)
    shd = tcPr.makeelement(qn("w:shd"), {})
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), h)
    tcPr.append(shd)
    _reorder_tcPr(tcPr)


def _set_cell_borders(cell, has_border: bool = True,
                       color_hex: str = "", width_pt: float = 0.0) -> None:
    tcPr = cell._element.find(qn("w:tcPr"))
    if tcPr is None:
        tcPr = cell._element.makeelement(qn("w:tcPr"), {})
        cell._element.insert(0, tcPr)
    for old in tcPr.findall(qn("w:tcBorders")):
        tcPr.remove(old)
    tcB = tcPr.makeelement(qn("w:tcBorders"), {})
    val = "single" if has_border else "nil"
    # 用 PDFTruth.drawings 抽出的 dominant stroke color；空 / 黑 → 預設淺灰
    use_color = (color_hex or "").lstrip("#").lower()
    if not use_color or use_color in ("000000", "000"):
        use_color = "999999"
    # OOXML w:sz 單位是 1/8 pt；PDF stroke_width 為 pt
    # 預設 0.5pt = sz="4"；常見 invoice 用 0.25pt sz="2"
    if width_pt > 0:
        sz_val = max(2, min(48, int(round(width_pt * 8))))  # 限 [2, 48] = 0.25-6pt
    else:
        sz_val = 4
    for tag in ("top", "left", "bottom", "right"):
        el = tcB.makeelement(qn(f"w:{tag}"), {})
        el.set(qn("w:val"), val)
        if has_border:
            el.set(qn("w:sz"), str(sz_val))
            el.set(qn("w:color"), use_color)
        tcB.append(el)
    tcPr.append(tcB)
    _reorder_tcPr(tcPr)


def _add_table(doc: Document, table_model: TableModel) -> object:
    """加表格 — 用 region.cells 形狀建 docx table。"""
    n_rows = len(table_model.cell_blocks)
    n_cols = len(table_model.cell_blocks[0]) if n_rows > 0 else 0
    if n_rows == 0 or n_cols == 0:
        return None
    tbl = doc.add_table(rows=n_rows, cols=n_cols)
    # 表格 layout: fixed (用 PDF 真值寬度，不讓 docx 自動 resize)
    tblPr = tbl._element.find(qn("w:tblPr"))
    if tblPr is None:
        tblPr = tbl._element.makeelement(qn("w:tblPr"), {})
        tbl._element.insert(0, tblPr)
    # 清掉舊 tblW / tblLayout 重建
    for old in tblPr.findall(qn("w:tblLayout")):
        tblPr.remove(old)
    for old in tblPr.findall(qn("w:tblW")):
        tblPr.remove(old)
    col_xs = table_model.region.col_xs
    col_widths_pt = [col_xs[i + 1] - col_xs[i] for i in range(len(col_xs) - 1)]
    total_w_dxa = int(sum(col_widths_pt) * 20)
    tblW = tblPr.makeelement(qn("w:tblW"), {})
    tblW.set(qn("w:type"), "dxa")
    tblW.set(qn("w:w"), str(total_w_dxa))
    layout = tblPr.makeelement(qn("w:tblLayout"), {})
    layout.set(qn("w:type"), "fixed")
    # **OOXML schema 強制順序**：tblW (idx 7) → tblLayout (idx 12) → tblLook (idx 14)
    # 若 tblLayout 出現在 tblW 之前，OxOffice / LibreOffice 不認 fixed layout 會
    # 自動縮 cell width 至 content 寬度。正確順序透過 reorder 子元素確保:
    _CHILD_ORDER = [
        "tblStyle", "tblpPr", "tblOverlap", "bidiVisual",
        "tblStyleRowBandSize", "tblStyleColBandSize",
        "tblW", "jc", "tblCellSpacing", "tblInd", "tblBorders",
        "tblLayout", "tblCellMar", "tblLook",
        "tblCaption", "tblDescription",
    ]
    # 加入 tblW + tblLayout，然後對所有子元素排序
    tblPr.append(tblW)
    tblPr.append(layout)
    _existing = list(tblPr)
    for ch in _existing:
        tblPr.remove(ch)
    def _child_idx(elem):
        tag = elem.tag.split("}")[-1]
        return _CHILD_ORDER.index(tag) if tag in _CHILD_ORDER else 999
    for ch in sorted(_existing, key=_child_idx):
        tblPr.append(ch)
    tbl_el = tbl._element
    # 移除舊 tblGrid
    for old in tbl_el.findall(qn("w:tblGrid")):
        tbl_el.remove(old)
    grid = tbl_el.makeelement(qn("w:tblGrid"), {})
    for w_pt in col_widths_pt:
        gc = grid.makeelement(qn("w:gridCol"), {qn("w:w"): str(int(w_pt * 20))})
        grid.append(gc)
    # tblPr 之後插入 grid
    tblPr.addnext(grid)

    # 套 row 高度（從 row_ys 推 pt → docx trHeight twips）
    row_ys = table_model.region.row_ys
    for r in range(n_rows):
        if r + 1 < len(row_ys):
            row_h_pt = row_ys[r + 1] - row_ys[r]
            if row_h_pt > 0:
                tr_el = tbl.rows[r]._element
                trPr = tr_el.find(qn("w:trPr"))
                if trPr is None:
                    trPr = tr_el.makeelement(qn("w:trPr"), {})
                    tr_el.insert(0, trPr)
                # 移除舊 trHeight
                for old in trPr.findall(qn("w:trHeight")):
                    trPr.remove(old)
                trHeight = trPr.makeelement(qn("w:trHeight"), {})
                # at least: 內容可長可短，但 row 不小於 PDF 真值
                trHeight.set(qn("w:val"), str(int(row_h_pt * 20)))
                trHeight.set(qn("w:hRule"), "atLeast")
                trPr.append(trHeight)

    # vmerge / hmerge table
    vmerge = getattr(table_model.region, "vmerge", None) or []
    hmerge = getattr(table_model.region, "hmerge", None) or []

    # 處理 hMerge — 對每 row，先依 hmerge 移除被合併的 tc element（保留 first，
    # 在 first 加 <w:gridSpan>）；merge 完 row.cells 數量 reduce 對應 gridSpan
    # 注意：python-docx Cell 物件對 hMerge 的支援差，直接動 XML 比較穩
    if hmerge:
        for r in range(n_rows):
            if r >= len(hmerge):
                continue
            tr_el = tbl.rows[r]._element
            tcs = tr_el.findall(qn("w:tc"))
            # 從右往左移除被合併 (hmerge[r][c] == 0)，並在合併 first 加 gridSpan
            for c in range(n_cols - 1, -1, -1):
                if c >= len(hmerge[r]):
                    continue
                if hmerge[r][c] == 0:
                    # 該 tc 被合進左側 — 移除
                    if c < len(tcs):
                        tr_el.remove(tcs[c])
                elif hmerge[r][c] > 1:
                    # 該 tc 跨 hmerge[r][c] 欄
                    if c < len(tcs):
                        tc = tcs[c]
                        tcPr = tc.find(qn("w:tcPr"))
                        if tcPr is None:
                            tcPr = tc.makeelement(qn("w:tcPr"), {})
                            tc.insert(0, tcPr)
                        for old in tcPr.findall(qn("w:gridSpan")):
                            tcPr.remove(old)
                        gs = tcPr.makeelement(qn("w:gridSpan"), {})
                        gs.set(qn("w:val"), str(hmerge[r][c]))
                        tcPr.append(gs)
                        # **tcW 必須等於跨欄總寬**（OOXML 規範）— 否則 OxOffice / Word
                        # 看到 tcW vs gridSpan 矛盾會 fallback auto layout，整個 table
                        # col widths 被縮（user 之前看到 col 全變直書即此原因）。
                        span_w_dxa = sum(int(col_widths_pt[cc] * 20)
                                          for cc in range(c, min(c + hmerge[r][c], n_cols)))
                        for old in tcPr.findall(qn("w:tcW")):
                            tcPr.remove(old)
                        tcW = tcPr.makeelement(qn("w:tcW"), {})
                        tcW.set(qn("w:type"), "dxa")
                        tcW.set(qn("w:w"), str(span_w_dxa))
                        tcPr.append(tcW)
                        _reorder_tcPr(tcPr)

    # 填 cell 內容（注意：hMerge 後 row.cells 可能變少）
    for r in range(n_rows):
        # 重新從 row 取 tc — hmerge 可能已移除某些 tc
        tr_el = tbl.rows[r]._element
        tcs_now = tr_el.findall(qn("w:tc"))
        # 每個 tc 對應 hmerge 後的 visible col
        # 取 hmerge[r] 內 hmerge[r][c] > 0 的 c index list
        visible_cols = []
        if hmerge and r < len(hmerge):
            visible_cols = [c for c in range(n_cols) if hmerge[r][c] > 0]
        else:
            visible_cols = list(range(n_cols))
        # 直接從 tcs_now 取對應的 tc element 建 _Cell wrapper
        # （Row.cells tuple 在 hmerge 後對 spanned cell 重複指向 same wrapper —
        # 不能用 cells[c] 或 cells[tc_idx] 取，會寫入錯位）
        from docx.table import _Cell as _DocxCell
        for tc_idx, c in enumerate(visible_cols):
            if tc_idx >= len(tcs_now):
                break
            cell = _DocxCell(tcs_now[tc_idx], tbl.rows[r])
            blocks_in_cell = table_model.cell_blocks[r][c]
            # 取該 cell PDF bbox 推 line alignment 用
            cell_bbox_pdf = None
            try:
                cell_bbox_pdf = table_model.region.cells[r][c]
            except Exception:
                pass
            # 清原 paragraph
            for p in cell.paragraphs:
                for run in p._element.findall(qn("w:r")):
                    p._element.remove(run)
            if blocks_in_cell:
                # 把所有 blocks 的所有 lines 拍平成 line list，每 line 一個段落
                first_para = True
                for blk in blocks_in_cell:
                    lines = list(blk.lines) if blk.lines else []
                    if not lines:
                        txt = (blk.text or "").strip()
                        if txt:
                            p = cell.paragraphs[0] if first_para else cell.add_paragraph()
                            first_para = False
                            run = p.add_run(txt)
                            sz = _block_dom_font_size(blk)
                            if sz > 0:
                                run.font.size = Pt(sz)
                        continue
                    for ln in lines:
                        ln_text = (ln.text or "").strip()
                        if not ln_text:
                            continue
                        p = cell.paragraphs[0] if first_para else cell.add_paragraph()
                        first_para = False
                        # 推 cell 內段落 alignment (line bbox vs cell bbox)
                        if cell_bbox_pdf:
                            cx0, _, cx1, _ = cell_bbox_pdf
                            lx0, _, lx1, _ = ln.bbox
                            cell_w = cx1 - cx0
                            line_w = lx1 - lx0
                            if cell_w > 0 and line_w < cell_w * 0.9:
                                # cell 寬比 line 寬，看 line 偏左 / 中 / 右
                                line_cx = (lx0 + lx1) / 2.0
                                cell_cx = (cx0 + cx1) / 2.0
                                offset = (line_cx - cell_cx) / cell_w
                                if offset > 0.10:
                                    _apply_paragraph_alignment(p, "right")
                                elif offset < -0.10:
                                    _apply_paragraph_alignment(p, "left")
                                elif abs(offset) <= 0.08:
                                    _apply_paragraph_alignment(p, "center")
                        link = _find_link_for_line(ln)
                        size_pt = float(ln.dominant_size or 0)
                        color_hex = _line_dominant_color(ln)
                        if link and link.get("uri"):
                            try:
                                _wrap_run_as_hyperlink(
                                    p, ln_text, link["uri"],
                                    size_pt=size_pt,
                                    font_name=(ln.dominant_font or "").lstrip("+"),
                                )
                                continue
                            except Exception as e:
                                log.debug("wrap hyperlink (cell) failed: %s", e)
                        run = p.add_run(ln_text)
                        bold, italic = _line_dominant_bold_italic(ln)
                        _set_run_font(run, size_pt,
                                       (ln.dominant_font or "").lstrip("+"),
                                       color_hex=color_hex,
                                       bold=bold, italic=italic)
            _set_cell_vertical_center(cell)
            # cell 背景色（從 PDF rect fill 推）
            try:
                fill_hex = table_model.cell_fill[r][c]
                _set_cell_shading(cell, fill_hex)
            except Exception:
                pass
            # virtual table（無框線推出來）— 不畫框，避免「<樣本 A>上方 date row 多加框」
            is_virtual_tbl = getattr(table_model.region, "is_virtual", False)
            _set_cell_borders(cell,
                              has_border=not is_virtual_tbl,
                              color_hex=getattr(table_model.region,
                                                  "border_color_hex", ""),
                              width_pt=getattr(table_model.region,
                                                "border_width_pt", 0.0))
            # vMerge：在 tcPr 內加 <w:vMerge w:val="restart"/> 或 <w:vMerge/>
            if vmerge and r < len(vmerge) and c < len(vmerge[r]):
                vm_state = vmerge[r][c]
                if vm_state in ("restart", "continue"):
                    tcPr = cell._element.find(qn("w:tcPr"))
                    if tcPr is None:
                        tcPr = cell._element.makeelement(qn("w:tcPr"), {})
                        cell._element.insert(0, tcPr)
                    for old in tcPr.findall(qn("w:vMerge")):
                        tcPr.remove(old)
                    vme = tcPr.makeelement(qn("w:vMerge"), {})
                    if vm_state == "restart":
                        vme.set(qn("w:val"), "restart")
                    tcPr.append(vme)
            # cell tcW 按 col_widths_pt 設 — **gridSpan > 1 時必須是跨欄總寬**
            # 否則 OxOffice/Word 看到 tcW vs gridSpan 矛盾會 fallback auto layout
            tcPr = cell._element.find(qn("w:tcPr"))
            if tcPr is None:
                tcPr = cell._element.makeelement(qn("w:tcPr"), {})
                cell._element.insert(0, tcPr)
            span = (hmerge[r][c] if hmerge and r < len(hmerge)
                     and c < len(hmerge[r]) and hmerge[r][c] > 0 else 1)
            span_w_pt = sum(col_widths_pt[cc]
                              for cc in range(c, min(c + span, n_cols)))
            tcW = tcPr.find(qn("w:tcW"))
            if tcW is None:
                tcW = tcPr.makeelement(qn("w:tcW"), {})
                tcPr.append(tcW)
            tcW.set(qn("w:type"), "dxa")
            tcW.set(qn("w:w"), str(int(span_w_pt * 20)))
            _reorder_tcPr(tcPr)
    return tbl


def _add_floating_polygon_image(doc: Document, x_pt: float, y_pt: float,
                                  w_pt: float, h_pt: float,
                                  fill_hex: str,
                                  path_points: list) -> None:
    """用 PIL 把實際 polygon 形狀渲染成 PNG（透明背景 + fill 色）→ 加 wp:anchor
    浮動圖。比 _add_floating_color_image 多保留 banner 的切角 / 圓角等真實形狀。
    """
    color = (fill_hex or "").lstrip("#").lower()
    if not color or len(color) < 6 or color in ("ffffff", "fff"):
        return
    try:
        r = int(color[0:2], 16)
        g = int(color[2:4], 16)
        b = int(color[4:6], 16)
    except Exception:
        return
    if w_pt <= 0 or h_pt <= 0 or len(path_points) < 3:
        return
    try:
        from PIL import Image, ImageDraw
        # 渲染解析度：1pt=4px (高 DPI 對 LibreOffice / Word 都 OK)
        SCALE = 4
        canvas_w = max(8, int(w_pt * SCALE))
        canvas_h = max(8, int(h_pt * SCALE))
        img = Image.new("RGBA", (canvas_w, canvas_h), (0, 0, 0, 0))
        d = ImageDraw.Draw(img)
        # path_points 是 PDF 座標（相對 page 原點），轉成 PNG 內相對座標
        pts = [(int((px - x_pt) * SCALE), int((py - y_pt) * SCALE))
                for (px, py) in path_points]
        d.polygon(pts, fill=(r, g, b, 255))
        png_io = BytesIO()
        img.save(png_io, format="PNG")
        png_io.seek(0)
        _insert_floating_image_anchor(doc, png_io, x_pt, y_pt, w_pt, h_pt)
    except Exception as e:
        log.debug("polygon banner png failed: %s", e)


def _insert_floating_image_anchor(doc: Document, png_io: BytesIO,
                                    x_pt: float, y_pt: float,
                                    w_pt: float, h_pt: float) -> None:
    """把 PNG io 以 wp:anchor (behindDoc=1) 浮動圖 insert，位置 relative to page。
    抽出 _add_floating_color_image 內 anchor 建構邏輯供 polygon 共用。"""
    from docx.oxml import OxmlElement
    cx_emu = max(1, int(w_pt * EMU_PER_PT))
    cy_emu = max(1, int(h_pt * EMU_PER_PT))
    x_emu = int(x_pt * EMU_PER_PT)
    y_emu = int(y_pt * EMU_PER_PT)
    p = doc.add_paragraph()
    p_pPr = p._element.get_or_add_pPr()
    sp = OxmlElement("w:spacing")
    sp.set(qn("w:before"), "0"); sp.set(qn("w:after"), "0")
    sp.set(qn("w:line"), "0"); sp.set(qn("w:lineRule"), "atLeast")
    p_pPr.append(sp)
    run = p.add_run()
    run.add_picture(png_io, width=Emu(cx_emu), height=Emu(cy_emu))
    drawing_el = run._element.find(qn("w:drawing"))
    if drawing_el is None:
        return
    NS_WP = "http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing"
    NS_A = "http://schemas.openxmlformats.org/drawingml/2006/main"
    inline = drawing_el.find(f"{{{NS_WP}}}inline")
    if inline is None:
        return
    graphic = inline.find(f"{{{NS_A}}}graphic")
    docPr = inline.find(f"{{{NS_WP}}}docPr")
    cNvGraphicFramePr = inline.find(f"{{{NS_WP}}}cNvGraphicFramePr")
    anchor = drawing_el.makeelement(f"{{{NS_WP}}}anchor", {
        "distT": "0", "distB": "0", "distL": "0", "distR": "0",
        "simplePos": "0", "relativeHeight": "1",
        "behindDoc": "1", "locked": "0", "layoutInCell": "1",
        "allowOverlap": "1",
    })
    anchor.append(anchor.makeelement(f"{{{NS_WP}}}simplePos", {"x": "0", "y": "0"}))
    posH = anchor.makeelement(f"{{{NS_WP}}}positionH", {"relativeFrom": "page"})
    ph_off = posH.makeelement(f"{{{NS_WP}}}posOffset", {})
    ph_off.text = str(x_emu)
    posH.append(ph_off)
    anchor.append(posH)
    posV = anchor.makeelement(f"{{{NS_WP}}}positionV", {"relativeFrom": "page"})
    pv_off = posV.makeelement(f"{{{NS_WP}}}posOffset", {})
    pv_off.text = str(y_emu)
    posV.append(pv_off)
    anchor.append(posV)
    extent = anchor.makeelement(f"{{{NS_WP}}}extent",
                                   {"cx": str(cx_emu), "cy": str(cy_emu)})
    anchor.append(extent)
    anchor.append(anchor.makeelement(f"{{{NS_WP}}}effectExtent",
                                       {"l": "0", "t": "0", "r": "0", "b": "0"}))
    anchor.append(anchor.makeelement(f"{{{NS_WP}}}wrapNone", {}))
    if docPr is not None:
        anchor.append(docPr)
    if cNvGraphicFramePr is not None:
        anchor.append(cNvGraphicFramePr)
    if graphic is not None:
        anchor.append(graphic)
    drawing_el.remove(inline)
    drawing_el.append(anchor)


def _add_floating_color_image(doc: Document, x_pt: float, y_pt: float,
                                w_pt: float, h_pt: float,
                                fill_hex: str) -> None:
    """用 PIL 生純色 PNG → 加 wp:anchor 浮動圖（PDF banner / 色塊）。
    比 wps:wsp shape 在 LibreOffice / Word 都更可靠。"""
    color = (fill_hex or "").lstrip("#").lower()
    if not color or len(color) < 6:
        return
    if color in ("ffffff", "fff"):
        return
    try:
        r = int(color[0:2], 16)
        g = int(color[2:4], 16)
        b = int(color[4:6], 16)
    except Exception:
        return
    cx_emu = max(1, int(w_pt * EMU_PER_PT))
    cy_emu = max(1, int(h_pt * EMU_PER_PT))
    x_emu = int(x_pt * EMU_PER_PT)
    y_emu = int(y_pt * EMU_PER_PT)
    try:
        from PIL import Image
        # 用 8×8 PNG，由 Word/LibreOffice stretch 到 anchor 大小
        img = Image.new("RGB", (8, 8), color=(r, g, b))
        png_io = BytesIO()
        img.save(png_io, format="PNG")
        png_io.seek(0)
        p = doc.add_paragraph()
        # spacing 0 避免段落本身佔位
        from docx.oxml import OxmlElement
        p_pPr = p._element.get_or_add_pPr()
        sp = OxmlElement("w:spacing")
        sp.set(qn("w:before"), "0"); sp.set(qn("w:after"), "0")
        sp.set(qn("w:line"), "0"); sp.set(qn("w:lineRule"), "atLeast")
        p_pPr.append(sp)
        run = p.add_run()
        run.add_picture(png_io, width=Emu(cx_emu), height=Emu(cy_emu))
        drawing_el = run._element.find(qn("w:drawing"))
        if drawing_el is None:
            return
        # inline → anchor (behindDoc=1)
        NS_WP = "http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing"
        NS_A = "http://schemas.openxmlformats.org/drawingml/2006/main"
        inline = drawing_el.find(f"{{{NS_WP}}}inline")
        if inline is None:
            return
        graphic = inline.find(f"{{{NS_A}}}graphic")
        docPr = inline.find(f"{{{NS_WP}}}docPr")
        cNvGraphicFramePr = inline.find(f"{{{NS_WP}}}cNvGraphicFramePr")
        anchor = drawing_el.makeelement(f"{{{NS_WP}}}anchor", {
            "distT": "0", "distB": "0", "distL": "0", "distR": "0",
            "simplePos": "0", "relativeHeight": "1",
            "behindDoc": "1", "locked": "0", "layoutInCell": "1",
            "allowOverlap": "1",
        })
        anchor.append(anchor.makeelement(f"{{{NS_WP}}}simplePos",
                                          {"x": "0", "y": "0"}))
        posH = anchor.makeelement(f"{{{NS_WP}}}positionH",
                                   {"relativeFrom": "page"})
        ph_off = posH.makeelement(f"{{{NS_WP}}}posOffset", {})
        ph_off.text = str(x_emu)
        posH.append(ph_off)
        anchor.append(posH)
        posV = anchor.makeelement(f"{{{NS_WP}}}positionV",
                                   {"relativeFrom": "page"})
        pv_off = posV.makeelement(f"{{{NS_WP}}}posOffset", {})
        pv_off.text = str(y_emu)
        posV.append(pv_off)
        anchor.append(posV)
        anchor.append(anchor.makeelement(f"{{{NS_WP}}}extent",
                                          {"cx": str(cx_emu), "cy": str(cy_emu)}))
        anchor.append(anchor.makeelement(f"{{{NS_WP}}}effectExtent",
                                          {"l": "0", "t": "0", "r": "0", "b": "0"}))
        anchor.append(anchor.makeelement(f"{{{NS_WP}}}wrapNone", {}))
        if docPr is not None:
            anchor.append(docPr)
        if cNvGraphicFramePr is not None:
            anchor.append(cNvGraphicFramePr)
        if graphic is not None:
            anchor.append(graphic)
        drawing_el.remove(inline)
        drawing_el.append(anchor)
    except Exception as e:
        log.debug("add floating color image failed: %s", e)


def _add_floating_rect_shape(doc: Document, x_pt: float, y_pt: float,
                               w_pt: float, h_pt: float,
                               fill_hex: str) -> None:
    """加 wp:anchor floating rect shape — behindDoc=1, wrapNone (PDF banner / 色塊)。"""
    from docx.oxml import OxmlElement
    color = (fill_hex or "").lstrip("#").lower()
    if not color or len(color) < 6:
        return
    if color in ("ffffff", "fff"):
        return
    x_emu = int(x_pt * EMU_PER_PT)
    y_emu = int(y_pt * EMU_PER_PT)
    cx_emu = int(w_pt * EMU_PER_PT)
    cy_emu = int(h_pt * EMU_PER_PT)
    # 建一個空段落 host shape
    p = doc.add_paragraph()
    # pPr 設 spacing 0 (避免段落自身佔位)
    p_pPr = p._element.get_or_add_pPr()
    sp = OxmlElement("w:spacing")
    sp.set(qn("w:before"), "0")
    sp.set(qn("w:after"), "0")
    sp.set(qn("w:line"), "0")
    sp.set(qn("w:lineRule"), "atLeast")
    p_pPr.append(sp)
    r = p.add_run()
    # 構造 wp:anchor XML
    NS = {
        "w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main",
        "wp": "http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing",
        "a": "http://schemas.openxmlformats.org/drawingml/2006/main",
        "wps": "http://schemas.microsoft.com/office/word/2010/wordprocessingShape",
    }
    xml = f'''<w:drawing xmlns:w="{NS['w']}" xmlns:wp="{NS['wp']}" xmlns:a="{NS['a']}" xmlns:wps="{NS['wps']}">
  <wp:anchor distT="0" distB="0" distL="0" distR="0" simplePos="0" relativeHeight="1" behindDoc="1" locked="0" layoutInCell="1" allowOverlap="1">
    <wp:simplePos x="0" y="0"/>
    <wp:positionH relativeFrom="page"><wp:posOffset>{x_emu}</wp:posOffset></wp:positionH>
    <wp:positionV relativeFrom="page"><wp:posOffset>{y_emu}</wp:posOffset></wp:positionV>
    <wp:extent cx="{cx_emu}" cy="{cy_emu}"/>
    <wp:effectExtent l="0" t="0" r="0" b="0"/>
    <wp:wrapNone/>
    <wp:docPr id="100" name="Banner"/>
    <wp:cNvGraphicFramePr/>
    <a:graphic>
      <a:graphicData uri="http://schemas.microsoft.com/office/word/2010/wordprocessingShape">
        <wps:wsp>
          <wps:cNvSpPr/>
          <wps:spPr>
            <a:xfrm><a:off x="0" y="0"/><a:ext cx="{cx_emu}" cy="{cy_emu}"/></a:xfrm>
            <a:prstGeom prst="rect"><a:avLst/></a:prstGeom>
            <a:solidFill><a:srgbClr val="{color}"/></a:solidFill>
            <a:ln><a:noFill/></a:ln>
          </wps:spPr>
          <wps:bodyPr/>
        </wps:wsp>
      </a:graphicData>
    </a:graphic>
  </wp:anchor>
</w:drawing>'''
    from lxml import etree
    drawing_el = etree.fromstring(xml)
    r._element.append(drawing_el)


def _convert_inline_to_anchor(drawing_el, x_emu: int, y_emu: int,
                                cx_emu: int, cy_emu: int) -> bool:
    """把 w:drawing 內的 wp:inline element 改成 wp:anchor with 絕對定位。
    回 True 表示成功，False 表示找不到 inline 結構 fallback 維持 inline。"""
    NS_WP = "http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing"
    NS_A = "http://schemas.openxmlformats.org/drawingml/2006/main"
    inline = drawing_el.find(f"{{{NS_WP}}}inline")
    if inline is None:
        return False
    # 抽出 inline 內子元素：a:graphic 整段保留 (內含 pic:pic + a:blip rId)
    graphic = inline.find(f"{{{NS_A}}}graphic")
    if graphic is None:
        return False
    docPr = inline.find(f"{{{NS_WP}}}docPr")
    cNvGraphicFramePr = inline.find(f"{{{NS_WP}}}cNvGraphicFramePr")
    # 建 anchor
    anchor = drawing_el.makeelement(f"{{{NS_WP}}}anchor", {
        "distT": "0", "distB": "0", "distL": "0", "distR": "0",
        "simplePos": "0", "relativeHeight": "251660288",
        "behindDoc": "0", "locked": "0", "layoutInCell": "1",
        "allowOverlap": "1",
    })
    simplePos = anchor.makeelement(f"{{{NS_WP}}}simplePos",
                                    {"x": "0", "y": "0"})
    anchor.append(simplePos)
    posH = anchor.makeelement(f"{{{NS_WP}}}positionH",
                               {"relativeFrom": "page"})
    posH_off = posH.makeelement(f"{{{NS_WP}}}posOffset", {})
    posH_off.text = str(x_emu)
    posH.append(posH_off)
    anchor.append(posH)
    posV = anchor.makeelement(f"{{{NS_WP}}}positionV",
                               {"relativeFrom": "page"})
    posV_off = posV.makeelement(f"{{{NS_WP}}}posOffset", {})
    posV_off.text = str(y_emu)
    posV.append(posV_off)
    anchor.append(posV)
    extent = anchor.makeelement(f"{{{NS_WP}}}extent",
                                 {"cx": str(cx_emu), "cy": str(cy_emu)})
    anchor.append(extent)
    effectExtent = anchor.makeelement(f"{{{NS_WP}}}effectExtent",
                                       {"l": "0", "t": "0", "r": "0", "b": "0"})
    anchor.append(effectExtent)
    # wrapNone — 不繞文字，浮在最上層（最像 PDF 渲染）
    wrapNone = anchor.makeelement(f"{{{NS_WP}}}wrapNone", {})
    anchor.append(wrapNone)
    if docPr is not None:
        anchor.append(docPr)
    if cNvGraphicFramePr is not None:
        anchor.append(cNvGraphicFramePr)
    anchor.append(graphic)
    # 替換
    drawing_el.remove(inline)
    drawing_el.append(anchor)
    return True


def _add_image_for_page(doc: Document, image, pdf_path: Path,
                          float_anchor: bool = True) -> None:
    """從 PDF 提圖 → 用 wp:anchor 絕對定位（page-relative）插入。

    float_anchor=True：image 浮在頁面上，X/Y 對齊 PDF 真實位置（wrapNone 不繞文字）
    float_anchor=False：傳統 wp:inline，跟著段落流
    """
    try:
        pdf_doc = fitz.open(str(pdf_path))
        # 用 extract_image 拿原始嵌入圖（保留 alpha / 透明背景）
        # 若有 smask → 合成 alpha mask 後輸出 PNG with alpha
        img_info = pdf_doc.extract_image(image.xref)
        smask_xref = int(img_info.get("smask") or 0)
        base_pix = fitz.Pixmap(pdf_doc, image.xref)
        if smask_xref:
            try:
                mask_pix = fitz.Pixmap(pdf_doc, smask_xref)
                # 合成 alpha：result = (base, mask) → RGBA
                combo = fitz.Pixmap(base_pix, mask_pix)
                png_bytes = combo.tobytes("png")
                combo = None
                mask_pix = None
            except Exception as e:
                log.debug("smask compose failed: %s; fallback no alpha", e)
                png_bytes = base_pix.tobytes("png")
        else:
            # 若 PyMuPDF Pixmap 本身就含 alpha（CMYK / Grayscale 不算），直接用
            png_bytes = base_pix.tobytes("png")
        base_pix = None
        pdf_doc.close()
    except Exception as e:
        log.debug("extract image xref=%s failed: %s", image.xref, e)
        return
    x0_pt, y0_pt, x1_pt, y1_pt = image.bbox
    width_pt = max(1.0, x1_pt - x0_pt)
    height_pt = max(1.0, y1_pt - y0_pt)
    cx_emu = int(width_pt * EMU_PER_PT)
    cy_emu = int(height_pt * EMU_PER_PT)
    try:
        p = doc.add_paragraph()
        r = p.add_run()
        r.add_picture(BytesIO(png_bytes),
                        width=Emu(cx_emu), height=Emu(cy_emu))
        if not float_anchor:
            return
        # 改 inline → anchor (絕對定位 page-relative)
        drawing_el = r._element.find(qn("w:drawing"))
        if drawing_el is None:
            return
        x_emu = int(x0_pt * EMU_PER_PT)
        y_emu = int(y0_pt * EMU_PER_PT)
        ok = _convert_inline_to_anchor(drawing_el, x_emu, y_emu, cx_emu, cy_emu)
        if not ok:
            log.debug("convert inline → anchor failed; keep inline")
    except Exception as e:
        log.debug("add_picture failed: %s", e)


def _add_single_hyperlink(doc: Document, link: dict) -> None:
    """單一 link 加成獨立段落（藍底線 run 模擬 hyperlink）。"""
    uri = link.get("uri") or link.get("file") or ""
    if not uri:
        return
    p = doc.add_paragraph()
    run = p.add_run(uri)
    rPr = run._element.get_or_add_rPr()
    color = rPr.makeelement(qn("w:color"), {qn("w:val"): HYPERLINK_COLOR})
    rPr.append(color)
    u = rPr.makeelement(qn("w:u"), {qn("w:val"): "single"})
    rPr.append(u)


def _add_hyperlinks_for_page(doc: Document, page_links: list,
                              page_num: int) -> int:
    """補 hyperlinks — 對每個 link 找最近文字段落，把該段第一個 run 包成 hyperlink。

    簡化版：直接在文末加一段 hyperlink，文字 = uri，給使用者參考。完整版要對應
    到原 text run，需要 link annotation bbox 跟段落 mapping，這版 MVP 暫不做。
    """
    n = 0
    for link in page_links:
        uri = link.get("uri") or link.get("file") or ""
        if not uri:
            continue
        p = doc.add_paragraph()
        # 用 plain run + 藍色底線模擬 hyperlink（不真設 w:hyperlink，避免動 rels）
        run = p.add_run(uri)
        run.font.color.rgb = None
        rPr = run._element.get_or_add_rPr()
        color = rPr.makeelement(qn("w:color"), {qn("w:val"): HYPERLINK_COLOR})
        rPr.append(color)
        u = rPr.makeelement(qn("w:u"), {qn("w:val"): "single"})
        rPr.append(u)
        n += 1
    return n


_PAGE_LINK_BY_BBOX: dict = {}  # 暫存當前 page 的 links（給 cell line 寫入時查詢）


def _find_link_for_line(line) -> dict | None:
    """看 line bbox 中心點是否落在某 page link bbox 內，是則回 link dict。"""
    if not _PAGE_LINK_BY_BBOX:
        return None
    links = _PAGE_LINK_BY_BBOX.get("links") or []
    lcx = float((line.bbox[0] + line.bbox[2]) / 2.0)
    lcy = float((line.bbox[1] + line.bbox[3]) / 2.0)
    for link in links:
        x0, y0, x1, y1 = link["bbox"]
        if x0 - 2 <= lcx <= x1 + 2 and y0 - 2 <= lcy <= y1 + 2:
            return link
    return None


def build_docx(doc_model: DocumentModel, pdf_path: Path,
                output_path: Path,
                page_links_by_page: dict | None = None) -> dict:
    """主入口：DocumentModel → docx 寫到 output_path。回 build stats dict。"""
    output_path = Path(output_path)
    output_path.parent.mkdir(parents=True, exist_ok=True)
    doc = Document()
    if doc_model.pages:
        first = doc_model.pages[0]
        _set_page_geometry(doc, first.width, first.height,
                            margin_left_pt=first.margin_left,
                            margin_top_pt=first.margin_top,
                            margin_right_pt=first.margin_right,
                            margin_bottom_pt=first.margin_bottom)

    total_tables = 0
    total_free_paras = 0
    total_images = 0
    total_links = 0

    for page_idx, pm in enumerate(doc_model.pages):
        # 第二頁起加 page break
        if page_idx > 0:
            try:
                from docx.enum.text import WD_BREAK
                br_p = doc.add_paragraph()
                br_p.add_run().add_break(WD_BREAK.PAGE)
            except Exception as e:
                log.debug("add page break failed: %s", e)
        # 設目前頁 link bbox cache（給 cell line 寫入時 wrap hyperlink 用）
        _PAGE_LINK_BY_BBOX.clear()
        _PAGE_LINK_BY_BBOX["links"] = (page_links_by_page or {}).get(pm.page_num) or []

        # 先加 banner floating shapes (behindDoc=1)
        # 改用 PIL 生純色 PNG + wp:anchor 浮動圖（比 wps:wsp shape LibreOffice 認）
        for banner in getattr(pm, "banner_rects", None) or []:
            # 兼容 5-tuple (舊) 與 6-tuple (新含 path_points)
            if len(banner) >= 6:
                x0, y0, x1, y1, fill_hex, path_points = banner[:6]
            else:
                x0, y0, x1, y1, fill_hex = banner[:5]
                path_points = []
            try:
                if path_points and len(path_points) >= 3:
                    _add_floating_polygon_image(doc, x0, y0, x1 - x0, y1 - y0,
                                                  fill_hex, path_points)
                else:
                    _add_floating_color_image(doc, x0, y0, x1 - x0, y1 - y0, fill_hex)
            except Exception as e:
                log.debug("add banner image failed: %s", e)

        # 同頁內按 Y 序混排 tables (real + virtual) / free blocks / images / links
        elements: list[tuple[float, str, object]] = []
        for fb in pm.free_blocks:
            elements.append((float(fb.block.bbox[1]), "para", fb))
        for tm in pm.tables:
            elements.append((float(tm.region.bbox[1]), "table", tm))
        for img in pm.images:
            elements.append((float(img.bbox[1]), "image", img))
        if page_links_by_page:
            # 收集本頁所有 line 文字（normalized）做 hyperlink dedup
            # 若 link uri 對應的文字已在某個 text line 內出現，就 skip 該 link
            # (避免「電子信箱 <email>」cell 之後又補一段 mailto:<email>)
            page_text_norm: set[str] = set()
            import re as _re
            for blk in pm.tables:
                for r in (blk.cell_lines or []):
                    for cell in r:
                        for ln in cell:
                            if ln.text:
                                page_text_norm.add(_re.sub(r"\s+", "", ln.text))
            for fb in pm.free_blocks:
                for ln in (fb.block.lines or []):
                    if ln.text:
                        page_text_norm.add(_re.sub(r"\s+", "", ln.text))
            for link in (page_links_by_page.get(pm.page_num) or []):
                uri = link.get("uri") or ""
                # uri 內取核心字串（去掉 mailto: / http://）
                core = uri
                for prefix in ("mailto:", "http://", "https://", "ftp://"):
                    if core.startswith(prefix):
                        core = core[len(prefix):]
                        break
                core_norm = _re.sub(r"\s+", "", core).rstrip("/")
                # 若該 uri 核心字串已在頁面 text 內 → skip (避免重覆補)
                if core_norm and any(core_norm in t for t in page_text_norm):
                    continue
                elements.append((float(link["bbox"][1]), "link", link))
        elements.sort(key=lambda e: e[0])

        for y, kind, payload in elements:
            try:
                if kind == "para":
                    _add_paragraph_for_block(doc, payload.block,
                                              is_heading=payload.is_heading,
                                              is_footer=payload.is_footer,
                                              page_width=float(pm.width or 0))
                    total_free_paras += 1
                elif kind == "table":
                    # **重要**：連續兩個 table 之間必須有 paragraph 分隔，否則
                    # OxOffice / LibreOffice 把它們 join 並把主表 widths 按小 table
                    # ratio 縮（觸發 bug：主表 4 col 變 1/3 寬度，內文變直書）
                    if doc.element.body[-1].tag == qn("w:tbl"):
                        spacer = doc.add_paragraph()
                        # spacer paragraph 高度設 0
                        p_pPr = spacer._element.get_or_add_pPr()
                        sp = OxmlElement("w:spacing")
                        sp.set(qn("w:before"), "0")
                        sp.set(qn("w:after"), "0")
                        sp.set(qn("w:line"), "20")  # 1pt 最小行高
                        sp.set(qn("w:lineRule"), "exact")
                        p_pPr.append(sp)
                    if _add_table(doc, payload) is not None:
                        total_tables += 1
                elif kind == "image":
                    _add_image_for_page(doc, payload, pdf_path)
                    total_images += 1
                elif kind == "link":
                    _add_single_hyperlink(doc, payload)
                    total_links += 1
            except Exception as e:
                log.debug("Y-order build element %s failed: %s", kind, e)

    doc.save(str(output_path))
    return {
        "ok": True,
        "pages": len(doc_model.pages),
        "tables": total_tables,
        "free_paragraphs": total_free_paras,
        "images": total_images,
        "hyperlinks": total_links,
    }
