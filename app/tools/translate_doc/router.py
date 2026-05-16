"""逐句翻譯 — backend.

API:
    GET  /                         → tool 頁面
    POST /extract-text             → 上傳 PDF / DOCX / TXT，回 sentences[]
    POST /translate-batch          → 對一批 sentences 翻譯，回 translations[]
    POST /translate-one            → 單句重新翻譯（為了 UI 重生成單句）
    POST /api/translate-doc        → 一次性 API：吃 text + target_lang，回對齊結果
                                      （符合「所有功能須有 API」規範）

LLM 設定來自 admin (`/admin/llm-settings`)。LLM 沒啟用時 endpoints 一律
回 503 + 提示去設定。我們不在這裡再開設定 UI。

注意：long doc 切句後會逐句送到 LLM，timeout 會放寬到 admin 設定值；
但建議單次 < 500 句以避免 UI hang 太久。500 句以上請自行分段。
"""
from __future__ import annotations

import io
import json
import re
import uuid
from pathlib import Path
from typing import Optional

from fastapi import APIRouter, File, Form, HTTPException, Request, UploadFile
from fastapi.responses import HTMLResponse, JSONResponse, Response

from ...config import settings
from ...core.llm_settings import llm_settings


router = APIRouter()


# 切句粒度：句點 / 中文句號 / 換行 / 問號 / 驚嘆號。中文與英文混合時
# 句尾標點要連帶吃進句子內，避免「。」獨立成句。
_SENT_SPLIT_RE = re.compile(
    r"(?<=[\.!?。！？])\s+|(?<=[\.!?。！？])(?=[A-Z一-鿿])|\n+"
)
# 上限保護：超過這數量就不全送 LLM（會 timeout / 沒意義）
MAX_SENTENCES = 800
MAX_TEXT_BYTES = 2 * 1024 * 1024  # 2 MB raw text


def _split_sentences(text: str) -> list[str]:
    """把長文字切成「逐句」的 list。空白行 / 純 whitespace 略過。
    保留行序，每句去頭尾空白。"""
    if not text:
        return []
    # Normalize line endings
    text = text.replace("\r\n", "\n").replace("\r", "\n")
    raw = _SENT_SPLIT_RE.split(text)
    out = []
    for s in raw:
        s = (s or "").strip()
        if s:
            out.append(s)
    return out


_LIST_MARKER_RE = re.compile(
    r"^("
    r"[IVXLCM]+\.?"          # Roman numerals: I, II, III, IV.
    r"|[ivxlcm]+\.?"         # lower-case roman: i, ii.
    r"|\d+[.)]?"             # 1, 1., 1)
    r"|[A-Za-z][.)]"         # A., a), b)
    r"|[•◦▪■◆●○\-*]"         # bullet glyphs
    r")$"
)


def _merge_list_markers(paras: list[str]) -> list[str]:
    """Merge bare "list marker" paragraphs (e.g. "I.", "1)", "•") into the
    next paragraph. Common case: original document put the marker on its
    own line — we'd otherwise translate them in isolation, losing context
    and giving the LLM nothing to translate from.
    """
    merged: list[str] = []
    pending: str | None = None
    for p in paras:
        s = p.strip()
        if not s:
            continue
        if pending is not None:
            merged.append(f"{pending} {s}")
            pending = None
            continue
        if len(s) <= 6 and _LIST_MARKER_RE.match(s):
            pending = s
            continue
        merged.append(s)
    if pending is not None:
        merged.append(pending)
    return merged


def _extract_text_from_odf(data: bytes, kind: str) -> str:
    """Parse OpenDocument (ODT / ODS / ODP) `content.xml` directly.

    ODF files are zips containing `content.xml`; the text we want lives in
    `<text:p>` / `<text:h>` elements (declared in the urn:oasis text
    namespace). We strip namespaces by hand instead of pulling in `lxml`.

    We don't bother going through soffice — that would require the binary,
    cost a subprocess + temp files per upload, and lose paragraph breaks
    when round-tripping through PDF.
    """
    import zipfile
    from xml.etree import ElementTree as ET
    try:
        with zipfile.ZipFile(io.BytesIO(data)) as zf:
            with zf.open("content.xml") as fp:
                tree = ET.parse(fp)
    except (zipfile.BadZipFile, KeyError) as e:
        raise HTTPException(400, f"{kind.upper()} parse failed: {e}")
    text_ns = "{urn:oasis:names:tc:opendocument:xmlns:text:1.0}"
    raw_paras: list[str] = []
    for el in tree.iter():
        if el.tag in (text_ns + "p", text_ns + "h"):
            # itertext walks all descendants — joins text-runs split by spans
            txt = "".join(el.itertext()).strip()
            if txt:
                raw_paras.append(txt)
    return "\n\n".join(_merge_list_markers(raw_paras))


def _extract_text_from_file(filename: str, data: bytes) -> str:
    """支援 PDF / DOCX / ODT / ODS / ODP / TXT / MD。其他副檔名拋 400。"""
    name = (filename or "").lower()
    if name.endswith(".txt") or name.endswith(".md"):
        # Try utf-8 first, then fall back to common encodings.
        for enc in ("utf-8", "utf-8-sig", "big5", "cp950", "latin-1"):
            try:
                return data.decode(enc)
            except UnicodeDecodeError:
                continue
        return data.decode("utf-8", errors="replace")
    if name.endswith(".pdf"):
        try:
            import fitz
        except ImportError:
            raise HTTPException(500, "PyMuPDF not available")
        try:
            with fitz.open(stream=data, filetype="pdf") as doc:
                paras: list[str] = []
                for page in doc:
                    text = page.get_text("text") or ""
                    # PyMuPDF separates paragraphs with blank lines; split on
                    # 2+ newlines to keep paragraph boundaries intact while
                    # merging soft line wraps within a paragraph.
                    for chunk in re.split(r"\n\s*\n+", text):
                        chunk = chunk.replace("\n", " ").strip()
                        if chunk:
                            paras.append(chunk)
                return "\n\n".join(_merge_list_markers(paras))
        except Exception as e:
            raise HTTPException(400, f"PDF parse failed: {e}")
    # Office / ODF：交給 soffice 匯出 UTF-8 文字（等同「打開→另存為純文字」），
    # 段落結構跟使用者在 OxOffice/LibreOffice 看到的一致。比直接 parse XML
    # 多 ~1-2 秒 subprocess，但結果穩定 — 列表編號、表格、註腳都正常。
    office_exts = (".docx", ".doc", ".odt", ".ods", ".odp", ".rtf")
    if name.endswith(office_exts):
        from ...core import office_convert
        import tempfile as _tf
        suffix = "." + name.rsplit(".", 1)[-1]
        with _tf.NamedTemporaryFile(suffix=suffix, delete=False) as tf:
            tf.write(data)
            src_path = Path(tf.name)
        try:
            text = office_convert.convert_to_text(src_path)
        except Exception as e:
            raise HTTPException(400, f"office 檔解析失敗：{e}")
        finally:
            try:
                src_path.unlink()
            except Exception:
                pass
        # soffice TXT 輸出本來就有清楚的段落分行；直接重排即可
        paras = [ln.strip() for ln in re.split(r"\n\s*\n+", text) if ln.strip()]
        # 把每段內的單一換行折回去（保留段落感）
        paras = [re.sub(r"\s*\n\s*", " ", p) for p in paras]
        return "\n\n".join(_merge_list_markers(paras))
    raise HTTPException(400, f"unsupported file type: {filename}")


def _detect_language(text: str) -> str:
    """超輕量語言偵測。看前 N 個字符的 unicode block 比例：
        - 30%+ CJK chars → 'zh' (粗略；繁中簡中沒分)
        - 否則 → 'en' (假設使用者翻的多半是中英對譯)
    回 ISO 639-1 二字碼，未知回 'auto'。"""
    sample = text[:2000]
    if not sample:
        return "auto"
    cjk_n = sum(1 for c in sample if "一" <= c <= "鿿")
    total_letters = sum(1 for c in sample if c.isalpha() or "一" <= c <= "鿿")
    if total_letters == 0:
        return "auto"
    if cjk_n / total_letters > 0.3:
        return "zh"
    return "en"


_LANG_NAMES = {
    "zh-TW": "繁體中文", "zh-CN": "簡體中文", "zh": "中文",
    "en": "英文", "ja": "日文", "ko": "韓文",
    "fr": "法文", "de": "德文", "es": "西班牙文", "vi": "越南文",
    "th": "泰文", "id": "印尼文", "ru": "俄文",
}


def _build_prompt(src_text: str, source_lang: str, target_lang: str,
                  domain: str = "") -> str:
    src_name = _LANG_NAMES.get(source_lang, source_lang or "原文")
    tgt_name = _LANG_NAMES.get(target_lang, target_lang or "目標語言")
    # 針對台灣繁體要求 LLM 用台灣慣用 IT 術語（避免大陸用語滲入）
    tw_terminology = ""
    if target_lang in ("zh-TW", "zh"):
        tw_terminology = (
            "**重要**：請使用「台灣繁體中文」用語，IT / 技術術語請用台灣業界習慣翻譯，**禁止**用大陸 / 香港用詞。對照表："
            "kernel→核心（不是「內核」）、software→軟體（不是「軟件」）、hardware→硬體（不是「硬件」）、"
            "image→圖片 / 影像（不是「圖像」）、video→影片（不是「視頻」）、network→網路（不是「網絡」）、"
            "server→伺服器（不是「服務器」）、menu→選單（不是「菜單」）、screen→螢幕（不是「屏幕」）、"
            "save→儲存（不是「保存」）、default→預設（不是「默認」）、setting→設定（不是「設置」）、"
            "file→檔案（不是「文件」）、information→訊息 / 資訊（不是「信息」）、"
            "print→列印（不是「打印」）、font→字型（不是「字體」）、document→文件（不是「文檔」）、"
            "code→程式碼（不是「代碼」）、program→程式（不是「程序」）、function→函式 / 功能（不是「函數」當動詞）、"
            "data→資料（不是「數據」）、object→物件（不是「對象」當程式術語）、array→陣列（不是「數組」）、"
            "queue→佇列（不是「隊列」）、cache→快取（不是「緩存」）、download→下載（OK 兩岸通用）、"
            "upload→上傳、login→登入（不是「登錄」）、logout→登出、user→使用者（不是「用戶」當主語）、"
            "browser→瀏覽器、driver→驅動程式、interface→介面（不是「界面」）、"
            "framework→框架、library→程式庫 / 函式庫（IT 上下文，不是「圖書館」）、"
            "feature→功能 / 特色（不是「特性」）、bug→錯誤 / 臭蟲、release→版本 / 釋出、deploy→部署、"
            "click→點擊 / 按、support→支援（不是「支持」當技術動詞）、"
            "performance→效能（不是「性能」）、optimize→最佳化（不是「優化」當形容詞）、"
            "address→位址（IP/記憶體上下文，不是「地址」）、port→連接埠 / 通訊埠（不是「端口」）、"
            "container→容器（OK）、virtualization→虛擬化、virtual machine→虛擬機（OK）。"
        )
    # v1.5.15: domain hint（選填）— user 在 UI 填的「文件領域」直接放進
    # prompt，讓 LLM 知道這是法律 / 醫療 / 技術 / 財務等場景，挑對應專業用詞
    domain_hint = ""
    d = (domain or "").strip()
    if d:
        # 截斷防 prompt-injection，留 80 chars 已足夠描述領域
        d = d[:80].replace("\n", " ").replace("\r", " ")
        domain_hint = (
            f"**文件領域**：{d}。請依此領域的慣用術語、文體與專業用詞翻譯，"
            "縮寫保持原貌（如 SLA / API / GDPR），技術名詞不要過度本地化。"
        )
    return (
        f"請把下面這句從「{src_name}」翻譯為「{tgt_name}」，"
        "翻譯要忠實、通順、符合該語言慣用的書寫風格。"
        + tw_terminology
        + domain_hint +
        "只輸出翻譯結果，不要附上原文、不要加任何標記、不要解釋、"
        "不要 markdown、不要行首符號、不要後綴。"
        "如果原文是專有名詞或無法翻譯，照原樣輸出即可。\n\n"
        f"原文：{src_text}"
    )


_FILLER_RE = re.compile(r"^[\s_\-=·.•◦▪■◇◆●○※—–…]+$")
# 「不需翻譯的純標記行」— 例：markdown code fence (``` / ~~~ 可帶語言名)、
# horizontal rule (--- / *** / ___)、純符號 / emoji、URL only。LLM 收到這
# 類輸入會掰一句「Please provide the text to translate」回來，造成譯文欄
# 雜訊。直接 passthrough 回原文，譯文留空。
_NOTRANSLATE_RES = (
    re.compile(r"^\s*(`{3,}|~{3,})\s*[A-Za-z0-9_+\-.#]*\s*$"),  # ``` python / ~~~ bash
    re.compile(r"^\s*(\*\s*){3,}\s*$"),                          # *** thematic break
    re.compile(r"^\s*[`~*#>]\s*$"),                              # 單個 markdown 標記
    re.compile(r"^\s*<[^>]+>\s*$"),                              # 純 HTML tag 一行
    re.compile(r"^\s*https?://\S+\s*$"),                          # 純 URL 一行
    re.compile(r"^[\s|+\-=]{3,}$"),                              # 表格分隔線 |---|---|
)


def _is_no_translate(src: str) -> bool:
    """回 True 表示這行只是 markdown / 標點 / 符號，不該送 LLM。"""
    if _FILLER_RE.match(src):
        return True
    for r in _NOTRANSLATE_RES:
        if r.match(src):
            return True
    return False


# 抽出行首 markdown / 縮排行首符號（list bullet、blockquote、heading、checkbox 等）
# 翻完後再補回譯文，免得 LLM 把 "- 裝置雙向同步" 翻成 "Two-way device sync" 漏掉「- 」
_LINE_PREFIX_RE = re.compile(
    r"^(\s*"                          # 縮排
    r"(?:[-*+]\s+(?:\[[ xX]\]\s+)?"   # - / * / + bullet，可帶 [x] checkbox
    r"|\d+\.\s+"                      # 1. 2. 3. 列表
    r"|>\s+"                          # > blockquote
    r"|#{1,6}\s+"                     # # / ## / ### 標題
    r")?)(.*)$",
    re.DOTALL,
)


def _split_line_prefix(src: str) -> tuple[str, str]:
    """回 (prefix, body)。沒行首符號時 prefix=''，body=src。"""
    m = _LINE_PREFIX_RE.match(src)
    if not m:
        return "", src
    prefix, body = m.group(1) or "", m.group(2) or ""
    return prefix, body


def _looks_like_proxy_error_page(text: str) -> str:
    """偵測譯文其實是 reverse proxy (nginx / cloudflare / haproxy) 的錯誤頁，
    LLM proxy 上游 timeout 時有些設定會回 200 + 錯誤 HTML body（LiteLLM 某些
    `fallbacks` / `return_response_on_error` 設定），jtdt 端不可把這當譯文存。
    回 status 字串（"504"/"502"/"...") 或 "" 表示正常。"""
    if not text:
        return ""
    t = text.lstrip()[:500].lower()
    # 1) 明顯的 HTML wrapper
    if not (t.startswith("<html") or t.startswith("<!doctype") or
            t.startswith("<head") or "<title>" in t[:200]):
        return ""
    # 2) 找錯誤碼 — 5xx / 4xx gateway / bad gateway / timeout
    import re as _re
    m = _re.search(r"(\b50[02-4]\b|\b40[0-9]\b)\s*(gateway|server error|timeout|unavailable|bad|forbidden|too many)", t)
    if m:
        return m.group(1)
    if "gateway time-out" in t or "bad gateway" in t or "service unavailable" in t:
        m2 = _re.search(r"\b(50[02-4])\b", t)
        return m2.group(1) if m2 else "5xx"
    if "cloudflare" in t and ("error" in t or "timeout" in t):
        return "cf-error"
    # 3) HTML 但不認得錯誤類型 — 仍視為異常（譯文絕不會回 HTML）
    return "html-body"


_MAX_TRANSLATE_RETRY = 2


def _translate_one(client, model: str, src: str,
                   source_lang: str, target_lang: str,
                   domain: str = "") -> dict:
    """單句翻譯 worker（給並行 executor 用）。
    空字串直接回 empty，不發 LLM call。任何 exception 回 error 字串，
    不 raise — caller 用 list 收集所有結果。
    收到 reverse proxy HTML error page (如 nginx 504) 自動 retry 一次，
    再失敗就回 error 不污染譯文欄。"""
    if not src.strip():
        return {"src": src, "translated": "", "error": ""}
    if _is_no_translate(src):
        return {"src": src, "translated": "", "error": "", "skipped": "filler"}
    prefix, body = _split_line_prefix(src)
    if not body.strip():
        return {"src": src, "translated": "", "error": "", "skipped": "filler"}
    prompt = _build_prompt(body, source_lang, target_lang, domain=domain)
    last_error = ""
    for attempt in range(_MAX_TRANSLATE_RETRY + 1):
        try:
            resp = client.text_query(
                prompt=prompt, model=model, temperature=0.0, think=False,
            )
            translated = (resp or "").strip()
            translated = re.sub(r"^(翻譯[:：]?\s*|Translation:\s*)",
                                "", translated, flags=re.IGNORECASE)
            # === 防 proxy error page 污染譯文（v1.8.58+）===
            proxy_err = _looks_like_proxy_error_page(translated)
            if proxy_err:
                last_error = f"LLM proxy 回錯誤頁 ({proxy_err})；嘗試 {attempt + 1}/{_MAX_TRANSLATE_RETRY + 1}"
                if attempt < _MAX_TRANSLATE_RETRY:
                    continue  # retry
                return {"src": src, "translated": "",
                        "error": f"⚠ LLM 上游錯誤 ({proxy_err})：請檢查 LLM proxy 設定 (gateway timeout)；此句未翻譯"}
            # 防 LLM 自己又加上同樣的行首符號 → 重複
            if prefix and translated.startswith(prefix.strip()):
                translated = translated[len(prefix.strip()):].lstrip()
            if prefix:
                translated = prefix + translated
            return {"src": src, "translated": translated, "error": ""}
        except Exception as e:
            last_error = f"LLM 失敗：{e}"
            if attempt < _MAX_TRANSLATE_RETRY:
                continue
            return {"src": src, "translated": "", "error": last_error}
    return {"src": src, "translated": "", "error": last_error or "未知失敗"}


def _warmup_llm(client, model: str) -> None:
    """送一句極短 ping 強制 Ollama 把 model load 進記憶體。

    Ollama 第一次呼叫某個模型時會先從 disk 載入到 VRAM（26B 模型常 30-90 秒）。
    平行送 batch 時，4 個 worker 同時等 cold load，httpx 預設 60s timeout 會通通
    一起 fire → 整個 batch 失敗。先送一個 sync warm-up call 阻塞等模型 ready，
    後續 batch 就跑 hot path 不會 timeout。

    失敗不 raise — 只是 best-effort 預熱；真翻譯還是會跑，最壞回到原本行為。
    """
    try:
        client.text_query(prompt="hi", model=model, temperature=0.0,
                          max_tokens=4, think=False)
    except Exception:
        # 預熱失敗就放行，讓真翻譯自己處理錯誤（user 看得到具體 error）
        pass


def _translate_sentences(
    sentences: list[str], source_lang: str, target_lang: str,
    domain: str = "",
) -> list[dict]:
    client = llm_settings.make_client()
    if client is None:
        raise HTTPException(503, "LLM 服務未啟用，請到「設定 → LLM 設定」開啟")
    # Per-tool 模型覆寫優先；admin 在 LLM 設定頁可以給 translate-doc 指定
    # 不同模型（例如純文字翻譯用 gemma4:26b，校驗仍用 gemma4:26b）
    model = llm_settings.get_model_for("translate-doc")
    conf = llm_settings.get()
    # 並行數 — admin 在 LLM 設定可調，預設 4。對 1-句 case 退化成同步呼叫。
    concurrency = max(1, min(16, int(conf.get("translate_concurrency", 4))))
    # **預熱 Ollama** — 先送一個 sync 短 ping 強制 model load 進 VRAM，避免 4 個
    # 平行 worker 一起卡 cold load → httpx 60s timeout 全部 fire 整個 batch 死。
    # 客戶 v1.8.31 回報的「卡住」根因。
    _warmup_llm(client, model)
    n = len(sentences)
    if n <= 1 or concurrency == 1:
        return [_translate_one(client, model, src, source_lang, target_lang, domain=domain)
                for src in sentences]
    # ThreadPoolExecutor.map 保證 output 順序對應 input 順序（重要 — UI
    # 是依索引貼回原文位置）。內部 LLMClient 的 httpx 是同步的，所以用
    # thread pool 而不是 asyncio.gather。
    from concurrent.futures import ThreadPoolExecutor
    with ThreadPoolExecutor(max_workers=concurrency) as ex:
        results = list(ex.map(
            lambda src: _translate_one(client, model, src, source_lang, target_lang, domain=domain),
            sentences,
        ))
    return results


# ---- routes -----------------------------------------------------------------

@router.get("/", response_class=HTMLResponse)
async def index(request: Request):
    templates = request.app.state.templates
    s = llm_settings.get()
    from ...core.office_convert import detect_engine
    return templates.TemplateResponse(
        "translate_doc.html",
        {
            "request": request,
            "llm_enabled": bool(s.get("enabled")),
            "llm_model": llm_settings.get_model_for("translate-doc"),
            "llm_default_model": s.get("model", ""),
            "llm_url": s.get("base_url", ""),
            "office_engine": detect_engine(),
        },
    )


@router.post("/extract-text")
async def extract_text(file: UploadFile = File(...)):
    data = await file.read()
    if not data:
        raise HTTPException(400, "empty file")
    if len(data) > 50 * 1024 * 1024:
        raise HTTPException(400, "file too large (limit 50 MB)")
    text = _extract_text_from_file(file.filename or "", data)
    sentences = _split_sentences(text)
    detected = _detect_language(text)
    return {
        "filename": file.filename,
        "char_count": len(text),
        "sentence_count": len(sentences),
        "sentences": sentences[:MAX_SENTENCES],
        "truncated": len(sentences) > MAX_SENTENCES,
        "detected_lang": detected,
    }


@router.post("/translate-batch")
async def translate_batch(request: Request):
    if not llm_settings.is_enabled():
        raise HTTPException(503, "LLM 服務未啟用")
    body = await request.json()
    sentences = body.get("sentences") or []
    if not isinstance(sentences, list):
        raise HTTPException(400, "sentences must be array")
    if len(sentences) > MAX_SENTENCES:
        raise HTTPException(
            400, f"too many sentences ({len(sentences)} > {MAX_SENTENCES})")
    sentences = [str(s) for s in sentences]
    source_lang = str(body.get("source_lang") or "auto")
    target_lang = str(body.get("target_lang") or "zh-TW")
    domain = str(body.get("domain") or "")
    if source_lang == "auto":
        source_lang = _detect_language("\n".join(sentences[:50]))
    # CRITICAL: 不能直接呼叫 _translate_sentences — 它是同步的 (內部
    # ThreadPoolExecutor + 阻塞 .map())，會卡住整個 async event loop →
    # 翻譯期間其他 request 全部排隊（v1.4.13 客戶回報：翻譯時開新分頁
    # 進其他工具都打不開）。用 asyncio.to_thread 把它送到 default
    # executor 跑，async loop 就能繼續處理其他請求。
    import asyncio as _asyncio
    results = await _asyncio.to_thread(
        _translate_sentences, sentences, source_lang, target_lang, domain)
    return {
        "source_lang": source_lang,
        "target_lang": target_lang,
        "results": results,
    }


@router.post("/translate-one")
async def translate_one(request: Request):
    if not llm_settings.is_enabled():
        raise HTTPException(503, "LLM 服務未啟用")
    body = await request.json()
    src = str(body.get("src") or "").strip()
    if not src:
        raise HTTPException(400, "src is empty")
    source_lang = str(body.get("source_lang") or "auto")
    target_lang = str(body.get("target_lang") or "zh-TW")
    domain = str(body.get("domain") or "")
    if source_lang == "auto":
        source_lang = _detect_language(src)
    import asyncio as _asyncio
    results = await _asyncio.to_thread(
        _translate_sentences, [src], source_lang, target_lang, domain)
    return results[0] if results else {"src": src, "translated": "",
                                       "error": "no result"}


# ---- public API endpoint (符合「所有功能須有 API」規範) ---------------------

@router.post("/api/translate-doc")
async def api_translate_doc(request: Request):
    """One-shot translate API — 吃 raw text，回對齊好的中英並排 array。

    Body (JSON):
        {
          "text": "要翻譯的文字（必填）",
          "source_lang": "auto" | "en" | "zh" | "zh-TW" | ...   (default: auto)
          "target_lang": "zh-TW"   (default: zh-TW)
        }

    Response:
        {
          "source_lang": "...", "target_lang": "...",
          "results": [{"src": "...", "translated": "...", "error": ""}, ...]
        }
    """
    if not llm_settings.is_enabled():
        raise HTTPException(503, "LLM 服務未啟用 — 請到 admin / LLM 設定開啟")
    body = await request.json()
    text = str(body.get("text") or "")
    if not text.strip():
        raise HTTPException(400, "text is empty")
    if len(text.encode("utf-8")) > MAX_TEXT_BYTES:
        raise HTTPException(400, f"text too large (limit {MAX_TEXT_BYTES} bytes)")
    sentences = _split_sentences(text)
    if not sentences:
        return {"source_lang": "auto", "target_lang": "zh-TW", "results": []}
    if len(sentences) > MAX_SENTENCES:
        raise HTTPException(
            400, f"too many sentences after split ({len(sentences)} > {MAX_SENTENCES})")
    source_lang = str(body.get("source_lang") or "auto")
    target_lang = str(body.get("target_lang") or "zh-TW")
    domain = str(body.get("domain") or "")
    if source_lang == "auto":
        source_lang = _detect_language(text)
    # CRITICAL: 不能直接呼叫 _translate_sentences — 它是同步的 (內部
    # ThreadPoolExecutor + 阻塞 .map())，會卡住整個 async event loop →
    # 翻譯期間其他 request 全部排隊（v1.4.13 客戶回報：翻譯時開新分頁
    # 進其他工具都打不開）。用 asyncio.to_thread 把它送到 default
    # executor 跑，async loop 就能繼續處理其他請求。
    import asyncio as _asyncio
    results = await _asyncio.to_thread(
        _translate_sentences, sentences, source_lang, target_lang, domain)
    return {
        "source_lang": source_lang,
        "target_lang": target_lang,
        "results": results,
    }


# ============ 匯出格式（v1.7.51）============
# 接受 {format, pairs:[{source, target}]}，產對應檔回傳。
# 三組 dropdown 對應 8 個格式：
#   文字檔: txt / md / csv
#   文件檔: docx / odt / pdf
#   試算表: xlsx / ods
_EXPORT_FORMATS = {"txt", "md", "csv", "docx", "odt", "pdf", "xlsx", "ods"}
_EXPORT_MIME = {
    "txt":  "text/plain; charset=utf-8",
    "md":   "text/markdown; charset=utf-8",
    "csv":  "text/csv; charset=utf-8",
    "docx": "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    "odt":  "application/vnd.oasis.opendocument.text",
    "pdf":  "application/pdf",
    "xlsx": "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
    "ods":  "application/vnd.oasis.opendocument.spreadsheet",
}


def _build_txt(pairs: list[dict]) -> bytes:
    """純文字「原文／譯文」對照，每對中間空行分隔。"""
    out = []
    for p in pairs:
        out.append((p.get("source") or "").rstrip())
        out.append((p.get("target") or "").rstrip())
        out.append("")
    return "\n".join(out).encode("utf-8")


def _build_md(pairs: list[dict]) -> bytes:
    """Markdown 表格（兩欄：原文 / 譯文）。"""
    def esc(s: str) -> str:
        return (s or "").replace("|", "\\|").replace("\n", "<br>")
    out = ["| 原文 | 譯文 |", "|---|---|"]
    for p in pairs:
        out.append(f"| {esc(p.get('source', ''))} | {esc(p.get('target', ''))} |")
    return "\n".join(out).encode("utf-8")


def _build_csv(pairs: list[dict]) -> bytes:
    """CSV — 兩欄 source,target，UTF-8 BOM 讓 Excel 正確識別中文。"""
    import csv as _csv
    buf = io.StringIO()
    w = _csv.writer(buf, quoting=_csv.QUOTE_MINIMAL)
    w.writerow(["source", "target"])
    for p in pairs:
        w.writerow([p.get("source", ""), p.get("target", "")])
    return b"\xef\xbb\xbf" + buf.getvalue().encode("utf-8")


def _meta_lines(meta: dict) -> list[str]:
    """頁首 meta 文字 — 上傳檔名（如有）+ 翻譯時間。"""
    out = []
    src_name = (meta.get("source_filename") or "").strip()
    if src_name:
        out.append(f"原檔：{src_name}")
    ts = (meta.get("translated_at") or "").strip()
    if ts:
        # ISO → 友善格式
        try:
            from datetime import datetime as _dt
            dt = _dt.fromisoformat(ts.replace("Z", "+00:00"))
            ts_pretty = dt.strftime("%Y-%m-%d %H:%M")
        except Exception:
            ts_pretty = ts
        out.append(f"翻譯時間：{ts_pretty}")
    out.append(f"共 {meta.get('count', '?')} 對")
    return out


def _build_docx(pairs: list[dict], meta: dict) -> bytes:
    """Word 兩欄表格，加樣式：標題列藍底白字粗體、表格框線、欄寬均分、
    交替橫條淡底，文件最上方加 heading + 檔名 + 翻譯時間。"""
    from docx import Document
    from docx.shared import Pt, Cm, RGBColor
    from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement

    doc = Document()
    # heading
    h = doc.add_heading("逐句翻譯對照", level=1)
    for r in h.runs:
        r.font.color.rgb = RGBColor(0x1E, 0x3A, 0x8A)
    # meta
    sub = doc.add_paragraph()
    r = sub.add_run("　・　".join(_meta_lines({**meta, "count": len(pairs)})))
    r.font.size = Pt(9)
    r.font.color.rgb = RGBColor(0x64, 0x74, 0x8B)

    table = doc.add_table(rows=1 + len(pairs), cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    for col in table.columns:
        for cell in col.cells:
            cell.width = Cm(8)

    def _set_cell_bg(cell, hex_color):
        tc_pr = cell._tc.get_or_add_tcPr()
        shd = OxmlElement('w:shd')
        shd.set(qn('w:val'), 'clear')
        shd.set(qn('w:color'), 'auto')
        shd.set(qn('w:fill'), hex_color)
        tc_pr.append(shd)

    tbl_pr = table._tbl.tblPr
    borders = OxmlElement('w:tblBorders')
    for tag in ('top', 'left', 'bottom', 'right', 'insideH', 'insideV'):
        b = OxmlElement(f'w:{tag}')
        b.set(qn('w:val'), 'single')
        b.set(qn('w:sz'), '4')
        b.set(qn('w:color'), 'CBD5E1')
        borders.append(b)
    tbl_pr.append(borders)

    hdr_cells = table.rows[0].cells
    for i, txt in enumerate(("原文", "譯文")):
        cell = hdr_cells[i]
        cell.text = ""
        run = cell.paragraphs[0].add_run(txt)
        run.bold = True
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        run.font.size = Pt(11)
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        _set_cell_bg(cell, "2563EB")

    for i, p in enumerate(pairs, start=1):
        cells = table.rows[i].cells
        for j, key in enumerate(("source", "target")):
            cell = cells[j]
            cell.text = ""
            run = cell.paragraphs[0].add_run(p.get(key) or "")
            run.font.size = Pt(10)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.TOP
            if i % 2 == 0:
                _set_cell_bg(cell, "F1F5F9")

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def _build_odt(pairs: list[dict], meta: dict) -> bytes:
    """ODT 兩欄表格 + 樣式：標題列藍底白字、淡灰交替列底、欄寬 8cm。"""
    from odf.opendocument import OpenDocumentText
    from odf.style import (Style, TableColumnProperties, TableCellProperties,
                           TextProperties, ParagraphProperties)
    from odf.table import Table, TableColumn, TableRow, TableCell
    from odf.text import P, H

    doc = OpenDocumentText()
    # styles
    col_st = Style(name="ColW8", family="table-column")
    col_st.addElement(TableColumnProperties(columnwidth="8cm"))
    doc.automaticstyles.addElement(col_st)

    hdr_cell_st = Style(name="HdrCell", family="table-cell")
    hdr_cell_st.addElement(TableCellProperties(
        backgroundcolor="#2563EB", padding="0.15cm",
        border="0.05pt solid #1E3A8A"))
    doc.automaticstyles.addElement(hdr_cell_st)

    hdr_text_st = Style(name="HdrText", family="paragraph")
    hdr_text_st.addElement(TextProperties(
        color="#FFFFFF", fontweight="bold", fontsize="11pt"))
    hdr_text_st.addElement(ParagraphProperties(textalign="center"))
    doc.automaticstyles.addElement(hdr_text_st)

    even_cell_st = Style(name="EvenCell", family="table-cell")
    even_cell_st.addElement(TableCellProperties(
        backgroundcolor="#F1F5F9", padding="0.15cm",
        border="0.05pt solid #CBD5E1"))
    doc.automaticstyles.addElement(even_cell_st)

    odd_cell_st = Style(name="OddCell", family="table-cell")
    odd_cell_st.addElement(TableCellProperties(
        backgroundcolor="#FFFFFF", padding="0.15cm",
        border="0.05pt solid #CBD5E1"))
    doc.automaticstyles.addElement(odd_cell_st)

    body_text_st = Style(name="BodyText", family="paragraph")
    body_text_st.addElement(TextProperties(fontsize="10pt"))
    doc.automaticstyles.addElement(body_text_st)

    meta_text_st = Style(name="MetaText", family="paragraph")
    meta_text_st.addElement(TextProperties(fontsize="9pt", color="#64748B"))
    doc.automaticstyles.addElement(meta_text_st)

    # Heading + meta
    h = H(outlinelevel=1, text="逐句翻譯對照")
    doc.text.addElement(h)
    meta_p = P(stylename=meta_text_st, text="　・　".join(_meta_lines({**meta, "count": len(pairs)})))
    doc.text.addElement(meta_p)

    table = Table(name="translations")
    table.addElement(TableColumn(stylename=col_st))
    table.addElement(TableColumn(stylename=col_st))

    hdr = TableRow()
    for txt in ("原文", "譯文"):
        cell = TableCell(stylename=hdr_cell_st)
        cell.addElement(P(stylename=hdr_text_st, text=txt))
        hdr.addElement(cell)
    table.addElement(hdr)

    for i, p in enumerate(pairs):
        row = TableRow()
        cell_st = even_cell_st if (i % 2 == 1) else odd_cell_st
        for key in ("source", "target"):
            cell = TableCell(stylename=cell_st)
            cell.addElement(P(stylename=body_text_st, text=p.get(key) or ""))
            row.addElement(cell)
        table.addElement(row)
    doc.text.addElement(table)
    buf = io.BytesIO()
    doc.write(buf)
    return buf.getvalue()


def _build_pdf(pairs: list[dict], meta: dict) -> bytes:
    """A4 兩欄表格：標題列藍底白字、交替橫條淡底、首頁標題 + 檔名 + 翻譯時間。"""
    import fitz
    from app.core import font_catalog
    doc = fitz.open()
    # v1.7.53：修 CJK 字型載入錯誤 — best_cjk_path 簽章是 cjk:str
    # （"traditional" / "simplified"），之前傳 True 變成 falsy lookup → 永遠
    # 回 None → 落到 helv 字型 → 中文全變方框 / 缺字。回傳是 (Path, ttc_idx)
    # tuple，要正確解包；TTC 字型用 ttc_idx 指定 sub-font。
    cjk_font_path = None
    cjk_ttc_idx = 0
    try:
        result = font_catalog.best_cjk_path("sans", "traditional")
        if result:
            cjk_font_path, cjk_ttc_idx = result
    except Exception:
        pass
    font_alias = "cjk"
    def _register(p):
        nonlocal font_alias
        if cjk_font_path:
            try:
                # PyMuPDF insert_font 支援 TTC 指定 face index
                kwargs = {"fontname": font_alias, "fontfile": str(cjk_font_path)}
                if cjk_ttc_idx:
                    kwargs["set_simple"] = False  # CJK 用 CID 字型
                p.insert_font(**kwargs)
                return
            except Exception:
                pass
        font_alias = "helv"

    page = doc.new_page(width=595, height=842)
    _register(page)
    margin = 40
    col_w = (page.rect.width - margin * 2) / 2
    # v1.7.54：line_h 提到 20 — 實測 fontsize=10 時 insert_textbox 需 rect
    # height >= 18pt 才會渲染（h=16 仍回 -1.36 表示空間不足靜默失敗），
    # 所以單行 row_h = line_h(20) + pad*2(8) = 28，text rect 內高 20pt 才穩
    line_h = 20
    pad = 4
    y = margin

    # 文件標題
    page.insert_text((margin, y + 16), "逐句翻譯對照",
                     fontname=font_alias, fontsize=16, color=(0.12, 0.23, 0.54))
    y += 24
    # meta
    meta_text = "　・　".join(_meta_lines({**meta, "count": len(pairs)}))
    page.insert_text((margin, y + 12), meta_text,
                     fontname=font_alias, fontsize=9, color=(0.39, 0.45, 0.55))
    y += 22

    # 標題列藍底白字
    hdr_h = line_h + pad * 2
    page.draw_rect(fitz.Rect(margin, y, page.rect.width - margin, y + hdr_h),
                   color=(0.12, 0.23, 0.54), fill=(0.15, 0.39, 0.92))
    page.insert_text((margin + pad, y + line_h + 1),
                     "原文", fontname=font_alias, fontsize=11, color=(1, 1, 1))
    page.insert_text((margin + col_w + pad, y + line_h + 1),
                     "譯文", fontname=font_alias, fontsize=11, color=(1, 1, 1))
    y += hdr_h
    page_h = page.rect.height - margin

    def _est_lines(s: str, max_chars_ascii: int = 28, max_chars_cjk: int = 18) -> int:
        # CJK 字寬約 ASCII 1.6x；分開估行數取大者
        cjk_count = sum(1 for c in s if ord(c) > 0x2E80)
        ascii_count = len(s) - cjk_count
        # 等效 ASCII 寬度
        eff = ascii_count + cjk_count * (max_chars_ascii / max_chars_cjk)
        return max(1, int(eff // max_chars_ascii) + (1 if eff % max_chars_ascii else 0))

    for idx, p in enumerate(pairs):
        src = (p.get("source") or "").strip()
        tgt = (p.get("target") or "").strip()
        approx_lines = max(1, _est_lines(src), _est_lines(tgt))
        row_h = approx_lines * line_h + pad * 2
        if y + row_h > page_h:
            page = doc.new_page(width=595, height=842)
            _register(page)
            y = margin
        # 交替列底色
        if idx % 2 == 1:
            page.draw_rect(fitz.Rect(margin, y, page.rect.width - margin, y + row_h),
                           color=None, fill=(0.945, 0.957, 0.973))
        page.insert_textbox(
            fitz.Rect(margin + pad, y + pad, margin + col_w - pad, y + row_h - pad),
            src, fontname=font_alias, fontsize=10, color=(0, 0, 0),
        )
        page.insert_textbox(
            fitz.Rect(margin + col_w + pad, y + pad, page.rect.width - margin - pad, y + row_h - pad),
            tgt, fontname=font_alias, fontsize=10, color=(0, 0, 0),
        )
        # 邊框 + 中欄分隔
        page.draw_rect(fitz.Rect(margin, y, page.rect.width - margin, y + row_h),
                       color=(0.80, 0.84, 0.88), width=0.4, fill=None)
        page.draw_line(fitz.Point(margin + col_w, y),
                       fitz.Point(margin + col_w, y + row_h),
                       color=(0.80, 0.84, 0.88), width=0.4)
        y += row_h
    out = doc.tobytes(garbage=4, deflate=True)
    doc.close()
    return out


def _build_xlsx(pairs: list[dict], meta: dict) -> bytes:
    """Excel 樣式：標題列藍底白字粗體 freeze；A/B 欄寬 60；換行；交替列淡底；
    A1:B2 兩列頁首（檔名 / 翻譯時間 / 共 N 對）。"""
    import openpyxl
    from openpyxl.styles import Font, PatternFill, Alignment, Border, Side
    wb = openpyxl.Workbook()
    ws = wb.active
    ws.title = "translations"

    meta_lines = _meta_lines({**meta, "count": len(pairs)})
    meta_row_count = len(meta_lines)
    # 第 1～N 列是 meta（合併兩欄），N+1 列是標題列，N+2 起是資料
    for i, line in enumerate(meta_lines, start=1):
        ws.cell(row=i, column=1, value=line).font = Font(size=10, color="64748B")
        ws.merge_cells(start_row=i, start_column=1, end_row=i, end_column=2)
    hdr_row = meta_row_count + 1

    ws.cell(row=hdr_row, column=1, value="原文")
    ws.cell(row=hdr_row, column=2, value="譯文")
    hdr_font = Font(bold=True, color="FFFFFF", size=11)
    hdr_fill = PatternFill("solid", fgColor="2563EB")
    hdr_align = Alignment(horizontal="center", vertical="center")
    border_thin = Border(
        left=Side(style="thin", color="CBD5E1"),
        right=Side(style="thin", color="CBD5E1"),
        top=Side(style="thin", color="CBD5E1"),
        bottom=Side(style="thin", color="CBD5E1"),
    )
    for c in (1, 2):
        cell = ws.cell(row=hdr_row, column=c)
        cell.font = hdr_font
        cell.fill = hdr_fill
        cell.alignment = hdr_align
        cell.border = border_thin

    body_align = Alignment(wrap_text=True, vertical="top")
    # v1.7.53：原文 / 譯文兩欄用不同色系區分（暖黃 vs 冷綠），交替列再加深一階
    src_odd  = PatternFill("solid", fgColor="FFFBEB")  # 原文奇數列：淺黃
    src_even = PatternFill("solid", fgColor="FEF3C7")  # 原文偶數列：較深黃
    tgt_odd  = PatternFill("solid", fgColor="ECFDF5")  # 譯文奇數列：淺綠
    tgt_even = PatternFill("solid", fgColor="D1FAE5")  # 譯文偶數列：較深綠
    for i, p in enumerate(pairs, start=1):
        row = hdr_row + i
        ws.cell(row=row, column=1, value=p.get("source") or "")
        ws.cell(row=row, column=2, value=p.get("target") or "")
        for c in (1, 2):
            cell = ws.cell(row=row, column=c)
            cell.alignment = body_align
            cell.font = Font(size=10)
            cell.border = border_thin
            if c == 1:
                cell.fill = src_even if (i % 2 == 0) else src_odd
            else:
                cell.fill = tgt_even if (i % 2 == 0) else tgt_odd

    ws.column_dimensions["A"].width = 60
    ws.column_dimensions["B"].width = 60
    # freeze 標題列
    ws.freeze_panes = ws.cell(row=hdr_row + 1, column=1)
    buf = io.BytesIO()
    wb.save(buf)
    return buf.getvalue()


def _build_ods(pairs: list[dict], meta: dict) -> bytes:
    """ODF 試算表：標題列藍底白字、欄寬 8cm、交替列淡底、頁首 meta 列。"""
    from odf.opendocument import OpenDocumentSpreadsheet
    from odf.style import (Style, TableColumnProperties, TableCellProperties,
                           TextProperties)
    from odf.table import Table, TableColumn, TableRow, TableCell
    from odf.text import P

    doc = OpenDocumentSpreadsheet()
    col_st = Style(name="ColW8", family="table-column")
    col_st.addElement(TableColumnProperties(columnwidth="8cm"))
    doc.automaticstyles.addElement(col_st)

    hdr_cell_st = Style(name="HdrCell", family="table-cell")
    hdr_cell_st.addElement(TableCellProperties(
        backgroundcolor="#2563EB", border="0.05pt solid #1E3A8A"))
    doc.automaticstyles.addElement(hdr_cell_st)
    hdr_text_st = Style(name="HdrText", family="paragraph")
    hdr_text_st.addElement(TextProperties(
        color="#FFFFFF", fontweight="bold", fontsize="11pt"))
    doc.automaticstyles.addElement(hdr_text_st)

    # v1.7.53：原文 / 譯文兩欄不同色系（暖黃 vs 冷綠），交替列再加深
    src_odd_st  = Style(name="SrcOdd",  family="table-cell")
    src_odd_st.addElement(TableCellProperties(
        backgroundcolor="#FFFBEB", border="0.05pt solid #CBD5E1"))
    doc.automaticstyles.addElement(src_odd_st)
    src_even_st = Style(name="SrcEven", family="table-cell")
    src_even_st.addElement(TableCellProperties(
        backgroundcolor="#FEF3C7", border="0.05pt solid #CBD5E1"))
    doc.automaticstyles.addElement(src_even_st)
    tgt_odd_st  = Style(name="TgtOdd",  family="table-cell")
    tgt_odd_st.addElement(TableCellProperties(
        backgroundcolor="#ECFDF5", border="0.05pt solid #CBD5E1"))
    doc.automaticstyles.addElement(tgt_odd_st)
    tgt_even_st = Style(name="TgtEven", family="table-cell")
    tgt_even_st.addElement(TableCellProperties(
        backgroundcolor="#D1FAE5", border="0.05pt solid #CBD5E1"))
    doc.automaticstyles.addElement(tgt_even_st)

    meta_text_st = Style(name="MetaText", family="paragraph")
    meta_text_st.addElement(TextProperties(fontsize="9pt", color="#64748B"))
    doc.automaticstyles.addElement(meta_text_st)

    table = Table(name="translations")
    table.addElement(TableColumn(stylename=col_st))
    table.addElement(TableColumn(stylename=col_st))

    # meta 列（合併效果用單格 + 空格代替；ODS row 沒簡單 merge API）
    for line in _meta_lines({**meta, "count": len(pairs)}):
        meta_row = TableRow()
        cell = TableCell(valuetype="string")
        cell.addElement(P(stylename=meta_text_st, text=line))
        meta_row.addElement(cell)
        meta_row.addElement(TableCell(valuetype="string"))   # 空 B 欄
        table.addElement(meta_row)

    hdr = TableRow()
    for txt in ("原文", "譯文"):
        cell = TableCell(valuetype="string", stylename=hdr_cell_st)
        cell.addElement(P(stylename=hdr_text_st, text=txt))
        hdr.addElement(cell)
    table.addElement(hdr)

    for i, p in enumerate(pairs):
        row = TableRow()
        # v1.7.53：原文 / 譯文兩欄不同色系
        for key in ("source", "target"):
            if key == "source":
                st = src_even_st if (i % 2 == 1) else src_odd_st
            else:
                st = tgt_even_st if (i % 2 == 1) else tgt_odd_st
            cell = TableCell(valuetype="string", stylename=st)
            cell.addElement(P(text=p.get(key) or ""))
            row.addElement(cell)
        table.addElement(row)
    doc.spreadsheet.addElement(table)
    buf = io.BytesIO()
    doc.write(buf)
    return buf.getvalue()


# 文字檔不需要 meta 頁首；文件 / 試算表才會帶上原檔名 + 翻譯時間
_TEXT_FORMATS = {"txt", "md", "csv"}
_BUILDERS_TEXT = {"txt": _build_txt, "md": _build_md, "csv": _build_csv}
_BUILDERS_RICH = {
    "docx": _build_docx, "odt": _build_odt, "pdf": _build_pdf,
    "xlsx": _build_xlsx, "ods": _build_ods,
}


def _output_filename(fmt: str, source_filename: str) -> str:
    """匯出檔名 — 若有原檔名，取 stem 加 _translated.<fmt>；否則 translated.<fmt>。"""
    from pathlib import Path as _P
    if source_filename:
        stem = _P(source_filename).stem or "translated"
        # 過濾 path traversal / 只留檔名 stem 部分
        stem = stem.replace("/", "").replace("\\", "").replace("..", "").strip()[:80] or "translated"
        return f"{stem}_translated.{fmt}"
    return f"translated.{fmt}"


@router.post("/export")
async def export_translations(request: Request):
    """匯出翻譯結果到指定格式。Body: {format, pairs:[{source,target}],
    source_filename?, translated_at?}。文字檔（txt/md/csv）忽略 meta；文件 /
    試算表會在頁首帶原檔名 + 翻譯時間。"""
    try:
        body = await request.json()
    except Exception:
        raise HTTPException(400, "invalid JSON body")
    fmt = str(body.get("format") or "").strip().lower()
    if fmt not in _EXPORT_FORMATS:
        raise HTTPException(400, f"unsupported format: {fmt}")
    pairs = body.get("pairs") or []
    if not isinstance(pairs, list):
        raise HTTPException(400, "pairs must be a list")
    if len(pairs) > 10000:
        raise HTTPException(413, "pairs exceeds 10000 limit")
    cleaned = [p for p in pairs
               if isinstance(p, dict) and (p.get("source") or p.get("target"))]
    if not cleaned:
        raise HTTPException(400, "no non-empty pairs to export")
    source_filename = str(body.get("source_filename") or "").strip()[:200]
    translated_at = str(body.get("translated_at") or "").strip()[:64]
    meta = {"source_filename": source_filename, "translated_at": translated_at}

    # 日誌中 fmt 是 user input；雖然 _TEXT_FORMATS / _BUILDERS_RICH 是
    # whitelist 過的 enum，CodeQL 仍會旗 log injection。日誌只描述「哪個
    # builder 失敗」，不直接帶 fmt — exception() 已含完整 stack trace。
    import asyncio as _asyncio
    if fmt in _TEXT_FORMATS:
        builder = _BUILDERS_TEXT[fmt]
        try:
            data = await _asyncio.to_thread(builder, cleaned)
        except Exception as e:
            import logging as _lg
            _lg.getLogger("app.translate_doc").exception("text-format export builder failed")
            raise HTTPException(500, f"匯出 {fmt} 失敗：{e}")
    else:
        builder = _BUILDERS_RICH[fmt]
        try:
            data = await _asyncio.to_thread(builder, cleaned, meta)
        except Exception as e:
            import logging as _lg
            _lg.getLogger("app.translate_doc").exception("rich-format export builder failed")
            raise HTTPException(500, f"匯出 {fmt} 失敗：{e}")

    from app.core.http_utils import content_disposition
    fname = _output_filename(fmt, source_filename)
    return Response(
        content=data,
        media_type=_EXPORT_MIME[fmt],
        headers={"Content-Disposition": content_disposition(fname)},
    )
