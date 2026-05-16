"""Endpoints for the 擷取文字 tool."""
from __future__ import annotations

import logging
import re
import subprocess
import time as _time
import uuid
from pathlib import Path
from typing import Optional

logger = logging.getLogger("app.pdf_extract_text")

import fitz
from fastapi import APIRouter, File, Form, HTTPException, Request, UploadFile
from fastapi.responses import FileResponse, HTMLResponse, StreamingResponse

from ...config import settings
from ...core.llm_settings import llm_settings


router = APIRouter()


@router.get("/", response_class=HTMLResponse)
async def index(request: Request):
    templates = request.app.state.templates
    return templates.TemplateResponse("pdf_extract_text.html", {
        "request": request,
        # 只有在 admin 那邊勾「啟用 LLM」才顯示 LLM 重排功能；沒啟用的話
        # 整段 hint + 按鈕都不該出現，避免使用者點了發現要先去設定。
        "llm_enabled": llm_settings.is_enabled(),
    })


# ------------------------------------------------------------------- model

def _work_dir(bid: str) -> Path:
    d = settings.temp_dir / f"ext_text_{bid}"
    d.mkdir(parents=True, exist_ok=True)
    return d


def _extract_structured(src: Path) -> dict:
    """Read every page and return a structured doc:

        {
            "pages": [
                {"page": 1, "blocks": [
                    {"type": "heading"|"paragraph",
                     "level": 1|2|3|0,
                     "text": "...",
                     "size": 14.5,
                     "bold": False}
                ]},
                ...
            ],
            "median_size": 11.0,
        }

    Heading detection is heuristic: font sizes noticeably above the median
    size get mapped to H1/H2/H3. Line breaks inside the same block that
    don't end with terminal punctuation are joined into a paragraph.
    """
    pages = []
    all_sizes: list[float] = []

    with fitz.open(str(src)) as doc:
        for pno in range(doc.page_count):
            page = doc[pno]
            td = page.get_text("dict")
            page_blocks: list[dict] = []
            for block in td.get("blocks", []):
                if block.get("type") != 0:
                    continue
                lines = block.get("lines", [])
                if not lines:
                    continue
                # Flatten lines → joined text + average font size for this block.
                joined_lines: list[str] = []
                sizes: list[float] = []
                bolds: list[bool] = []
                for line in lines:
                    parts: list[str] = []
                    for span in line.get("spans", []):
                        t = span.get("text", "")
                        if t:
                            parts.append(t)
                        if span.get("size"):
                            sizes.append(float(span.get("size", 0)))
                            all_sizes.append(float(span.get("size", 0)))
                        # "flags": 16 = bold, 2 = italic (PyMuPDF)
                        flags = int(span.get("flags", 0) or 0)
                        bolds.append(bool(flags & 16))
                    ltxt = "".join(parts).strip()
                    if ltxt:
                        joined_lines.append(ltxt)
                if not joined_lines:
                    continue
                avg_size = (sum(sizes) / len(sizes)) if sizes else 0.0
                is_bold = (sum(bolds) / len(bolds)) > 0.5 if bolds else False
                # Join lines within a block into a single paragraph; keep
                # soft newline for lines ending with terminal punctuation
                # (Chinese/English 。！？./!?).
                merged = _join_block_lines(joined_lines)
                page_blocks.append({
                    "text": merged,
                    "size": avg_size,
                    "bold": is_bold,
                })
            pages.append({"page": pno + 1, "blocks": page_blocks})

    median_size = _median(all_sizes) if all_sizes else 11.0
    # Classify block type now that median is known.
    for p in pages:
        for b in p["blocks"]:
            b["type"], b["level"] = _classify_block(b["text"], b["size"],
                                                    b["bold"], median_size)
    # Second pass: merge adjacent paragraph blocks on the same page when the
    # previous block doesn't end with terminal punctuation AND both blocks
    # share similar font size. PyMuPDF's visual blocks are greedy about
    # splitting lines, so form-like documents end up with every visual line
    # its own "block". This pass reconstructs reading-flow paragraphs.
    for p in pages:
        merged: list[dict] = []
        for b in p["blocks"]:
            if (merged
                and b["type"] == "paragraph"
                and merged[-1]["type"] == "paragraph"
                and abs((merged[-1]["size"] or 0) - (b["size"] or 0)) < 0.5
                and not re.search(r"[。！？：\.\!\?]\s*$", merged[-1]["text"])
                and not re.match(r"^[\-\*•●◆]|^\d+[\.\)、]|^[A-Z][A-Z\s]+$",
                                  b["text"])):
                prev = merged[-1]["text"]
                cur = b["text"]
                if _is_cjk_char(prev[-1]) or _is_cjk_char(cur[0]):
                    merged[-1]["text"] = prev + cur
                else:
                    merged[-1]["text"] = prev + " " + cur
            else:
                merged.append(dict(b))
        p["blocks"] = merged
    return {"pages": pages, "median_size": median_size}


def _join_block_lines(lines: list[str]) -> str:
    """Join lines inside a visual block, deciding between soft-break (same
    paragraph, broken by wrap) and hard-break (genuine line separation)."""
    if not lines:
        return ""
    out = [lines[0]]
    for ln in lines[1:]:
        prev = out[-1]
        # Hard break if previous line ends with terminal punctuation
        # (paragraph break already in source), or current starts with a
        # list marker. Else soft-join.
        if re.search(r"[。！？：\.\!\?]$", prev) or re.match(r"^[\-\*•●◆●‣▪▫◦▶►]|^\d+[\.\)、]", ln):
            out.append(ln)
        else:
            # CJK lines should have no space inserted; Latin lines get a space.
            if _is_cjk_char(prev[-1]) or _is_cjk_char(ln[0]):
                out[-1] = prev + ln
            else:
                out[-1] = prev + " " + ln
    return "\n".join(out)


def _is_cjk_char(ch: str) -> bool:
    return ("一" <= ch <= "鿿" or
            "㐀" <= ch <= "䶿" or
            "　" <= ch <= "〿" or
            "＀" <= ch <= "￯")


def _median(vals: list[float]) -> float:
    if not vals:
        return 0.0
    s = sorted(vals)
    n = len(s)
    return s[n // 2] if n % 2 else (s[n // 2 - 1] + s[n // 2]) / 2


def _classify_block(text: str, size: float, bold: bool,
                    median: float) -> tuple[str, int]:
    """Return (type, level) for a block. Headings: size noticeably above
    median. Level 1 = largest, 3 = slightly larger. Otherwise paragraph."""
    if not text.strip():
        return "paragraph", 0
    if median <= 0:
        return "paragraph", 0
    ratio = size / median if median else 1.0
    if ratio >= 1.6:
        return "heading", 1
    if ratio >= 1.3:
        return "heading", 2
    if ratio >= 1.15 and bold:
        return "heading", 3
    return "paragraph", 0


# ----------------------------------------------------------------- outputs

def _render_txt(doc: dict) -> str:
    lines: list[str] = []
    for p in doc["pages"]:
        for b in p["blocks"]:
            lines.append(b["text"])
            lines.append("")
    return "\n".join(lines).rstrip() + "\n"


def _render_md(doc: dict) -> str:
    lines: list[str] = []
    for p in doc["pages"]:
        for b in p["blocks"]:
            if b["type"] == "heading":
                prefix = "#" * max(1, min(6, b["level"]))
                lines.append(f"{prefix} {b['text']}")
            else:
                lines.append(b["text"])
            lines.append("")
    return "\n".join(lines).rstrip() + "\n"


def _render_docx(doc: dict, out: Path) -> bool:
    """Best-effort docx via python-docx if available."""
    try:
        from docx import Document  # type: ignore
    except Exception:
        return False
    d = Document()
    for p in doc["pages"]:
        for b in p["blocks"]:
            if b["type"] == "heading":
                d.add_heading(b["text"], level=max(1, min(9, b["level"])))
            else:
                d.add_paragraph(b["text"])
    d.save(str(out))
    return True


def _render_odt_from_docx(docx_path: Path, out_dir: Path) -> Optional[Path]:
    """Convert .docx → .odt using soffice (OxOffice / LibreOffice)."""
    from ...core.office_convert import find_soffice
    exe = find_soffice()
    if not exe or not Path(exe).exists():
        return None
    try:
        subprocess.run(
            [exe, "--headless", "--convert-to", "odt",
             "--outdir", str(out_dir), str(docx_path)],
            check=True, capture_output=True, timeout=60,
        )
    except Exception:
        return None
    odt = out_dir / (docx_path.stem + ".odt")
    return odt if odt.exists() else None


# ----------------------------------------------------------------- routes

@router.post("/extract")
async def extract(request: Request, file: UploadFile = File(...)):
    if not (file.filename or "").lower().endswith(".pdf"):
        raise HTTPException(400, "只支援 PDF")
    data = await file.read()
    if not data:
        raise HTTPException(400, "empty file")

    bid = uuid.uuid4().hex
    from ...core import upload_owner as _uo
    _uo.record(bid, request)
    wdir = _work_dir(bid)
    src = wdir / "src.pdf"
    src.write_bytes(data)
    stem = Path(file.filename or "document").stem
    request.state.upload_filename = file.filename or ""

    # CRITICAL: PyMuPDF / python-docx / soffice calls are sync C/IO work.
    # Running them inline blocks the asyncio event loop → every other user
    # on the server gets stuck waiting. We hit this with a >100MB PDF —
    # one core pinned at 99% would freeze the entire site for everyone.
    # Push the heavy work into a worker thread via asyncio.to_thread() so
    # the event loop stays free for other requests (sidebar nav, healthz,
    # other tools, …).
    import asyncio as _asyncio

    def _do_extract():
        doc = _extract_structured(src)
        import json as _json
        (wdir / "model.json").write_text(_json.dumps(doc, ensure_ascii=False),
                                         encoding="utf-8")
        (wdir / "stem.txt").write_text(stem, encoding="utf-8")
        txt = _render_txt(doc)
        md = _render_md(doc)
        (wdir / f"{stem}.txt").write_text(txt, encoding="utf-8")
        (wdir / f"{stem}.md").write_text(md, encoding="utf-8")
        docx_available = _render_docx(doc, wdir / f"{stem}.docx")
        odt_available = False
        if docx_available:
            try:
                odt = _render_odt_from_docx(wdir / f"{stem}.docx", wdir)
                odt_available = bool(odt)
            except Exception:
                odt_available = False
        return doc, md, docx_available, odt_available

    doc, md, docx_available, odt_available = await _asyncio.to_thread(_do_extract)

    total_chars = sum(len(b["text"]) for p in doc["pages"] for b in p["blocks"])
    return {
        "batch_id": bid,
        "filename": file.filename,
        "page_count": len(doc["pages"]),
        "block_count": sum(len(p["blocks"]) for p in doc["pages"]),
        "char_count": total_chars,
        "preview_md": md[:5000],   # first 5000 chars for on-page preview
        "has_docx": docx_available,
        "has_odt": odt_available,
        "downloads": {
            "txt": f"/tools/pdf-extract-text/download/{bid}/txt",
            "md":  f"/tools/pdf-extract-text/download/{bid}/md",
            "docx": f"/tools/pdf-extract-text/download/{bid}/docx" if docx_available else None,
            "odt":  f"/tools/pdf-extract-text/download/{bid}/odt"  if odt_available else None,
        },
    }


@router.get("/download/{batch_id}/{fmt}")
async def download(batch_id: str, fmt: str, request: Request):
    from app.core.safe_paths import require_uuid_hex
    from ...core import upload_owner
    require_uuid_hex(batch_id, "batch_id")
    upload_owner.require(batch_id, request)
    wdir = settings.temp_dir / f"ext_text_{batch_id}"
    if not wdir.exists():
        raise HTTPException(404, "batch 不存在或已過期")
    stem = (wdir / "stem.txt").read_text(encoding="utf-8").strip() or "document"
    ext_map = {"txt": ("txt", "text/plain"),
               "md":  ("md",  "text/markdown"),
               "docx": ("docx", "application/vnd.openxmlformats-officedocument.wordprocessingml.document"),
               "odt":  ("odt",  "application/vnd.oasis.opendocument.text")}
    if fmt not in ext_map:
        raise HTTPException(400, "unsupported format")
    ext, mime = ext_map[fmt]
    fp = wdir / f"{stem}.{ext}"
    if not fp.exists():
        raise HTTPException(404, "format 尚未產生")
    return FileResponse(str(fp), media_type=mime, filename=fp.name)


# ----------------------------------------------- optional LLM paragraph reflow

@router.post("/llm-reflow")
async def llm_reflow(request: Request):
    """Re-join paragraphs split by PDF line wrapping via the configured LLM.
    Streams NDJSON progress events so the frontend can draw a progress bar:

        {"type":"start","total":36}
        {"type":"progress","done":1,"total":36,"ok":1,"skipped":0}
        ...
        {"type":"done","preview_md":"..."}

    Any fatal error emits ``{"type":"error","message":"..."}`` as the last
    line (HTTP status still 200 because the stream already started).
    """
    import json as _json
    body = await request.json()
    batch_id = str(body.get("batch_id") or "")
    if not batch_id:
        raise HTTPException(400, "batch_id required")
    wdir = settings.temp_dir / f"ext_text_{batch_id}"
    if not wdir.exists():
        raise HTTPException(404, "batch 不存在或已過期")

    try:
        from ...core.llm_settings import llm_settings
    except Exception:
        raise HTTPException(503, "LLM 未設定")
    if not llm_settings.is_enabled():
        raise HTTPException(503, "LLM 功能未啟用；請到「LLM 設定」開啟。")

    model_path = wdir / "model.json"
    if not model_path.exists():
        raise HTTPException(410, "結構資料已過期，請重新擷取。")
    doc = _json.loads(model_path.read_text(encoding="utf-8"))

    client = llm_settings.make_client()
    if client is None:
        raise HTTPException(503, "LLM client 未就緒")
    # Per-tool 模型覆寫優先（admin 在 LLM 設定頁可指定）
    model_name = llm_settings.get_model_for("pdf-extract-text")
    if not model_name:
        raise HTTPException(503, "尚未設定 LLM 模型。")

    # Build a worklist of every paragraph. We don't pre-filter by "\n"
    # anymore because the extractor already merges visual-block line wraps,
    # so paragraphs broken ACROSS blocks (the common case in form-like
    # PDFs) look single-line to us but still need LLM merging with the
    # next block. Headings are left alone.
    worklist: list[tuple[int, int]] = []
    for pi, p in enumerate(doc["pages"]):
        for bi, b in enumerate(p["blocks"]):
            if b["type"] != "paragraph":
                continue
            # Skip obviously trivial lines (single word, numbers-only, etc.)
            if len((b["text"] or "").strip()) < 4:
                continue
            worklist.append((pi, bi))
    total = len(worklist)

    def _sse(event: dict) -> bytes:
        return (_json.dumps(event, ensure_ascii=False) + "\n").encode("utf-8")

    def _iter():
        yield _sse({"type": "start", "total": total})
        ok = 0
        skipped = 0
        if total == 0:
            stem = (wdir / "stem.txt").read_text(encoding="utf-8").strip() or "document"
            md = _render_md(doc)
            yield _sse({"type": "done", "preview_md": md[:5000],
                        "ok": 0, "skipped": 0, "total": 0})
            return
        from app.core.log_safe import safe_log
        logger.info(
            "LLM reflow start: batch=%s model=%s total=%d",
            safe_log(batch_id), safe_log(model_name), total,
        )
        for done, (pi, bi) in enumerate(worklist, start=1):
            b = doc["pages"][pi]["blocks"][bi]
            # Emit a "working" event BEFORE the LLM call so the UI stops
            # looking frozen during the (often multi-second) text_query.
            snippet = b["text"].replace("\n", " ")[:40]
            yield _sse({
                "type": "working", "done": done - 1, "total": total,
                "page": pi + 1, "snippet": snippet,
                "length": len(b["text"]),
            })
            prompt = (
                "以下是從 PDF 擷取的一段文字，因為排版關係有些地方斷行錯誤，"
                "請在不改變任何字詞、只處理換行的前提下，把因版面斷行而錯斷"
                "的句子重新合併成自然的段落；保留真正的段落分隔。"
                "只回傳重排後的內文，不要加任何前後說明：\n\n"
                + b["text"]
            )
            t0 = _time.time()
            logger.info(
                "LLM reflow [%d/%d] p%d len=%d snippet=%r",
                done, total, pi + 1, len(b["text"]), snippet,
            )
            last_err: str | None = None
            try:
                reply = client.text_query(prompt, model=model_name,
                                          temperature=0.1, max_tokens=2048,
                                          think=False)
                elapsed = _time.time() - t0
                # Belt-and-braces: strip any <think>…</think> blocks that
                # a reasoning model may emit despite our no_think flags.
                reply = re.sub(r"<think>.*?</think>", "", reply or "",
                               flags=re.DOTALL).strip()
                if reply:
                    b["text"] = reply
                    ok += 1
                    logger.info(
                        "LLM reflow [%d/%d] OK in %.1fs (in=%d → out=%d chars)",
                        done, total, elapsed, len(b["text"]), len(reply),
                    )
                else:
                    skipped += 1
                    last_err = "LLM returned empty reply"
                    logger.warning(
                        "LLM reflow [%d/%d] SKIP (empty reply) in %.1fs",
                        done, total, elapsed,
                    )
            except Exception as exc:
                elapsed = _time.time() - t0
                skipped += 1
                last_err = f"{type(exc).__name__}: {exc}"
                logger.exception(
                    "LLM reflow [%d/%d] FAIL in %.1fs — %s",
                    done, total, elapsed, last_err,
                )
            ev: dict = {"type": "progress", "done": done, "total": total,
                        "ok": ok, "skipped": skipped}
            if last_err:
                ev["last_error"] = last_err
            yield _sse(ev)

        # Re-render outputs once every paragraph has been processed.
        stem = (wdir / "stem.txt").read_text(encoding="utf-8").strip() or "document"
        (wdir / f"{stem}.txt").write_text(_render_txt(doc), encoding="utf-8")
        md = _render_md(doc)
        (wdir / f"{stem}.md").write_text(md, encoding="utf-8")
        try:
            _render_docx(doc, wdir / f"{stem}.docx")
        except Exception:
            pass
        model_path.write_text(_json.dumps(doc, ensure_ascii=False),
                              encoding="utf-8")
        yield _sse({"type": "done", "preview_md": md[:5000],
                    "ok": ok, "skipped": skipped, "total": total})

    return StreamingResponse(_iter(), media_type="application/x-ndjson")


# ---- 對外 API：單次 upload + JSON 回傳所有文字 ----
@router.post("/api/pdf-extract-text", include_in_schema=True)
async def api_pdf_extract_text(request: Request, file: UploadFile = File(...)):
    """單次上傳 PDF，回 JSON：每頁文字 + 結構化區塊。"""
    if not (file.filename or "").lower().endswith(".pdf"):
        raise HTTPException(400, "只支援 PDF")
    data = await file.read()
    if not data or data[:4] != b"%PDF":
        raise HTTPException(400, "不是有效的 PDF")
    bid = uuid.uuid4().hex
    from ...core import upload_owner as _uo
    _uo.record(bid, request)
    wdir = _work_dir(bid)
    src = wdir / "src.pdf"
    src.write_bytes(data)
    import asyncio as _asyncio
    doc_model = await _asyncio.to_thread(_extract_structured, src)
    # Build per-page text (joined blocks) for callers who just need text.
    pages_text: list[dict] = []
    for pg in doc_model["pages"]:
        text = "\n".join(b["text"] for b in pg["blocks"])
        pages_text.append({"page": pg["page"], "text": text,
                           "blocks": pg["blocks"]})
    return {
        "filename": file.filename,
        "page_count": len(doc_model["pages"]),
        "pages": pages_text,
    }
