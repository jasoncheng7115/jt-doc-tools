"""v1.8.63 jtdt-reform engine 單元 + 端對端測試。

B 路線 — 完全不靠 pdf2docx 重建 docx 的新引擎。

模組覆蓋：
- table_detector: drawings → grid → cells
- paragraph_grouper: blocks → tables / free paragraphs
- docx_builder: DocumentModel → docx
- engine.convert_via_jtdt_reform: 端對端
"""
from __future__ import annotations

from pathlib import Path

import pytest
from PIL import Image
from docx import Document
from docx.oxml.ns import qn

from app.tools.pdf_to_office.engines.jtdt_reform import (
    convert_via_jtdt_reform,
)
from app.tools.pdf_to_office.engines.jtdt_reform.table_detector import (
    TableRegion, _classify_lines, _cluster_coords, block_in_table,
    detect_tables,
)
from app.tools.pdf_to_office.engines.jtdt_reform.paragraph_grouper import (
    build_document_model,
)
from app.tools.pdf_to_office.engines.jtdt_reform.docx_builder import build_docx
from app.tools.pdf_to_office.pdf_truth.models import (
    PDFBlock, PDFDrawing, PDFImage, PDFLine, PDFPage, PDFTruth,
)


def _line(text, bbox, size=12):
    return PDFLine(chars=[], bbox=bbox, text=text,
                   dominant_font="Helvetica", dominant_size=size)


def _block(text, bbox, page_num=0, size=12):
    lines = [_line(text, bbox, size=size)]
    return PDFBlock(lines=lines, bbox=bbox, text=text, block_type="text",
                    page_num=page_num, dominant_font="Helvetica",
                    dominant_size=size)


def _page(blocks=None, drawings=None, images=None, width=612, height=792,
          page_num=0):
    return PDFPage(page_num=page_num, width=width, height=height,
                   margin_top=0, margin_bottom=0, margin_left=0, margin_right=0,
                   blocks=blocks or [], images=images or [],
                   drawings=drawings or [])


def _truth(pages, body_font_size=12.0):
    return PDFTruth(pages=pages, fonts=[], total_pages=len(pages),
                    language_guess="en", body_font_size=body_font_size,
                    body_font_name="Helvetica")


# --- table_detector --------------------------------------------------------

def test_cluster_coords():
    """容忍 ±2pt 群聚。"""
    assert _cluster_coords([100, 101, 200, 201.5, 300]) == pytest.approx(
        [100.5, 200.75, 300])
    assert _cluster_coords([]) == []


def test_classify_lines_horizontal_vertical():
    """line drawings 拆 h/v。"""
    h_line = PDFDrawing(type="line", bbox=(50, 100, 500, 100), page_num=0)
    v_line = PDFDrawing(type="line", bbox=(100, 50, 100, 500), page_num=0)
    rect = PDFDrawing(type="rect", bbox=(60, 60, 300, 300), page_num=0)
    h, v = _classify_lines([h_line, v_line, rect])
    assert len(h) >= 2  # h_line + rect top + bottom
    assert len(v) >= 2  # v_line + rect left + right


def test_detect_tables_simple_grid():
    """繪一個 3-row × 2-col 的 grid (用 rect drawings) → 偵測到 1 table。"""
    # 3 row × 2 col grid; horizontal lines at y=100, 130, 160, 190; vertical at x=50, 300, 550
    drawings = []
    for y in (100, 130, 160, 190):
        drawings.append(PDFDrawing(type="line", bbox=(50, y, 550, y), page_num=0))
    for x in (50, 300, 550):
        drawings.append(PDFDrawing(type="line", bbox=(x, 100, x, 190), page_num=0))
    page = _page(drawings=drawings)
    tables = detect_tables(_truth([page]))
    assert len(tables) == 1
    t = tables[0]
    assert len(t.row_ys) == 4  # 4 horiz lines
    assert len(t.col_xs) == 3  # 3 vert lines
    assert len(t.cells) == 3   # 3 rows
    assert len(t.cells[0]) == 2  # 2 cols


def test_block_in_table_assignment():
    """Block 中心點落入 cell bbox → 配到該 (row, col)。"""
    drawings = [
        PDFDrawing(type="line", bbox=(50, 100, 550, 100), page_num=0),
        PDFDrawing(type="line", bbox=(50, 150, 550, 150), page_num=0),
        PDFDrawing(type="line", bbox=(50, 100, 50, 150), page_num=0),
        PDFDrawing(type="line", bbox=(300, 100, 300, 150), page_num=0),
        PDFDrawing(type="line", bbox=(550, 100, 550, 150), page_num=0),
    ]
    page = _page(drawings=drawings)
    tables = detect_tables(_truth([page]))
    assert len(tables) == 1
    t = tables[0]
    # block at center (175, 125) → row 0, col 0
    pos = block_in_table((150, 110, 200, 140), t)
    assert pos == (0, 0)
    # block at center (425, 125) → row 0, col 1
    pos = block_in_table((400, 110, 450, 140), t)
    assert pos == (0, 1)
    # block outside table
    pos = block_in_table((20, 700, 50, 720), t)
    assert pos is None


# --- paragraph_grouper -----------------------------------------------------

def test_grouper_assigns_blocks_to_cells():
    """Blocks 落入 cell → cell_blocks 內；落外 → free_blocks。"""
    drawings = [
        PDFDrawing(type="line", bbox=(50, 100, 550, 100), page_num=0),
        PDFDrawing(type="line", bbox=(50, 150, 550, 150), page_num=0),
        PDFDrawing(type="line", bbox=(50, 100, 50, 150), page_num=0),
        PDFDrawing(type="line", bbox=(300, 100, 300, 150), page_num=0),
        PDFDrawing(type="line", bbox=(550, 100, 550, 150), page_num=0),
    ]
    blocks = [
        _block("Cell A", (150, 110, 200, 140)),    # inside col 0
        _block("Cell B", (400, 110, 450, 140)),    # inside col 1
        _block("FREE", (50, 700, 200, 720)),       # outside
    ]
    page = _page(blocks=blocks, drawings=drawings)
    model = build_document_model(_truth([page]))
    assert len(model.pages) == 1
    pm = model.pages[0]
    assert len(pm.tables) == 1
    tm = pm.tables[0]
    assert tm.cell_blocks[0][0][0].text == "Cell A"
    assert tm.cell_blocks[0][1][0].text == "Cell B"
    assert len(pm.free_blocks) == 1
    assert pm.free_blocks[0].block.text == "FREE"


def test_grouper_heading_detection():
    """字級 > body × 1.2 視為 heading。"""
    blocks = [_block("BIG TITLE", (50, 50, 500, 90), size=24)]
    page = _page(blocks=blocks)
    model = build_document_model(_truth([page], body_font_size=12.0))
    assert len(model.pages[0].free_blocks) == 1
    assert model.pages[0].free_blocks[0].is_heading is True


# --- docx_builder ----------------------------------------------------------

def test_build_docx_basic(tmp_path):
    """簡單 PDFTruth → docx 輸出檔存在 + 有 body paragraphs + table。"""
    drawings = [
        PDFDrawing(type="line", bbox=(50, 100, 550, 100), page_num=0),
        PDFDrawing(type="line", bbox=(50, 150, 550, 150), page_num=0),
        PDFDrawing(type="line", bbox=(50, 100, 50, 150), page_num=0),
        PDFDrawing(type="line", bbox=(300, 100, 300, 150), page_num=0),
        PDFDrawing(type="line", bbox=(550, 100, 550, 150), page_num=0),
    ]
    blocks = [
        _block("Title above", (50, 50, 500, 80)),
        _block("Cell A", (150, 110, 200, 140)),
        _block("Cell B", (400, 110, 450, 140)),
        _block("Footer below", (50, 700, 200, 720)),
    ]
    page = _page(blocks=blocks, drawings=drawings)
    model = build_document_model(_truth([page]))
    out_path = tmp_path / "out.docx"
    # build_docx 需要 pdf_path 拉圖片，這裡無圖傳 dummy 路徑
    stats = build_docx(model, tmp_path / "fake.pdf", out_path)
    assert stats["ok"] is True
    assert out_path.exists()

    doc = Document(str(out_path))
    # 至少有「Title above」「Footer below」段落
    body_texts = [p.text for p in doc.paragraphs]
    assert any("Title above" in t for t in body_texts)
    assert any("Footer below" in t for t in body_texts)
    # 有 1 個 table
    assert len(doc.tables) == 1
    t = doc.tables[0]
    assert t.rows[0].cells[0].text.strip() == "Cell A"
    assert t.rows[0].cells[1].text.strip() == "Cell B"


# --- engine.convert_via_jtdt_reform 端對端 -----------------------------------

def test_convert_via_jtdt_reform_end_to_end(tmp_path):
    """造一個簡單 PDF（用 PIL 寫成空白 PDF），跑 native engine。
    無 drawings → 全部成 free paragraphs。"""
    img = Image.new("RGB", (612, 792), color="white")
    pdf_path = tmp_path / "blank.pdf"
    img.save(str(pdf_path), "PDF")
    docx_path = tmp_path / "out.docx"
    res = convert_via_jtdt_reform(pdf_path, docx_path)
    assert res["ok"] is True
    assert res["engine"] == "jtdt-reform"
    assert docx_path.exists()


def test_convert_via_jtdt_reform_missing_pdf(tmp_path):
    with pytest.raises(FileNotFoundError):
        convert_via_jtdt_reform(tmp_path / "nope.pdf", tmp_path / "out.docx")


def test_extractor_polygon_with_fill_becomes_rect(tmp_path):
    """polygon path（4 條 line 組成封閉填色形狀）必須被視為一個 fill rect drawing。

    Regression: v1.8.78 之前 extractor 只認 type='re' 的子項，忽略 polygon path
    with fill。報價單橙色 banner 用 4-line polygon 實作 → docx render 看不到 banner。
    """
    import fitz

    from app.tools.pdf_to_office.pdf_truth.extractor import extract_pdf_truth

    pdf_path = tmp_path / "poly_banner.pdf"
    doc = fitz.open()
    page = doc.new_page(width=595, height=842)
    # 4 條 line 組成的封閉多邊形（仿報價單 banner）— 用 draw_polyline + fill
    rect_pts = [
        fitz.Point(0, 0),
        fitz.Point(595, 0),
        fitz.Point(595, 110),
        fitz.Point(0, 110),
        fitz.Point(0, 0),
    ]
    shape = page.new_shape()
    for i in range(len(rect_pts) - 1):
        shape.draw_line(rect_pts[i], rect_pts[i + 1])
    shape.finish(fill=(1.0, 0.957, 0.925), color=None)
    shape.commit()
    doc.save(str(pdf_path))
    doc.close()

    truth = extract_pdf_truth(pdf_path)
    fill_rects = [
        d for d in truth.pages[0].drawings
        if d.type == "rect" and d.fill_color
    ]
    # 必須至少抓到一個全寬填色 rect — 不能因為是 polygon 就跳過
    banners = [d for d in fill_rects
               if (d.bbox[2] - d.bbox[0]) > 500 and (d.bbox[3] - d.bbox[1]) > 50]
    assert banners, "polygon-with-fill 沒被偵測成 rect drawing"
    b = banners[0]
    # 顏色受 PDF round-trip rounding 影響（0xeb/0xec/0xed 都可能），比範圍
    assert b.fill_color.startswith("#"), f"fill_color 不像 hex: {b.fill_color}"
    r = int(b.fill_color[1:3], 16); g = int(b.fill_color[3:5], 16); bl = int(b.fill_color[5:7], 16)
    assert r >= 250 and 240 <= g <= 250 and 230 <= bl <= 240, \
        f"polygon fill 顏色不對: rgb=({r},{g},{bl})"
