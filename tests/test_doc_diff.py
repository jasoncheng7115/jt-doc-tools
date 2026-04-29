"""Tests for the renamed 文件差異比對 tool (formerly pdf-diff).

Covers: PDF×PDF compare, accept-list extension validation, redirect from
old `/tools/pdf-diff` to `/tools/doc-diff`, and the new Office-input path
(skipped if soffice isn't installed on the test runner)."""
from __future__ import annotations

import io

import fitz
import pytest
from fastapi.testclient import TestClient

from app.main import app
from app.core import office_convert


def _pdf_with_text(lines: list[str]) -> bytes:
    doc = fitz.open()
    p = doc.new_page(width=595, height=842)
    p.insert_text((50, 80), "\n".join(lines), fontsize=12, fontname="helv")
    out = io.BytesIO()
    doc.save(out); doc.close()
    return out.getvalue()


@pytest.fixture
def client():
    return TestClient(app)


def test_index_renders_new_name(client):
    r = client.get("/tools/doc-diff/")
    assert r.status_code == 200
    assert "文件差異比對" in r.text
    # The accept attribute should advertise the office extensions
    assert ".docx" in r.text and ".odt" in r.text


def test_legacy_pdf_diff_url_redirects(client):
    """Bookmarks pointing at /tools/pdf-diff/ must keep working (308 → doc-diff).

    308 is used (not 301/302) so the method + body are preserved on POST —
    legacy API callers hitting /tools/pdf-diff/compare keep working without
    silently downgrading to GET.
    """
    r = client.get("/tools/pdf-diff/", follow_redirects=False)
    assert r.status_code == 308
    assert r.headers["location"].endswith("/tools/doc-diff/")
    # Without trailing slash
    r = client.get("/tools/pdf-diff", follow_redirects=False)
    assert r.status_code == 308
    # Sub-path: a POST to the old /compare path must redirect (and the
    # client following it should land on the new compare endpoint).
    a = _pdf_with_text(["x"]); b = _pdf_with_text(["y"])
    r = client.post(
        "/tools/pdf-diff/compare", follow_redirects=False,
        files={
            "file_a": ("a.pdf", a, "application/pdf"),
            "file_b": ("b.pdf", b, "application/pdf"),
        },
    )
    assert r.status_code == 308
    assert r.headers["location"].endswith("/tools/doc-diff/compare")


def test_compare_two_pdfs_returns_diff(client):
    a = _pdf_with_text(["alpha", "beta", "gamma"])
    b = _pdf_with_text(["alpha", "BETA-changed", "gamma", "delta"])
    r = client.post(
        "/tools/doc-diff/compare",
        files={
            "file_a": ("a.pdf", a, "application/pdf"),
            "file_b": ("b.pdf", b, "application/pdf"),
        },
    )
    assert r.status_code == 200, r.text
    body = r.json()
    assert body["page_count_a"] == 1 and body["page_count_b"] == 1
    totals = body["totals"]
    # Some non-zero diff expected (changed line + inserted line)
    assert totals["added"] + totals["removed"] + totals["changed"] >= 2


def test_compare_rejects_unsupported_extension(client):
    a = _pdf_with_text(["x"])
    r = client.post(
        "/tools/doc-diff/compare",
        files={
            "file_a": ("a.pdf", a, "application/pdf"),
            "file_b": ("b.png", b"\x89PNG\r\n\x1a\n", "image/png"),
        },
    )
    assert r.status_code == 400
    assert "不支援" in r.json()["detail"]


@pytest.mark.skipif(
    office_convert.find_soffice() is None,
    reason="soffice (OxOffice/LibreOffice) not installed on this runner",
)
def test_compare_office_to_pdf_works(client):
    """End-to-end with a real Office input: build a tiny .docx via python-docx,
    compare it against a PDF version of (similar) text, expect a 200 response."""
    try:
        from docx import Document
    except Exception:
        pytest.skip("python-docx not available")
    docx_buf = io.BytesIO()
    d = Document()
    d.add_paragraph("hello world")
    d.add_paragraph("second line")
    d.save(docx_buf)

    pdf_b = _pdf_with_text(["hello world", "different line"])
    r = client.post(
        "/tools/doc-diff/compare",
        files={
            "file_a": ("a.docx", docx_buf.getvalue(),
                       "application/vnd.openxmlformats-officedocument.wordprocessingml.document"),
            "file_b": ("b.pdf", pdf_b, "application/pdf"),
        },
    )
    assert r.status_code == 200, r.text
    body = r.json()
    assert body["page_count_a"] >= 1 and body["page_count_b"] == 1


def test_migration_renames_pdf_diff_to_doc_diff(tmp_path):
    """`_m3_rename_pdf_diff_to_doc_diff` must rewrite role_perms +
    subject_perms rows on existing installs so users don't silently lose
    access to the renamed tool after upgrade."""
    import sqlite3
    from app.core.auth_db import _m1_initial, _m2_username_source_unique, \
        _m3_rename_pdf_diff_to_doc_diff

    conn = sqlite3.connect(":memory:")
    conn.row_factory = sqlite3.Row
    _m1_initial(conn)
    _m2_username_source_unique(conn)

    # Pre-populate with the OLD tool id, mimicking a v1.1.60-or-earlier install.
    import time
    now = time.time()
    conn.execute(
        "INSERT INTO roles(id, display_name, description, is_builtin, is_protected, created_at) "
        "VALUES (?,?,?,?,?,?)",
        ("legal-sec", "法務資安", "", 1, 0, now),
    )
    conn.execute("INSERT INTO role_perms(role_id, tool_id) VALUES (?, ?)",
                 ("legal-sec", "pdf-diff"))
    conn.execute(
        "INSERT INTO subject_perms(subject_type, subject_key, tool_id) "
        "VALUES (?, ?, ?)",
        ("user", "alice@local", "pdf-diff"),
    )
    conn.commit()

    # Apply the new migration.
    _m3_rename_pdf_diff_to_doc_diff(conn)

    # Old rows should be gone, new rows in their place.
    rp = conn.execute(
        "SELECT tool_id FROM role_perms WHERE role_id='legal-sec'").fetchall()
    assert {r["tool_id"] for r in rp} == {"doc-diff"}
    sp = conn.execute(
        "SELECT tool_id FROM subject_perms WHERE subject_key='alice@local'").fetchall()
    assert {r["tool_id"] for r in sp} == {"doc-diff"}


def test_migration_idempotent_when_doc_diff_already_exists(tmp_path):
    """If admin manually pre-granted `doc-diff` (e.g. on v1.1.61 with no upgrade),
    re-running the migration must not blow up on the unique constraint."""
    import sqlite3
    from app.core.auth_db import _m1_initial, _m2_username_source_unique, \
        _m3_rename_pdf_diff_to_doc_diff

    conn = sqlite3.connect(":memory:")
    conn.row_factory = sqlite3.Row
    _m1_initial(conn)
    _m2_username_source_unique(conn)

    import time
    now = time.time()
    conn.execute(
        "INSERT INTO roles(id, display_name, description, is_builtin, is_protected, created_at) "
        "VALUES (?,?,?,?,?,?)",
        ("legal-sec", "法務資安", "", 1, 0, now),
    )
    # Both old and new exist for the same role
    conn.executemany(
        "INSERT INTO role_perms(role_id, tool_id) VALUES (?, ?)",
        [("legal-sec", "pdf-diff"), ("legal-sec", "doc-diff")],
    )
    conn.commit()

    _m3_rename_pdf_diff_to_doc_diff(conn)  # must not raise

    rp = conn.execute(
        "SELECT tool_id FROM role_perms WHERE role_id='legal-sec'").fetchall()
    assert {r["tool_id"] for r in rp} == {"doc-diff"}
