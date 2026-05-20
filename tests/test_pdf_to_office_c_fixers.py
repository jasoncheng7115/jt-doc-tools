"""v1.8.61 C 階段強化 fixer 測試。

- C1: paragraph_line_split 走 table 內 w:p（含 table 內段）
- C2: text_recovery line-level missing 偵測（TEST / 台灣 regression）
- C4: table_borders_from_image — image-based 每表邊線偵測
- C5: table_empty_cell_recovery — empty cell PDFTruth y-mapping 補回（「1 單位」regression）
"""
from __future__ import annotations

import io
import os
import tempfile
from pathlib import Path

import numpy as np
import pytest
from PIL import Image, ImageDraw
from docx import Document
from docx.oxml.ns import qn

from app.tools.pdf_to_office.pdf_truth.models import (
    PDFBlock, PDFDrawing, PDFImage, PDFLine, PDFPage, PDFTruth,
)
from app.tools.pdf_to_office.postprocess.fixers import (
    paragraph_line_split, table_borders_from_image, table_empty_cell_recovery,
    text_recovery,
)


# --- helpers -----------------------------------------------------------------

def _make_line(text, bbox):
    return PDFLine(chars=[], bbox=bbox, text=text,
                   dominant_font="Helvetica", dominant_size=12)


def _make_block(lines_with_bbox, page_num=0):
    lines = [_make_line(t, b) for t, b in lines_with_bbox]
    if lines:
        x0 = min(b[0] for _, b in lines_with_bbox)
        y0 = min(b[1] for _, b in lines_with_bbox)
        x1 = max(b[2] for _, b in lines_with_bbox)
        y1 = max(b[3] for _, b in lines_with_bbox)
        bbox = (x0, y0, x1, y1)
    else:
        bbox = (0, 0, 0, 0)
    text = "\n".join(t for t, _ in lines_with_bbox)
    return PDFBlock(lines=lines, bbox=bbox, text=text,
                    block_type="text", page_num=page_num,
                    dominant_font="Helvetica", dominant_size=12)


def _make_page(blocks=None, drawings=None, images=None,
               width=612, height=792, page_num=0):
    return PDFPage(
        page_num=page_num, width=width, height=height,
        margin_top=0, margin_bottom=0, margin_left=0, margin_right=0,
        blocks=blocks or [], images=images or [], drawings=drawings or [],
    )


def _make_truth(pages):
    return PDFTruth(pages=pages, fonts=[], total_pages=len(pages),
                    language_guess="en")


# --- C1 paragraph_line_split 走 table 內 -------------------------------------

def test_paragraph_line_split_inside_table_cell():
    """段落「TCC 406Taiwan」在 docx table cell 內 → 該 cell 內 paragraph 也要被
    paragraph_line_split 拆。"""
    doc = Document()
    tbl = doc.add_table(rows=1, cols=1)
    cell = tbl.rows[0].cells[0]
    cell.text = "TCC 406Taiwan"

    blk = _make_block([
        ("Some other line at top", (400, 50, 580, 70)),
        ("Address line", (400, 80, 580, 100)),
        ("TCC 406", (400, 110, 480, 125)),
        ("Taiwan", (400, 130, 460, 145)),
    ])
    truth = _make_truth([_make_page(blocks=[blk])])
    res = paragraph_line_split.fix_paragraph_line_split(doc, truth, None)
    assert res["split"] >= 1
    # cell 內現在應該有 2 段：TCC 406 + Taiwan
    cell_paras = [p.text for p in cell.paragraphs]
    assert "TCC 406" in cell_paras
    assert "Taiwan" in cell_paras
    assert "TCC 406Taiwan" not in cell_paras


# --- C2 text_recovery line-level --------------------------------------------

def test_text_recovery_line_level_test_regression():
    """「TEST」是頁底飄浮 absolute-position 短行；PDFTruth 內有但 docx 漏抓 → 補。"""
    doc = Document()
    doc.add_paragraph("報價單抬頭")
    doc.add_paragraph("Some invoice content here.")

    blocks = [
        _make_block([("報價單抬頭", (50, 50, 500, 70))]),
        _make_block([("Some invoice content here.", (50, 80, 500, 100))]),
        _make_block([("TEST", (50, 600, 100, 620))]),  # short floating block missing
    ]
    truth = _make_truth([_make_page(blocks=blocks)])
    res = text_recovery.fix_text_recovery(doc, truth, None)
    assert res["recovered"] >= 1
    texts = [p.text for p in doc.paragraphs]
    assert "TEST" in texts


def test_text_recovery_line_within_block_top_band():
    """4 行 address block 第 4 行漏抓 + 整 block 位於頂部 band → D1 補回。
    這是 v1.8.61 引入 + v1.8.62 D1 保守化後仍能補回的情境（block 整段在 top band）。"""
    doc = Document()
    doc.add_paragraph("範例企業有限公司, 測試員")
    doc.add_paragraph("台中市北屯區松山街71巷3號1樓")
    doc.add_paragraph("TCC 406")
    # 「台灣」未進 docx

    address_block = _make_block([
        ("範例企業有限公司, 測試員", (50, 50, 350, 70)),
        ("台中市北屯區松山街71巷3號1樓", (50, 75, 350, 95)),
        ("TCC 406", (50, 100, 200, 115)),
        ("台灣", (50, 115, 150, 118)),  # y_top=115 / 792 = 0.145 → top band
    ])
    truth = _make_truth([_make_page(blocks=[address_block], height=792)])
    res = text_recovery.fix_text_recovery(doc, truth, None)
    assert res["recovered"] >= 1
    texts = [p.text for p in doc.paragraphs]
    assert "台灣" in texts


def test_text_recovery_line_within_block_mid_band_skipped():
    """D1 保守化：missing line 在 mid band (0.15 < y/h < 0.6) → skip。
    避免「台灣」黏到「2023 年 06 月 06 日」前面這類 v1.8.61 regression。"""
    doc = Document()
    for t in ("F.", "S.", "T."):
        doc.add_paragraph(t)
    blocks = [
        _make_block([("F.", (50, 50, 100, 70))]),
        _make_block([("S.", (50, 100, 100, 120))]),
        _make_block([("T.", (50, 600, 100, 620))]),
        _make_block([("MIDLINE", (50, 400, 150, 420))]),  # 400/792=0.5 mid
    ]
    truth = _make_truth([_make_page(blocks=blocks, height=792)])
    res = text_recovery.fix_text_recovery(doc, truth, None)
    assert res["recovered"] == 0
    assert "MIDLINE" not in [p.text for p in doc.paragraphs]


def test_text_recovery_skips_when_already_in_table_cell():
    """文字已在 table cell 內 → 不重複補成 body 段落。"""
    doc = Document()
    tbl = doc.add_table(rows=1, cols=1)
    tbl.rows[0].cells[0].text = "已在表格內"

    blk = _make_block([("已在表格內", (50, 50, 200, 70))])
    truth = _make_truth([_make_page(blocks=[blk])])
    res = text_recovery.fix_text_recovery(doc, truth, None)
    assert res["recovered"] == 0


# --- C4 table_borders_from_image -------------------------------------------

def _make_pdf_with_image_at_path(image: Image.Image, pdf_path: Path):
    """把 PIL image 寫成單頁 PDF。"""
    image.convert("RGB").save(str(pdf_path), "PDF")


def test_image_border_detection_no_lines(tmp_path):
    """造一個白底 PDF，render 後該區應該偵測不到實線。"""
    img = Image.new("RGB", (612, 792), color="white")
    pdf = tmp_path / "blank.pdf"
    _make_pdf_with_image_at_path(img, pdf)

    arr, scale = table_borders_from_image._render_page_gray(str(pdf), 0)
    # crop 中央區域
    crop = arr[100:300, 100:500]
    has, dirs = table_borders_from_image._table_has_visible_borders(crop)
    assert has is False
    assert dirs == {"horiz": False, "vert": False}


def test_image_border_detection_with_horizontal_line(tmp_path):
    """造一個白底 PDF 加一條水平黑線 → 偵測得到。"""
    img = Image.new("RGB", (612, 792), color="white")
    draw = ImageDraw.Draw(img)
    draw.line([(50, 400), (560, 400)], fill="black", width=2)
    pdf = tmp_path / "horiz.pdf"
    _make_pdf_with_image_at_path(img, pdf)

    arr, scale = table_borders_from_image._render_page_gray(str(pdf), 0)
    # render 後座標是 px = pt * 2 (144 DPI / 72 = 2x)
    # PIL 圖 612x792 (pt) → render 1224x1584; line at y=400 (pt) → y=800 (px)
    # 我們是直接把 PIL image 當 PDF page，所以 PDF 內部尺寸就是 image 尺寸（不是 612x792 pt）
    # 取整張 array 給 detector
    has, dirs = table_borders_from_image._table_has_visible_borders(arr)
    assert has is True
    assert dirs["horiz"] is True


def test_image_border_strip_no_match(tmp_path):
    """docx 表格 cell 內容找不到 PDFTruth 對應 block → 該表 skip。"""
    img = Image.new("RGB", (612, 792), color="white")
    pdf = tmp_path / "blank.pdf"
    _make_pdf_with_image_at_path(img, pdf)

    doc = Document()
    tbl = doc.add_table(rows=2, cols=2)
    tbl.rows[0].cells[0].text = "OnlyInDocx"

    truth = _make_truth([_make_page()])  # empty blocks → no match
    res = table_borders_from_image.fix_table_borders_from_image(
        doc, truth, None, pdf_path=str(pdf))
    assert res["stripped"] == 0
    assert res["failed"] == 1


def test_longest_run():
    arr = np.array([True, True, False, True, True, True, False, True])
    assert table_borders_from_image._longest_run(arr) == 3


# --- C5 table_empty_cell_recovery -------------------------------------------

def test_table_empty_cell_recovery_basic():
    """造 docx table 3x2 一個 empty cell；PDFTruth 該 region 內有「1 單位」未被 docx
    抽到 → 補進該 empty cell。"""
    doc = Document()
    tbl = doc.add_table(rows=3, cols=2)
    tbl.rows[0].cells[0].text = "編號"
    tbl.rows[0].cells[1].text = "數量"
    tbl.rows[1].cells[0].text = "1"
    # tbl.rows[1].cells[1] left empty — should be filled with "1 單位"
    tbl.rows[2].cells[0].text = "2"
    tbl.rows[2].cells[1].text = "2 單位"

    blocks = [
        _make_block([("編號", (50, 100, 100, 115))]),
        _make_block([("數量", (200, 100, 250, 115))]),
        _make_block([("1", (50, 130, 60, 145))]),
        _make_block([("1 單位", (200, 130, 260, 145))]),  # missing in docx
        _make_block([("2", (50, 160, 60, 175))]),
        _make_block([("2 單位", (200, 160, 260, 175))]),
    ]
    truth = _make_truth([_make_page(blocks=blocks)])

    res = table_empty_cell_recovery.fix_table_empty_cell_recovery(doc, truth, None)
    assert res["filled"] >= 1
    # 補回應該在 tbl.rows[1].cells[1]
    assert "1 單位" in tbl.rows[1].cells[1].text


def test_table_empty_cell_recovery_no_pdf_truth():
    doc = Document()
    doc.add_table(rows=1, cols=1)
    res = table_empty_cell_recovery.fix_table_empty_cell_recovery(doc, None, None)
    assert res["filled"] == 0


def test_table_empty_cell_recovery_no_empty_cells():
    """所有 cell 都有內容 → 不動。"""
    doc = Document()
    tbl = doc.add_table(rows=1, cols=2)
    tbl.rows[0].cells[0].text = "A"
    tbl.rows[0].cells[1].text = "B"
    blocks = [_make_block([("A", (50, 50, 60, 65))]),
              _make_block([("B", (100, 50, 110, 65))])]
    truth = _make_truth([_make_page(blocks=blocks)])
    res = table_empty_cell_recovery.fix_table_empty_cell_recovery(doc, truth, None)
    assert res["filled"] == 0
