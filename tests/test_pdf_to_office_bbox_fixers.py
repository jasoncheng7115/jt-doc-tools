"""Sprint B 新 fixer 單元測試：
- table_bbox_width：從 pdfplumber 真值欄寬覆寫 docx tblGrid + tcW
- bbox_layout：form-vs-article 偵測 + 安全 reorder
- image_position_fix：bbox-size fallback 配對 + 段落水平對齊

策略：用合成 docx + 假 PDFTruth + monkeypatch pdfplumber，不依賴實體 PDF。
"""
from __future__ import annotations

import io

import pytest
from docx import Document
from docx.oxml.ns import qn
from docx.shared import Pt

from app.tools.pdf_to_office.pdf_truth.models import (
    PDFBlock, PDFImage, PDFLine, PDFPage, PDFTruth, PDFDrawing,
)
from app.tools.pdf_to_office.postprocess.fixers import (
    bbox_layout, image_position_fix, table_bbox_width,
)


# --- 共用 helper -------------------------------------------------------------

def _make_pdf_block(text, bbox, page_num=0, n_lines=1):
    lines = []
    for i in range(n_lines):
        lines.append(PDFLine(
            chars=[], bbox=bbox, text=text,
            dominant_font="Helvetica", dominant_size=12,
        ))
    return PDFBlock(
        lines=lines, bbox=bbox, text=text, block_type="text",
        page_num=page_num, dominant_font="Helvetica", dominant_size=12,
    )


def _make_page(width=612, height=792, blocks=None, drawings=None,
               images=None, page_num=0):
    return PDFPage(
        page_num=page_num, width=width, height=height,
        margin_top=0, margin_bottom=0, margin_left=0, margin_right=0,
        blocks=blocks or [], images=images or [], drawings=drawings or [],
    )


def _make_truth(pages):
    return PDFTruth(pages=pages, fonts=[], total_pages=len(pages),
                    language_guess="en")


# --- table_bbox_width 測試 ---------------------------------------------------

def _make_docx_with_table(rows_text):
    """造一個 docx 含 1 個表格，內容由 rows_text 提供（list[list[str]]）。"""
    doc = Document()
    tbl = doc.add_table(rows=len(rows_text), cols=len(rows_text[0]))
    for r_idx, row_cells in enumerate(rows_text):
        for c_idx, txt in enumerate(row_cells):
            tbl.rows[r_idx].cells[c_idx].text = txt
    return doc, tbl


def test_table_bbox_width_basic_apply(monkeypatch):
    """pdfplumber 回欄寬 [50pt, 80pt, 120pt]，3 欄 docx 表格 → tblGrid + tcW
    應寫成對應 twips。"""
    rows = [["A", "B", "C"], ["d", "e", "f"]]
    doc, tbl = _make_docx_with_table(rows)

    # mock pdfplumber 抽結果
    def fake_extract(pdf_path):
        return [{
            "text": [["A", "B", "C"], ["d", "e", "f"]],
            "widths_pt": [50.0, 80.0, 120.0],
            "page_num": 0,
        }]
    monkeypatch.setattr(table_bbox_width,
                        "_extract_pdfplumber_tables_with_widths",
                        fake_extract)

    truth = _make_truth([_make_page()])
    res = table_bbox_width.fix_table_bbox_width(
        doc, truth, None, pdf_path="/fake/path.pdf")
    assert res["applied_tables"] == 1
    assert res["cells_changed"] == 6  # 2 rows × 3 cols

    # 驗 tblGrid
    grid = tbl._element.find(qn("w:tblGrid"))
    assert grid is not None
    cols = grid.findall(qn("w:gridCol"))
    widths = [int(c.get(qn("w:w"))) for c in cols]
    assert widths == [1000, 1600, 2400]  # 50/80/120 × 20 twips

    # 驗第一列 cell tcW
    first_row = tbl.rows[0]._element
    tcs = first_row.findall(qn("w:tc"))
    for tc, expected_w in zip(tcs, widths):
        tcW = tc.find(qn("w:tcPr")).find(qn("w:tcW"))
        assert tcW.get(qn("w:type")) == "dxa"
        assert int(tcW.get(qn("w:w"))) == expected_w


def test_table_bbox_width_col_count_mismatch_skipped(monkeypatch):
    """pdfplumber 給 5 欄但 docx 只有 3 欄（差 > 3）→ skip。"""
    rows = [["A", "B", "C"], ["d", "e", "f"]]
    doc, tbl = _make_docx_with_table(rows)

    def fake_extract(pdf_path):
        return [{
            "text": [["A", "B", "C", "D", "E"]],
            "widths_pt": [40, 40, 40, 40, 40],
            "page_num": 0,
        }]
    monkeypatch.setattr(table_bbox_width,
                        "_extract_pdfplumber_tables_with_widths",
                        fake_extract)

    truth = _make_truth([_make_page()])
    res = table_bbox_width.fix_table_bbox_width(
        doc, truth, None, pdf_path="/fake/path.pdf")
    # shape_match 因 col 差 2 仍可能過 → 但 col 對齊不上會 skip
    assert res["applied_tables"] == 0


def test_table_bbox_width_spacer_col_merged(monkeypatch):
    """pdfplumber 多出 1 個窄 spacer 欄（< 3 pt）→ 自動合進鄰欄。"""
    rows = [["A", "B", "C"], ["d", "e", "f"]]
    doc, tbl = _make_docx_with_table(rows)

    def fake_extract(pdf_path):
        return [{
            "text": [["A", "", "B", "C"]],
            "widths_pt": [50.0, 2.0, 78.0, 120.0],  # 2pt spacer 應合並
            "page_num": 0,
        }]
    monkeypatch.setattr(table_bbox_width,
                        "_extract_pdfplumber_tables_with_widths",
                        fake_extract)

    truth = _make_truth([_make_page()])
    res = table_bbox_width.fix_table_bbox_width(
        doc, truth, None, pdf_path="/fake/path.pdf")
    assert res["applied_tables"] == 1

    grid = tbl._element.find(qn("w:tblGrid"))
    widths = [int(c.get(qn("w:w"))) for c in grid.findall(qn("w:gridCol"))]
    # spacer 2pt 合到左 → [52pt, 78pt, 120pt] × 20 = [1040, 1560, 2400]
    assert widths == [1040, 1560, 2400]


def test_table_bbox_width_no_pdf_path():
    doc = Document()
    res = table_bbox_width.fix_table_bbox_width(doc, None, None, pdf_path=None)
    assert res["applied_tables"] == 0
    assert "no pdf_path" in res["skipped"]


# --- bbox_layout form-vs-article 測試 ----------------------------------------

def test_classify_form_drawings_density():
    """drawings > 60 → 判 form。"""
    drawings = [PDFDrawing(type="line", bbox=(0, 0, 100, 0), page_num=0)
                for _ in range(70)]
    page = _make_page(drawings=drawings)
    res = bbox_layout._classify_form_or_article(page, [], {})
    assert res["is_form"] is True
    assert any("drawings>" in r for r in res["reasons"])


def test_classify_form_short_blocks():
    """80% 以上 block 是短欄位（≤ 2 行 + < 30 字元）→ 判 form。"""
    blocks = []
    for i in range(8):
        blocks.append(_make_pdf_block("Label", (0, i * 20, 80, i * 20 + 15)))
    # 2 長 block 不影響 80% 比例
    blocks.append(_make_pdf_block("A very long paragraph with more than thirty characters here",
                                  (0, 200, 400, 215), n_lines=3))
    page = _make_page(blocks=blocks)
    res = bbox_layout._classify_form_or_article(page, blocks, {})
    assert res["is_form"] is True


def test_classify_article_with_long_paragraphs():
    """全是長段落 → 不判 form。"""
    blocks = []
    long_text = "This is a long paragraph with multi lines and well over thirty characters of body text. " * 3
    for i in range(6):
        blocks.append(_make_pdf_block(long_text, (50, i * 100, 500, i * 100 + 80),
                                       n_lines=4))
    page = _make_page(blocks=blocks)
    res = bbox_layout._classify_form_or_article(page, blocks, {})
    assert res["is_form"] is False


def test_classify_form_wide_block_crosses_boundary():
    """多欄被偵測，但有 block 橫跨 boundary 且寬度 > 頁寬 60% → 判 form。"""
    # 頁寬 600pt; boundary 假設在 300; 一個 block 從 50 到 550 跨界
    blocks = [_make_pdf_block("Header spanning entire width", (50, 50, 550, 80))]
    page = _make_page(width=600, height=800, blocks=blocks)
    mc_info = {"is_multi_column": True, "boundary_x": 300}
    res = bbox_layout._classify_form_or_article(page, blocks, mc_info)
    assert res["is_form"] is True
    assert any("wide_block_crosses_boundary" in r for r in res["reasons"])


def test_safe_reorder_runs_basic():
    """造 4 段落，後 2 段 Y 序錯亂，但都在連續 run 內 → 應 swap。"""
    doc = Document()
    p1 = doc.add_paragraph("AAAA first")   # sort_key Y=10
    p2 = doc.add_paragraph("BBBB second")  # sort_key Y=20
    p3 = doc.add_paragraph("DDDD fourth")  # sort_key Y=40, but appears 3rd
    p4 = doc.add_paragraph("CCCC third")   # sort_key Y=30, but appears 4th

    body = doc.element.body
    children = list(body)
    para_tag = qn("w:p")
    para_children = [(i, c) for i, c in enumerate(children) if c.tag == para_tag]
    # 對應 sort_key：p1=10, p2=20, p3=40, p4=30
    sort_keys = {0: (0, 0, 10), 1: (0, 0, 20), 2: (0, 0, 40), 3: (0, 0, 30)}
    matched = [(ci, None, sort_keys[idx])
               for idx, (ci, _) in enumerate(para_children)]

    moved = bbox_layout._safe_reorder_runs(body, children, matched)
    assert moved >= 2

    # 重讀 body 內段落順序 — 應是 AAAA, BBBB, CCCC, DDDD
    new_paras = [c for c in doc.element.body if c.tag == para_tag]
    texts = []
    for c in new_paras:
        ts = []
        for r in c.findall(qn("w:r")):
            for t in r.findall(qn("w:t")):
                if t.text:
                    ts.append(t.text)
        texts.append("".join(ts))
    assert texts == ["AAAA first", "BBBB second", "CCCC third", "DDDD fourth"]


def test_safe_reorder_runs_split_by_table():
    """段落 [A,B,C,D] 中間 idx 2 是 table → 兩 run [A,B] 和 [D] (single 不動)。
    這裡用 matched_idx_set 不含 idx 2 模擬 table 隔開。"""
    matched = [
        (0, None, (0, 0, 30)),
        (1, None, (0, 0, 20)),  # 跟 idx 0 反序
        # idx 2 = table (not in set)
        (3, None, (0, 0, 50)),
        (4, None, (0, 0, 60)),  # 正序
    ]
    matched_idx_set = {0, 1, 3, 4}
    runs = bbox_layout._group_consecutive_runs(matched, matched_idx_set)
    assert [[m[0] for m in r] for r in runs] == [[0, 1], [3, 4]]


def test_bbox_layout_form_page_skips_reorder():
    """form-classified 頁的段落不參與 reorder。"""
    doc = Document()
    doc.add_paragraph("First line here for form page.")
    doc.add_paragraph("Second line here for form page.")
    doc.add_paragraph("Third line here for form page.")

    # 假頁面有 70 個 drawings → 判 form
    drawings = [PDFDrawing(type="line", bbox=(0, 0, 100, 0), page_num=0)
                for _ in range(70)]
    blocks = [
        _make_pdf_block("First line here for form page.", (50, 100, 500, 120)),
        _make_pdf_block("Second line here for form page.", (50, 130, 500, 150)),
        _make_pdf_block("Third line here for form page.", (50, 160, 500, 180)),
    ]
    page = _make_page(blocks=blocks, drawings=drawings)
    truth = _make_truth([page])

    res = bbox_layout.fix_bbox_layout(doc, truth, None)
    assert res["matched_paragraphs"] == 0  # form page 全跳過
    assert 1 in res["form_pages"]
    assert res["reordered_paragraphs"] == 0


# --- image_position_fix 測試 -------------------------------------------------

def test_infer_alignment():
    img_center = PDFImage(bbox=(250, 100, 350, 200), page_num=0, xref=1,
                          width=100, height=100)
    img_right = PDFImage(bbox=(450, 100, 550, 200), page_num=0, xref=2,
                         width=100, height=100)
    img_left = PDFImage(bbox=(20, 100, 120, 200), page_num=0, xref=3,
                        width=100, height=100)
    assert image_position_fix._infer_alignment(img_center, 612) == "center"
    assert image_position_fix._infer_alignment(img_right, 612) == "right"
    assert image_position_fix._infer_alignment(img_left, 612) == "left"
    assert image_position_fix._infer_alignment(img_center, 0) is None


def test_match_by_bbox_size_aspect_ratio():
    """SHA1 失敗時，按長寬比 + 寬度近似配對。"""
    EMU = image_position_fix.EMU_PER_PT
    # docx image 100pt × 50pt = 2:1 aspect
    docx_img = {"cx": 100 * EMU, "cy": 50 * EMU, "hash": ""}
    # PDF candidates: a 200×100 (aspect match), b 100×100 (no match)
    pdf_match = PDFImage(bbox=(0, 0, 200, 100), page_num=0, xref=1,
                         width=200, height=100, image_hash="x")
    pdf_nope = PDFImage(bbox=(0, 0, 100, 100), page_num=0, xref=2,
                        width=100, height=100, image_hash="y")
    res = image_position_fix._match_by_bbox_size(docx_img, [pdf_nope, pdf_match], set())
    assert res is pdf_match


def test_match_by_bbox_size_no_match_when_too_different():
    EMU = image_position_fix.EMU_PER_PT
    docx_img = {"cx": 100 * EMU, "cy": 50 * EMU, "hash": ""}
    pdf_far = PDFImage(bbox=(0, 0, 100, 100), page_num=0, xref=1,
                       width=100, height=100, image_hash="z")
    res = image_position_fix._match_by_bbox_size(docx_img, [pdf_far], set())
    assert res is None
