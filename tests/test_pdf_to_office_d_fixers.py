"""v1.8.62 D 階段 fixer 測試。

D1: text_recovery 大保守化（只補 top/bot band，中段 skip）
D2: link_text_recovery（超連結文字補回）
D3: table_unmerge_with_pdf_labels（vMerge 拆解）
D4: table_normalize cell vAlign center（既有改動，已 covered by existing）
D5: table_empty_cell_recovery 條件放寬
"""
from __future__ import annotations

from pathlib import Path

import pytest
from docx import Document
from docx.oxml.ns import qn

from app.tools.pdf_to_office.pdf_truth.models import (
    PDFBlock, PDFLine, PDFPage, PDFTruth,
)
from app.tools.pdf_to_office.postprocess.fixers import (
    table_unmerge_with_pdf_labels, text_recovery,
)


def _make_line(text, bbox):
    return PDFLine(chars=[], bbox=bbox, text=text,
                   dominant_font="Helvetica", dominant_size=12)


def _make_block(lines_with_bbox, page_num=0):
    lines = [_make_line(t, b) for t, b in lines_with_bbox]
    if lines:
        bbox = (min(b[0] for _, b in lines_with_bbox),
                min(b[1] for _, b in lines_with_bbox),
                max(b[2] for _, b in lines_with_bbox),
                max(b[3] for _, b in lines_with_bbox))
    else:
        bbox = (0, 0, 0, 0)
    return PDFBlock(lines=lines, bbox=bbox,
                    text="\n".join(t for t, _ in lines_with_bbox),
                    block_type="text", page_num=page_num,
                    dominant_font="Helvetica", dominant_size=12)


def _make_page(blocks=None, drawings=None, images=None,
               width=612, height=792, page_num=0):
    return PDFPage(page_num=page_num, width=width, height=height,
                   margin_top=0, margin_bottom=0, margin_left=0, margin_right=0,
                   blocks=blocks or [], images=images or [], drawings=drawings or [])


def _make_truth(pages):
    return PDFTruth(pages=pages, fonts=[], total_pages=len(pages),
                    language_guess="en")


# --- D1 text_recovery 大保守化 ---------------------------------------------

def test_text_recovery_top_band_inserted_at_start():
    """y_top 在頂部 (< 0.15) → 補到 body 開頭。"""
    doc = Document()
    for t in ("First.", "Second.", "Third."):
        doc.add_paragraph(t)
    # 多放幾個 anchor 讓 miss_ratio < 50%
    blocks = [
        _make_block([("First.", (50, 200, 500, 220))]),
        _make_block([("Second.", (50, 250, 500, 270))]),
        _make_block([("Third.", (50, 300, 500, 320))]),
        _make_block([("TITLE", (50, 50, 500, 80))]),  # top band only missing
    ]
    truth = _make_truth([_make_page(blocks=blocks, height=792)])
    res = text_recovery.fix_text_recovery(doc, truth, None)
    assert res["recovered"] == 1
    assert res["band_breakdown"]["top"] == 1
    paras = [p.text for p in doc.paragraphs]
    assert paras[0] == "TITLE"


def test_text_recovery_bottom_band_appended_at_end():
    """y_top 在底部 (> 0.6) → append 到 body 末尾。"""
    doc = Document()
    for t in ("First.", "Second.", "Third."):
        doc.add_paragraph(t)
    blocks = [
        _make_block([("First.", (50, 50, 500, 70))]),
        _make_block([("Second.", (50, 100, 500, 120))]),
        _make_block([("Third.", (50, 200, 500, 220))]),
        _make_block([("TEST", (50, 700, 100, 720))]),  # bot
    ]
    truth = _make_truth([_make_page(blocks=blocks, height=792)])
    res = text_recovery.fix_text_recovery(doc, truth, None)
    assert res["recovered"] == 1
    assert res["band_breakdown"]["bot"] == 1
    paras = [p.text for p in doc.paragraphs]
    assert paras[-1] == "TEST"


def test_text_recovery_mid_band_skipped():
    """y_top 在中段 (0.15-0.6) → skip。"""
    doc = Document()
    for t in ("First.", "Second.", "Third."):
        doc.add_paragraph(t)
    blocks = [
        _make_block([("First.", (50, 50, 500, 70))]),
        _make_block([("Second.", (50, 100, 500, 120))]),
        _make_block([("Third.", (50, 200, 500, 220))]),
        _make_block([("MIDLINE", (50, 400, 100, 420))]),  # mid
    ]
    truth = _make_truth([_make_page(blocks=blocks, height=792)])
    res = text_recovery.fix_text_recovery(doc, truth, None)
    assert res["recovered"] == 0
    assert res["band_breakdown"]["mid_skipped"] == 1


def test_text_recovery_substring_skipped():
    """missing text 是 docx 任一 text 的 substring → skip。"""
    doc = Document()
    for t in ("First.", "Second.", "台灣 年 06 月 06 日"):
        doc.add_paragraph(t)
    blocks = [
        _make_block([("First.", (50, 50, 500, 70))]),
        _make_block([("Second.", (50, 100, 500, 120))]),
        _make_block([("台灣 年 06 月 06 日", (50, 300, 300, 320))]),
        _make_block([("台灣", (50, 100, 100, 115))]),  # top band but substring
    ]
    truth = _make_truth([_make_page(blocks=blocks, height=792)])
    res = text_recovery.fix_text_recovery(doc, truth, None)
    assert res["recovered"] == 0
    assert res["band_breakdown"]["substring_skipped"] >= 1


def test_text_recovery_placeholder_skipped():
    """填空線 / 日期留白模板 → skip。"""
    doc = Document()
    for t in ("First.", "Second.", "Third."):
        doc.add_paragraph(t)
    blocks = [
        _make_block([("First.", (50, 50, 500, 70))]),
        _make_block([("Second.", (50, 100, 500, 120))]),
        _make_block([("Third.", (50, 200, 500, 220))]),
        _make_block([("______", (50, 60, 200, 75))]),
        _make_block([("年 月 日", (50, 700, 200, 720))]),
    ]
    truth = _make_truth([_make_page(blocks=blocks, height=792)])
    res = text_recovery.fix_text_recovery(doc, truth, None)
    assert res["recovered"] == 0
    assert res["band_breakdown"]["placeholder_skipped"] >= 2


# --- D3 table_unmerge_with_pdf_labels --------------------------------------

def _make_table_with_vmerge():
    """造一個 docx table 第 2 列 col 1+2 被 vMerge 上方的 cell。"""
    doc = Document()
    tbl = doc.add_table(rows=2, cols=3)
    # row 0: 業務聯絡人 | 電話 | 信箱
    tbl.rows[0].cells[0].text = "業務聯絡人"
    tbl.rows[0].cells[1].text = "電話"
    tbl.rows[0].cells[2].text = "信箱"
    # row 1: 財會聯絡人 | (vmerge from row 0 col 1) | (vmerge from row 0 col 2)
    tbl.rows[1].cells[0].text = "財會聯絡人"
    # 手動加 vMerge 到 row 0 col 1+2 (restart) + row 1 col 1+2 (continue)
    for col_idx in (1, 2):
        for row_idx, val in ((0, "restart"), (1, "continue")):
            tc_el = tbl.rows[row_idx]._element.findall(qn("w:tc"))[col_idx]
            tcPr = tc_el.find(qn("w:tcPr"))
            if tcPr is None:
                tcPr = tc_el.makeelement(qn("w:tcPr"), {})
                tc_el.insert(0, tcPr)
            vMerge = tcPr.makeelement(qn("w:vMerge"), {})
            vMerge.set(qn("w:val"), val)
            tcPr.append(vMerge)
    # 清掉 row 1 col 1, 2 的文字（continue 該空）
    for c_idx in (1, 2):
        tc_el = tbl.rows[1]._element.findall(qn("w:tc"))[c_idx]
        for p in tc_el.findall(qn("w:p")):
            for r in p.findall(qn("w:r")):
                p.remove(r)
    return doc, tbl


def test_table_unmerge_basic():
    """PDF 該位置有獨立 label「電話」「信箱」→ 拆 vMerge + 補 label。"""
    doc, tbl = _make_table_with_vmerge()
    # PDFTruth：兩列各有獨立 cell text
    blocks = [
        _make_block([("業務聯絡人", (50, 100, 150, 120))]),
        _make_block([("電話", (200, 100, 250, 120))]),
        _make_block([("信箱", (400, 100, 450, 120))]),
        _make_block([("財會聯絡人", (50, 130, 150, 150))]),
        _make_block([("電話", (200, 130, 250, 150))]),
        _make_block([("信箱", (400, 130, 450, 150))]),
    ]
    truth = _make_truth([_make_page(blocks=blocks)])

    res = table_unmerge_with_pdf_labels.fix_table_unmerge_with_pdf_labels(
        doc, truth, None)
    assert res["unmerged"] >= 1
    assert res["filled"] >= 1


def test_table_unmerge_skip_when_no_pdf_text():
    """PDF 該位置無對應 text → 不拆 vMerge（真的應該合併的 cell）。"""
    doc, tbl = _make_table_with_vmerge()
    # PDF 只有 row 0 labels — row 1 沒對應 PDF text
    blocks = [
        _make_block([("業務聯絡人", (50, 100, 150, 120))]),
        _make_block([("電話", (200, 100, 250, 120))]),
        _make_block([("信箱", (400, 100, 450, 120))]),
        _make_block([("財會聯絡人", (50, 130, 150, 150))]),
        # row 1 col 1, 2 在 PDF 上是空白（真的 merged 上方）
    ]
    truth = _make_truth([_make_page(blocks=blocks)])
    res = table_unmerge_with_pdf_labels.fix_table_unmerge_with_pdf_labels(
        doc, truth, None)
    assert res["unmerged"] == 0


def test_table_unmerge_no_tables():
    doc = Document()
    doc.add_paragraph("no table")
    res = table_unmerge_with_pdf_labels.fix_table_unmerge_with_pdf_labels(
        doc, _make_truth([_make_page()]), None)
    assert res["unmerged"] == 0


# --- D2 link_text_recovery — 大部分需 PDF 檔，這裡只測 helper -----------------

def test_link_text_recovery_no_pdf_path():
    from app.tools.pdf_to_office.postprocess.fixers import link_text_recovery
    res = link_text_recovery.fix_link_text_recovery(
        Document(), _make_truth([_make_page()]), None, pdf_path=None)
    assert res["recovered"] == 0
    assert "no pdf_path" in res["skipped"]


def test_link_text_in_bbox():
    """從 PDFTruth 撈 bbox 內 line text。"""
    from app.tools.pdf_to_office.postprocess.fixers import link_text_recovery
    blk = _make_block([("jason@jason.tools", (200, 300, 400, 320)),
                       ("無關文字", (50, 700, 100, 720))])
    truth = _make_truth([_make_page(blocks=[blk])])
    # link bbox 包含 jason@jason.tools 但不含「無關文字」
    text = link_text_recovery._text_in_bbox(truth, 0, (190, 290, 410, 330))
    assert "jason@jason.tools" in text
    assert "無關文字" not in text


# --- D4 table_normalize cell vAlign center --------------------------------

def test_table_normalize_all_cells_vertical_center():
    """v1.8.62: 所有 cell 一律 vAlign center（含內文列）。"""
    from app.tools.pdf_to_office.postprocess.fixers.table_normalize import (
        fix_table_normalize,
    )
    doc = Document()
    tbl = doc.add_table(rows=2, cols=2)
    tbl.rows[0].cells[0].text = "Header"
    tbl.rows[1].cells[0].text = "Body"
    fix_table_normalize(doc, _make_truth([_make_page()]), None)
    # 檢查 row 1 cell 0 (內文列) 也是 center
    tc_el = tbl.rows[1]._element.findall(qn("w:tc"))[0]
    tcPr = tc_el.find(qn("w:tcPr"))
    vAlign = tcPr.find(qn("w:vAlign"))
    assert vAlign is not None
    assert vAlign.get(qn("w:val")) == "center"
