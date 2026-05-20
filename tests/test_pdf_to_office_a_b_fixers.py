"""Sprint B 二階段 5 個 fixer 單元測試（v1.8.60）：
- paragraph_line_split：PDF block 多行強拆 docx 單段
- table_cell_dedup_text：cell 內重複行去重
- fake_table_remove（強化）：mostly-empty + sparse 表格移除
- text_recovery：PDFTruth 有但 docx 無的 short block 補回
- image_position_fix（補插）：同 hash 圖在 N bbox 但 docx 只 1 張時補
"""
from __future__ import annotations

import pytest
from docx import Document
from docx.oxml.ns import qn

from app.tools.pdf_to_office.pdf_truth.models import (
    PDFBlock, PDFDrawing, PDFImage, PDFLine, PDFPage, PDFTruth,
)
from app.tools.pdf_to_office.postprocess.fixers import (
    fake_table_remove, paragraph_line_split, table_cell_dedup_text,
    text_recovery,
)
from app.tools.pdf_to_office.postprocess.fixers.image_position_fix import (
    fix_image_position_fix,
)


# --- helpers -----------------------------------------------------------------

def _make_line(text, bbox):
    return PDFLine(chars=[], bbox=bbox, text=text,
                   dominant_font="Helvetica", dominant_size=12)


def _make_block(lines_with_bbox, page_num=0):
    """lines_with_bbox = [(text, bbox), ...]。回 PDFBlock 含這些 line。"""
    lines = [_make_line(t, b) for t, b in lines_with_bbox]
    if lines:
        x0 = min(b[0] for _, b in lines_with_bbox)
        y0 = min(b[1] for _, b in lines_with_bbox)
        x1 = max(b[2] for _, b in lines_with_bbox)
        y1 = max(b[3] for _, b in lines_with_bbox)
        bbox = (x0, y0, x1, y1)
    else:
        bbox = (0, 0, 0, 0)
    text = "\n".join(t for t, _ in lines_with_bbox)
    return PDFBlock(lines=lines, bbox=bbox, text=text,
                    block_type="text", page_num=page_num,
                    dominant_font="Helvetica", dominant_size=12)


def _make_page(blocks=None, drawings=None, images=None,
               width=612, height=792, page_num=0):
    return PDFPage(
        page_num=page_num, width=width, height=height,
        margin_top=0, margin_bottom=0, margin_left=0, margin_right=0,
        blocks=blocks or [], images=images or [], drawings=drawings or [],
    )


def _make_truth(pages):
    return PDFTruth(pages=pages, fonts=[], total_pages=len(pages),
                    language_guess="en")


# --- paragraph_line_split ----------------------------------------------------

def test_paragraph_line_split_basic():
    """docx 段落 'TCC 406Taiwan' 對應 PDF block 兩行 ['TCC 406', 'Taiwan'] → 拆。"""
    doc = Document()
    doc.add_paragraph("範例企業有限公司")
    doc.add_paragraph("TCC 406Taiwan")

    blk = _make_block([
        ("範例企業有限公司", (400, 50, 580, 70)),
        ("台中市北屯區松山街71巷3號1樓", (400, 75, 580, 95)),
        ("TCC 406", (400, 100, 480, 115)),
        ("Taiwan", (400, 120, 460, 135)),
    ])
    truth = _make_truth([_make_page(blocks=[blk])])

    res = paragraph_line_split.fix_paragraph_line_split(doc, truth, None)
    assert res["split"] >= 1
    assert res["pieces_inserted"] >= 1

    paras = [p.text for p in doc.paragraphs]
    assert "TCC 406Taiwan" not in paras
    assert "TCC 406" in paras
    assert "Taiwan" in paras


def test_paragraph_line_split_no_match_skips():
    """docx 文字跟 PDF block 對不上 → skip。"""
    doc = Document()
    doc.add_paragraph("Unrelated content here.")

    blk = _make_block([
        ("Line one in PDF", (0, 0, 100, 20)),
        ("Line two in PDF", (0, 25, 100, 45)),
    ])
    truth = _make_truth([_make_page(blocks=[blk])])

    res = paragraph_line_split.fix_paragraph_line_split(doc, truth, None)
    assert res["split"] == 0


def test_paragraph_line_split_already_has_newline_skipped():
    """docx 段落內 w:t 文字含字面 \\n（pdf2docx 罕見但可能） → 不處理。"""
    doc = Document()
    p = doc.add_paragraph()
    run = p.add_run()
    # 手動寫 w:t 內含字面 newline（python-docx add_run("\n") 會走 w:br，繞過）
    t = run._element.find(qn("w:t"))
    if t is None:
        t = run._element.makeelement(qn("w:t"), {})
        run._element.append(t)
    t.text = "Line A\nLine B"
    t.set(qn("xml:space"), "preserve")

    blk = _make_block([("Line A", (0, 0, 100, 20)), ("Line B", (0, 25, 100, 45))])
    truth = _make_truth([_make_page(blocks=[blk])])
    res = paragraph_line_split.fix_paragraph_line_split(doc, truth, None)
    assert res["split"] == 0


# --- table_cell_dedup_text ---------------------------------------------------

def test_table_cell_dedup_consecutive():
    """cell 內連續兩段都是 '技術服務' → 清掉第二段。"""
    doc = Document()
    tbl = doc.add_table(rows=1, cols=2)
    cell = tbl.rows[0].cells[0]
    cell.text = "技術服務"   # 第一段
    cell.add_paragraph("技術服務")  # 第二段 = 重複

    res = table_cell_dedup_text.fix_table_cell_dedup_text(doc, None, None)
    assert res["cleared_paragraphs"] == 1

    # 重新讀 — 兩段都還在，但第二段文字應該被清空
    paras = cell.paragraphs
    texts = [p.text for p in paras]
    assert texts[0] == "技術服務"
    assert texts[1] == ""


def test_table_cell_dedup_no_dups():
    doc = Document()
    tbl = doc.add_table(rows=1, cols=1)
    cell = tbl.rows[0].cells[0]
    cell.text = "First"
    cell.add_paragraph("Second")
    cell.add_paragraph("Third")
    res = table_cell_dedup_text.fix_table_cell_dedup_text(doc, None, None)
    assert res["cleared_paragraphs"] == 0


def test_table_cell_dedup_non_consecutive_not_touched():
    """重複但非連續 (A, B, A) → 不動。"""
    doc = Document()
    tbl = doc.add_table(rows=1, cols=1)
    cell = tbl.rows[0].cells[0]
    cell.text = "A"
    cell.add_paragraph("B")
    cell.add_paragraph("A")
    res = table_cell_dedup_text.fix_table_cell_dedup_text(doc, None, None)
    assert res["cleared_paragraphs"] == 0


# --- fake_table_remove 強化 --------------------------------------------------

def test_fake_table_heavy_empty_no_pdf_table():
    """2x3 表格只有 1 個 cell 有文字，且 PDF 無真表格 → 移除。"""
    doc = Document()
    tbl = doc.add_table(rows=2, cols=3)
    tbl.rows[0].cells[0].text = "Hi"

    # PDFTruth 無 rect drawing → no real table
    truth = _make_truth([_make_page()])
    res = fake_table_remove.fix_fake_table_remove(doc, truth, None)
    assert res["removed_tables"] == 1
    assert "heavy_empty_no_pdf_table" in res["reasons"]


def test_fake_table_heavy_empty_with_real_pdf_table_strict():
    """PDF 有真表格 → empty_ratio >= 0.85 + non_empty <= 3 才動。"""
    doc = Document()
    # 6 cell 中 1 個有文字 → empty_ratio = 5/6 ≈ 0.83，不滿足 >= 0.85 → 不動
    tbl = doc.add_table(rows=2, cols=3)
    tbl.rows[0].cells[0].text = "Hi"

    # 加一個大 rect drawing → PDF has real table
    drw = PDFDrawing(type="rect", bbox=(50, 50, 550, 700), page_num=0)
    truth = _make_truth([_make_page(drawings=[drw])])
    res = fake_table_remove.fix_fake_table_remove(doc, truth, None)
    assert res["removed_tables"] == 0


def test_fake_table_real_table_preserved():
    """正常表格（70% 非空）→ 不動。"""
    doc = Document()
    tbl = doc.add_table(rows=2, cols=3)
    for i, row in enumerate(tbl.rows):
        for j, c in enumerate(row.cells):
            if not (i == 1 and j == 2):  # 5/6 cells filled
                c.text = f"r{i}c{j}"
    res = fake_table_remove.fix_fake_table_remove(doc, _make_truth([_make_page()]), None)
    assert res["removed_tables"] == 0


# --- text_recovery -----------------------------------------------------------

def test_text_recovery_simple_short_block_added():
    """PDF 有 'TEST' block + docx 沒對應段落 → 補。"""
    doc = Document()
    doc.add_paragraph("First paragraph appears in PDF.")
    doc.add_paragraph("Second paragraph also in PDF.")

    blocks = [
        _make_block([("First paragraph appears in PDF.", (50, 50, 500, 70))]),
        _make_block([("Second paragraph also in PDF.", (50, 80, 500, 100))]),
        _make_block([("TEST", (50, 700, 100, 720))]),  # short block, missing
    ]
    truth = _make_truth([_make_page(blocks=blocks)])

    res = text_recovery.fix_text_recovery(doc, truth, None)
    assert res["recovered"] == 1
    texts = [p.text for p in doc.paragraphs]
    assert "TEST" in texts


def test_text_recovery_bail_when_page_missing_too_high():
    """整頁 ≥ 50% block 找不到 → 該頁 skip 不補。"""
    doc = Document()
    doc.add_paragraph("only one match.")

    blocks = [
        _make_block([("only one match.", (50, 50, 500, 70))]),
        _make_block([("missing 1", (50, 80, 200, 100))]),
        _make_block([("missing 2", (50, 110, 200, 130))]),
        _make_block([("missing 3", (50, 140, 200, 160))]),
    ]
    truth = _make_truth([_make_page(blocks=blocks)])
    res = text_recovery.fix_text_recovery(doc, truth, None)
    # miss_ratio = 3/4 = 0.75 >= 0.5 → bail
    assert res["recovered"] == 0


def test_text_recovery_skips_long_blocks():
    """漏抓的 block 太長 → 不補（容易位置錯）。"""
    doc = Document()
    doc.add_paragraph("anchor here for layout.")

    long_text = "A" * 150
    blocks = [
        _make_block([("anchor here for layout.", (50, 50, 500, 70))]),
        _make_block([(long_text, (50, 80, 500, 200))]),  # > 100 chars
    ]
    truth = _make_truth([_make_page(blocks=blocks)])
    res = text_recovery.fix_text_recovery(doc, truth, None)
    assert res["recovered"] == 0


# --- image_position_fix 補插重複 hash 圖 -------------------------------------

def _make_docx_with_image(width_pt: float, height_pt: float, hash_hex: str):
    """造一個 docx 含一張 inline image。手動建構 drawing element 含 a:blip + a:ext。"""
    from io import BytesIO
    from PIL import Image
    doc = Document()
    p = doc.add_paragraph()
    img_io = BytesIO()
    Image.new("RGB", (50, 50), color=(200, 200, 200)).save(img_io, format="PNG")
    img_io.seek(0)
    run = p.add_run()
    # docx 的 add_picture 走特定 image part；hash 對不上我們手寫的 PDFImage hash
    # 改用 InlineShapes API
    run.add_picture(img_io, width=None, height=None)
    return doc


def test_image_position_fix_inserts_duplicate():
    """PDF 有 2 張同 hash starburst，docx 只放了 1 張 → 補一張。"""
    # 構造 docx 含 1 inline image + 對 image_position_fix 傳一個 PDFTruth 有 2 張
    # 同 hash 的 PDFImage。
    # 因 docx hash 是 SHA1(blob)[:16]，我們直接 fake PDFTruth hash = docx hash。
    import hashlib
    from io import BytesIO
    from PIL import Image

    doc = Document()
    p1 = doc.add_paragraph()
    p1.add_run("intro")
    img_io = BytesIO()
    Image.new("RGB", (60, 60), color=(150, 150, 150)).save(img_io, format="PNG")
    img_bytes = img_io.getvalue()
    img_io.seek(0)
    p2 = doc.add_paragraph()
    p2.add_run().add_picture(img_io)

    # 取出 docx 內 image part 的 hash
    image_hash = ""
    for rel in doc.part.rels.values():
        if "image" in (rel.reltype or "").lower():
            blob = rel.target_part.blob
            image_hash = hashlib.sha1(blob, usedforsecurity=False).hexdigest()[:16]
            break
    assert image_hash, "test setup failed: docx image part not found"

    pdf_img_1 = PDFImage(bbox=(50, 100, 200, 250), page_num=0, xref=1,
                         width=60, height=60, image_hash=image_hash)
    pdf_img_2 = PDFImage(bbox=(300, 100, 450, 250), page_num=0, xref=2,
                         width=60, height=60, image_hash=image_hash)
    truth = _make_truth([_make_page(images=[pdf_img_1, pdf_img_2])])

    res = fix_image_position_fix(doc, truth, None)
    assert res["matched_by_hash"] == 1
    assert res["inserted_duplicates"] == 1


def test_image_position_fix_no_leftover_no_insert():
    """N 張 docx ↔ N 張 PDF 同 hash → 不補。"""
    import hashlib
    from io import BytesIO
    from PIL import Image

    doc = Document()
    img_io = BytesIO()
    Image.new("RGB", (40, 40), color=(100, 100, 100)).save(img_io, format="PNG")
    img_io.seek(0)
    doc.add_paragraph().add_run().add_picture(img_io)

    image_hash = ""
    for rel in doc.part.rels.values():
        if "image" in (rel.reltype or "").lower():
            blob = rel.target_part.blob
            image_hash = hashlib.sha1(blob, usedforsecurity=False).hexdigest()[:16]
            break

    pdf_img = PDFImage(bbox=(50, 50, 150, 150), page_num=0, xref=1,
                       width=40, height=40, image_hash=image_hash)
    truth = _make_truth([_make_page(images=[pdf_img])])
    res = fix_image_position_fix(doc, truth, None)
    assert res["inserted_duplicates"] == 0
